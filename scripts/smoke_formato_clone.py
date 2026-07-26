"""Dev smoke test: template clone-and-fill pipeline against real example documents.

Runs the REAL Gemini detection (`detectar_campos_plantilla`) on each file, fills
every detected campo with dummy values, writes the filled copy, and verifies:
  (a) every injected value appears at its target address,
  (b) table/paragraph/section (or sheet/dimension) structure is preserved,
  (c) docx headers/footers XML are unchanged.

Usage (from cashing-backend, venv with .env configured):
    .venv\\Scripts\\python.exe scripts\\smoke_formato_clone.py [files...]
"""

from __future__ import annotations

import asyncio
import sys
import time
import zipfile
from collections import Counter
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree as ET

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
sys.stdout.reconfigure(encoding="utf-8", errors="replace")  # type: ignore[union-attr]

OUT_DIR = Path(
    r"C:\Users\User\AppData\Local\Temp\claude"
    r"\C--Users-User-Documents-workspace-cashing"
    r"\a5c114ed-854a-4707-8fbf-832112afed04\scratchpad\smoke_formato"
)

CTX = Path(r"C:\Users\User\Documents\workspace\cashing\context")
DEFAULT_FILES = [
    CTX / "SYJ/CTO3/CUOTA#4/6. INFORME DE SUPERVISION 4.docx",
    CTX / "SYJ/CTO3/CUOTA#4/12. INFORME DE ACTIVIDADES 4.docx",
    CTX / "DAGMA/COEMPRESAR/CUOTA2/1. CUENTA DE COBRO.docx",
    CTX / "DAGMA/COEMPRESAR/CUOTA2/2. FORMATO PARA INFORME DE ACTIVIDADES COEMPRESAR.docx",
    CTX / "DAGMA/CUOTA #4/INFORME SUPERVISION CUOTA #4.docx",
    CTX / "SYJ/CTO3/CUOTA#4/3. DS-4161-2518.xlsx",
]

MONEY_FIELDS = {
    "contrato.valor_total",
    "contrato.valor_mensual",
    "cuota.valor",
    "computed.acumulado",
    "computed.saldo",
    "computed.valor_total_con_adiciones",
}
DATE_FIELDS = {"contrato.fecha_inicio", "contrato.fecha_fin", "cuota.fecha"}


def _squash(text: str) -> str:
    return " ".join(text.split())


def build_valores(campos: list[dict]) -> dict[str, str]:
    valores: dict[str, str] = {}
    for campo in campos:
        key = campo.get("campo") or ""
        if not key or key in valores:
            continue
        if key in MONEY_FIELDS:
            valores[key] = "$ 9.999.999"
        elif key in DATE_FIELDS:
            valores[key] = "31/12/2099"
        elif key == "computed.valor_letras":
            valores[key] = "NUEVE MILLONES NOVECIENTOS NOVENTA Y NUEVE MIL NOVECIENTOS NOVENTA Y NUEVE PESOS M/CTE"
        elif key == "computed.tipo_informe":
            valores[key] = "alternativa"
        elif key.startswith("justificacion."):
            n = key.split(".", 1)[1]
            valores[key] = f"Texto de prueba para la obligación {n}.\n\nSegundo párrafo de prueba."
        else:
            valores[key] = f"PRUEBA-{key.split('.', 1)[-1]}"
    return valores


def _node_text_docx(node) -> str:
    from docx.table import _Cell

    if isinstance(node, _Cell):
        return _squash(" ".join(p.text for p in node.paragraphs))
    return _squash(node.text)


def verify_docx(original: bytes, filled: bytes, campos: list[dict], valores: dict[str, str]) -> list[str]:
    """Return a list of failure strings (empty = all checks pass)."""
    from docx import Document

    from app.services.docx_clone_service import resolver_direccion

    failures: list[str] = []
    doc = Document(BytesIO(filled))

    # (a) injected values present at their addresses
    for campo in campos:
        key = campo.get("campo")
        if key not in valores:
            continue
        valor = valores[key]
        modo = campo.get("modo", "cell")
        node = resolver_direccion(doc, campo.get("direccion", ""))
        if node is None:
            failures.append(f"[inject] {key}: direccion '{campo.get('direccion')}' no resuelve en el documento relleno")
            continue
        if modo == "checkbox":
            if valor != "alternativa":
                continue
            alt = resolver_direccion(doc, campo.get("direccion_alternativa", ""))
            marker = campo.get("valor_alternativo") or campo.get("valor_ejemplo") or "X"
            if alt is None:
                failures.append(f"[inject] {key}: direccion_alternativa no resuelve")
            elif _squash(marker) not in _node_text_docx(alt):
                failures.append(f"[inject] {key}: marcador '{marker}' ausente en {campo.get('direccion_alternativa')}")
            continue
        if modo == "justificacion":
            text = _node_text_docx(node)
            for line in (ln for chunk in valor.split("\n\n") for ln in chunk.split("\n") if ln.strip()):
                if _squash(line) not in text:
                    failures.append(f"[inject] {key}: línea '{line[:40]}...' ausente en {campo.get('direccion')}")
            continue
        if _squash(valor) not in _node_text_docx(node):
            failures.append(f"[inject] {key} ({modo}): valor '{valor}' ausente en {campo.get('direccion')}")

    # (b) structure preserved
    orig = Document(BytesIO(original))
    for label, a, b in [
        ("tables", len(orig.tables), len(doc.tables)),
        ("paragraphs", len(orig.paragraphs), len(doc.paragraphs)),
        ("sections", len(orig.sections), len(doc.sections)),
    ]:
        if a != b:
            failures.append(f"[structure] {label}: original={a} filled={b}")

    # (c) headers/footers XML unchanged (byte-equal, or XML-canonical-equal since
    # python-docx re-serializes parts on save)
    with zipfile.ZipFile(BytesIO(original)) as zin, zipfile.ZipFile(BytesIO(filled)) as zout:
        def hf(z: zipfile.ZipFile) -> dict[str, bytes]:
            return {
                n: z.read(n)
                for n in z.namelist()
                if n.startswith("word/header") or n.startswith("word/footer")
            }

        a_parts, b_parts = hf(zin), hf(zout)
        if set(a_parts) != set(b_parts):
            failures.append(f"[hdrftr] part sets differ: {sorted(set(a_parts) ^ set(b_parts))}")
        for name in sorted(set(a_parts) & set(b_parts)):
            if a_parts[name] == b_parts[name]:
                continue
            try:
                if ET.canonicalize(a_parts[name].decode("utf-8")) != ET.canonicalize(b_parts[name].decode("utf-8")):
                    failures.append(f"[hdrftr] {name}: XML content changed")
            except Exception as exc:
                failures.append(f"[hdrftr] {name}: no comparable ({exc})")
    return failures


def verify_xlsx(original: bytes, filled: bytes, campos: list[dict], valores: dict[str, str]) -> list[str]:
    from openpyxl import load_workbook

    from app.services.xlsx_clone_service import _resolver_celda

    failures: list[str] = []
    wb = load_workbook(BytesIO(filled), data_only=False)

    for campo in campos:
        key = campo.get("campo")
        if key not in valores:
            continue
        cell = _resolver_celda(wb, campo.get("direccion", ""))
        if cell is None:
            failures.append(f"[inject] {key}: direccion '{campo.get('direccion')}' no resuelve en el workbook relleno")
        elif _squash(str(cell.value or "")) != _squash(valores[key]):
            failures.append(f"[inject] {key}: '{cell.value}' != '{valores[key]}' en {campo.get('direccion')}")

    orig = load_workbook(BytesIO(original), data_only=False)
    if orig.sheetnames != wb.sheetnames:
        failures.append(f"[structure] sheets: {orig.sheetnames} != {wb.sheetnames}")
    for name in orig.sheetnames:
        if name not in wb.sheetnames:
            continue
        a, b = orig[name], wb[name]
        if (a.max_row, a.max_column) != (b.max_row, b.max_column):
            failures.append(
                f"[structure] '{name}' dims: {a.max_row}x{a.max_column} != {b.max_row}x{b.max_column}"
            )
    return failures


async def smoke_file(path: Path) -> dict:
    from app.services.docx_clone_service import rellenar
    from app.services.requisito_inference_service import detectar_campos_plantilla
    from app.services.xlsx_clone_service import rellenar_xlsx

    formato = "xlsx" if path.suffix.lower() == ".xlsx" else "docx"
    original = path.read_bytes()

    t0 = time.perf_counter()
    campos, avisos = await detectar_campos_plantilla(original, formato=formato)
    latency = time.perf_counter() - t0

    by_scope = Counter((c.get("campo") or "?").split(".")[0] for c in campos)
    by_modo = Counter(c.get("modo", "?") for c in campos)
    print(f"\n=== {path.name} ({formato}) ===")
    print(f"  detectados: {len(campos)}  |  LLM {latency:.1f}s")
    print(f"  por scope: {dict(by_scope)}")
    print(f"  por modo:  {dict(by_modo)}")
    for c in campos:
        print(f"    {c.get('direccion'):<28} {c.get('modo'):<13} {c.get('campo'):<38} ej='{str(c.get('valor_ejemplo'))[:45]}'")
    if avisos:
        print(f"  avisos ({len(avisos)} descartados):")
        for a in avisos:
            print(f"    - {a}")

    result = {
        "file": path.name,
        "formato": formato,
        "campos": len(campos),
        "scopes": dict(by_scope),
        "avisos": len(avisos),
        "latency": latency,
        "fill": "SKIP",
        "failures": [],
    }
    if not campos:
        result["fill"] = "FAIL (0 campos)"
        return result

    valores = build_valores(campos)
    try:
        if formato == "xlsx":
            filled = rellenar_xlsx(original, campos, valores)
        else:
            filled = rellenar(original, campos, valores)
    except Exception as exc:
        result["fill"] = f"FAIL ({type(exc).__name__}: {exc})"
        return result

    out_path = OUT_DIR / f"{path.stem}.FILLED{path.suffix}"
    out_path.write_bytes(filled)
    result["fill"] = "OK"
    result["out"] = str(out_path)

    verify = verify_xlsx if formato == "xlsx" else verify_docx
    result["failures"] = verify(original, filled, campos, valores)
    for f in result["failures"]:
        print(f"  FALLO {f}")
    return result


async def main(paths: list[Path]) -> int:
    from app.core.config import settings

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"LLM_EXTRACTION_MODEL = {settings.LLM_EXTRACTION_MODEL or settings.LLM_DEFAULT_MODEL}")
    print(f"Salida: {OUT_DIR}")

    results = []
    for path in paths:
        if not path.exists():
            print(f"\n=== {path} ===\n  NO EXISTE")
            results.append({"file": path.name, "formato": "?", "campos": 0, "scopes": {},
                            "avisos": 0, "latency": 0.0, "fill": "FAIL (missing)", "failures": []})
            continue
        results.append(await smoke_file(path))

    print("\n" + "=" * 100)
    print(f"{'archivo':<55} {'campos':>6} {'avisos':>6} {'LLM s':>6}  {'fill':<6} resultado")
    print("-" * 100)
    ok = True
    for r in results:
        status = "PASS" if r["fill"] == "OK" and not r["failures"] else f"FAIL ({len(r['failures'])} checks)" if r["fill"] == "OK" else "FAIL"
        if status != "PASS":
            ok = False
        print(f"{r['file']:<55} {r['campos']:>6} {r['avisos']:>6} {r['latency']:>6.1f}  {r['fill']:<6} {status}")
    print(f"\nLatencia LLM total: {sum(r['latency'] for r in results):.1f}s")
    return 0 if ok else 1


if __name__ == "__main__":
    args = [Path(a) for a in sys.argv[1:]] or DEFAULT_FILES
    sys.exit(asyncio.run(main(args)))
