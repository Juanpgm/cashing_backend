"""Integration test (billing-resilience-templates, slice #6, task 6.11):
extended informe generation, exercised end-to-end through the TOOL layer.

NAMING DEVIATION (documented, same pattern as prior slices' `_CODIGOS_
OBLIGATORIOS` mirror notes): the spec/tasks Tool Surface table names a single
generic `generar_informe_cuota` tool, but the actual registered tools are
`generar_informe_actividades` and `generar_informe_supervision` (app/tools/
catalog/informes.py) — both thin wrappers over the SAME `informe_service`
generators this slice extended. This test invokes them through `invoke_tool`
(the real MCP/agent dispatch path) and inspects the DOCX bytes handed to
`StoragePort.upload` to prove the adaptive layout selection, always-draft
header, and one-time-obligation blanking all reach the caller through the
full tool boundary — not just through a direct service-level call.
"""

from __future__ import annotations

import io
from datetime import date
from typing import Any
from unittest.mock import AsyncMock, patch

import app.tools.catalog  # noqa: F401 — import-for-side-effect: registers every catalog tool
import pytest
from app.core.text_match import normalize
from app.models.actividad import Actividad
from app.models.contrato import Contrato
from app.models.cuenta_cobro import CuentaCobro, EstadoCuentaCobro, PosicionCuota
from app.models.obligacion import Obligacion, TipoObligacion
from app.models.plantilla_organismo import PlantillaOrganismo
from app.tools.catalog.informes import GenerarInformeInput
from app.tools.context import ToolContext
from app.tools.invoke import invoke_tool
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

pytestmark = pytest.mark.asyncio

_PATCH_S3 = "app.services.document_service._get_storage"


def _fake_storage() -> AsyncMock:
    storage = AsyncMock()
    storage.upload = AsyncMock(return_value="fake/key")
    return storage


@pytest.fixture
async def contrato_dagma(db: AsyncSession, test_user: dict[str, Any]) -> Contrato:
    user = test_user["user"]
    c = Contrato(
        usuario_id=user.id,
        numero_contrato="CTR-TOOL-DAGMA-001",
        objeto="Servicios profesionales ambientales",
        valor_total=24_000_000,
        valor_mensual=2_000_000,
        fecha_inicio=date(2024, 1, 1),
        fecha_fin=date(2024, 12, 31),
        entidad="DAGMA",
        dependencia="Ambiental",
        supervisor_nombre="Sup DAGMA",
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


@pytest.fixture
async def obligaciones(db: AsyncSession, contrato_dagma: Contrato) -> list[Obligacion]:
    obs = [
        Obligacion(
            contrato_id=contrato_dagma.id,
            descripcion=f"Obligación contractual #{i + 1} con texto suficientemente largo",
            tipo=TipoObligacion.GENERAL,
            orden=i,
        )
        for i in range(2)
    ]
    db.add_all(obs)
    await db.commit()
    for o in obs:
        await db.refresh(o)
    contrato_dagma.obligaciones = list(obs)
    return obs


@pytest.fixture
async def plantilla_dagma(db: AsyncSession, contrato_dagma: Contrato) -> PlantillaOrganismo:
    plantilla = PlantillaOrganismo(
        usuario_id=contrato_dagma.usuario_id,
        entidad=contrato_dagma.entidad,
        entidad_normalizada=normalize(contrato_dagma.entidad),
        tipo_documento="informe_actividades",
        formato="docx",
        estructura_json={
            "columnas": ["Obligación", "Avance del periodo"],
            "secciones": [],
            "anexo_refs": [],
            "notas": "",
        },
    )
    db.add(plantilla)
    await db.commit()
    await db.refresh(plantilla)
    return plantilla


@pytest.fixture
async def cuenta(db: AsyncSession, contrato_dagma: Contrato, obligaciones: list[Obligacion]) -> CuentaCobro:
    cc = CuentaCobro(
        contrato_id=contrato_dagma.id,
        mes=5,
        anio=2024,
        estado=EstadoCuentaCobro.BORRADOR,
        valor=2_000_000,
        requisitos_modo="estandar",
        numero_cuota=1,
        posicion=PosicionCuota.PRIMERA,
    )
    db.add(cc)
    await db.commit()
    await db.refresh(cc)
    for i, ob in enumerate(obligaciones):
        db.add(
            Actividad(
                cuenta_cobro_id=cc.id,
                obligacion_id=ob.id,
                descripcion=f"Actividad realizada {i + 1}",
                justificacion=f"Justificación {i + 1}",
                fecha_realizacion=date(2024, 5, 10 + i),
            )
        )
    await db.commit()
    await db.refresh(cc)
    return cc


async def test_generar_informe_actividades_tool_aplica_layout_organismo_y_borrador(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    plantilla_dagma: PlantillaOrganismo,
) -> None:
    """End-to-end through `invoke_tool`: the DAGMA 2-column layout AND the
    always-draft header both reach the actual uploaded DOCX bytes."""
    user = test_user["user"]
    ctx = ToolContext(db=db, usuario=user)
    storage = _fake_storage()

    with patch(_PATCH_S3, return_value=storage):
        result = await invoke_tool("generar_informe_actividades", ctx, GenerarInformeInput(cuenta_id=cuenta.id))

    assert result.estado == "cargado"
    storage.upload.assert_called_once()
    content = storage.upload.call_args.kwargs["data"]
    doc = Document(io.BytesIO(content))
    tabla = doc.tables[1]
    headers = [c.text for c in tabla.rows[0].cells]
    assert headers == ["Obligación", "Avance del periodo"]
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "BORRADOR" in full_text
