"""Tests for third-person rewriting of the supervisión report's activity texts.

Covers:
- `parse_tercera_persona` parsing contract (round-trip, count mismatch, tolerant of prose).
- `_convertir_actividades_tercera_persona` batching + fail-open behavior.
- `_add_actividades_table` overrides rendering.
- Integration: supervisión report uses third-person text, actividades report is
  unaffected (still first person), and supervisión fails open (still generates) when
  the LLM raises.
"""

from __future__ import annotations

import io
import uuid
from datetime import date
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from app.agent.prompts.supervision_tercera_persona import parse_tercera_persona
from app.models.actividad import Actividad
from app.models.contrato import Contrato
from app.models.cuenta_cobro import CuentaCobro, EstadoCuentaCobro
from app.models.obligacion import Obligacion, TipoObligacion
from app.services import informe_service
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

# NOTE: no module-level `pytestmark = pytest.mark.asyncio` here — this file mixes sync
# (parser) and async (service) tests, and `asyncio_mode = "auto"` already runs async
# tests correctly without the mark.


# ── parse_tercera_persona ────────────────────────────────────────────────────


def test_parse_tercera_persona_round_trip() -> None:
    content = "1| El contratista elaboró el informe mensual.\n2| El contratista asistió a la reunión de seguimiento."
    result = parse_tercera_persona(content, expected=2)
    assert result == [
        "El contratista elaboró el informe mensual.",
        "El contratista asistió a la reunión de seguimiento.",
    ]


def test_parse_tercera_persona_count_mismatch_returns_none() -> None:
    content = "1| Único texto reescrito."
    assert parse_tercera_persona(content, expected=2) is None


def test_parse_tercera_persona_tolerates_surrounding_prose() -> None:
    content = (
        "Aquí están los textos reescritos:\n\n"
        "1| El contratista elaboró el informe mensual.\n"
        "2| El contratista asistió a la reunión de seguimiento.\n\n"
        "Espero que sea útil."
    )
    result = parse_tercera_persona(content, expected=2)
    assert result == [
        "El contratista elaboró el informe mensual.",
        "El contratista asistió a la reunión de seguimiento.",
    ]


def test_parse_tercera_persona_empty_expected_returns_none() -> None:
    assert parse_tercera_persona("1| algo", expected=0) is None


# ── _convertir_actividades_tercera_persona ──────────────────────────────────


def _make_actividad(descripcion: str, justificacion: str | None) -> Actividad:
    act = Actividad(
        id=uuid.uuid4(),
        cuenta_cobro_id=uuid.uuid4(),
        descripcion=descripcion,
        justificacion=justificacion,
        fecha_realizacion=date(2024, 5, 10),
    )
    return act


async def test_convertir_actividades_tercera_persona_ok() -> None:
    act1 = _make_actividad("Elaboré el informe mensual.", "Cumplo con la obligación de reportar.")
    act2 = _make_actividad("Asistí a la reunión de seguimiento.", "Doy cumplimiento a la obligación de acompañamiento.")

    fake_resp = MagicMock()
    fake_resp.content = (
        "1| El contratista elaboró el informe mensual.\n"
        "2| El contratista cumple con la obligación de reportar.\n"
        "3| El contratista asistió a la reunión de seguimiento.\n"
        "4| El contratista da cumplimiento a la obligación de acompañamiento."
    )
    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(return_value=fake_resp)

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        overrides = await informe_service._convertir_actividades_tercera_persona([act1, act2])

    assert overrides[act1.id] == (
        "El contratista elaboró el informe mensual.",
        "El contratista cumple con la obligación de reportar.",
    )
    assert overrides[act2.id] == (
        "El contratista asistió a la reunión de seguimiento.",
        "El contratista da cumplimiento a la obligación de acompañamiento.",
    )


async def test_convertir_actividades_tercera_persona_llm_error_fails_open() -> None:
    act1 = _make_actividad("Elaboré el informe mensual.", "Cumplo con la obligación de reportar.")

    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(side_effect=RuntimeError("llm down"))

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        overrides = await informe_service._convertir_actividades_tercera_persona([act1])

    assert overrides == {}


async def test_convertir_actividades_tercera_persona_count_mismatch_fails_open() -> None:
    act1 = _make_actividad("Elaboré el informe mensual.", "Cumplo con la obligación de reportar.")

    fake_resp = MagicMock()
    fake_resp.content = "1| Único texto reescrito, faltan los demás."
    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(return_value=fake_resp)

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        overrides = await informe_service._convertir_actividades_tercera_persona([act1])

    assert overrides == {}


async def test_convertir_actividades_tercera_persona_empty_batch_skips_llm() -> None:
    act1 = _make_actividad("", None)

    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock()

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        overrides = await informe_service._convertir_actividades_tercera_persona([act1])

    assert overrides == {}
    mock_llm.complete.assert_not_called()


# ── _add_actividades_table overrides ────────────────────────────────────────


async def test_add_actividades_table_uses_overrides_selectively() -> None:
    act1 = _make_actividad("Actividad original 1", "Justificación original 1")
    act2 = _make_actividad("Actividad original 2", "Justificación original 2")

    overrides = {act1.id: ("Actividad tercera persona 1", "Justificación tercera persona 1")}

    doc = Document()
    informe_service._add_actividades_table(doc, [act1, act2], {}, overrides=overrides)

    table = doc.tables[0]
    row1 = table.rows[1]
    row2 = table.rows[2]

    assert row1.cells[2].text == "Actividad tercera persona 1"
    assert row1.cells[3].text == "Justificación tercera persona 1"
    assert row2.cells[2].text == "Actividad original 2"
    assert row2.cells[3].text == "Justificación original 2"


# ── Integration fixtures ─────────────────────────────────────────────────────


@pytest.fixture
async def contrato(db: AsyncSession, test_user: dict[str, Any]) -> Contrato:
    user = test_user["user"]
    c = Contrato(
        usuario_id=user.id,
        numero_contrato="CTR-3P-001",
        objeto="Servicios profesionales de desarrollo de software",
        valor_total=36_000_000,
        valor_mensual=3_000_000,
        fecha_inicio=date(2024, 1, 1),
        fecha_fin=date(2024, 12, 31),
        entidad="Alcaldía",
        dependencia="TI",
        supervisor_nombre="Carlos Supervisor",
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


@pytest.fixture
async def obligaciones(db: AsyncSession, contrato: Contrato) -> list[Obligacion]:
    obs = [
        Obligacion(
            contrato_id=contrato.id,
            descripcion=f"Obligación contractual #{i + 1} con un texto razonablemente largo",
            tipo=TipoObligacion.GENERAL,
            orden=i,
        )
        for i in range(2)
    ]
    db.add_all(obs)
    await db.commit()
    for o in obs:
        await db.refresh(o)
    contrato.obligaciones = list(obs)
    return obs


@pytest.fixture
async def cuenta(db: AsyncSession, contrato: Contrato, obligaciones: list[Obligacion]) -> CuentaCobro:
    cc = CuentaCobro(
        contrato_id=contrato.id,
        mes=5,
        anio=2024,
        estado=EstadoCuentaCobro.BORRADOR,
        valor=3_000_000,
    )
    db.add(cc)
    await db.commit()
    await db.refresh(cc)

    for i, ob in enumerate(obligaciones):
        act = Actividad(
            cuenta_cobro_id=cc.id,
            obligacion_id=ob.id,
            descripcion=f"Elaboré y entregué el producto {i + 1} del período.",
            justificacion=f"Cumplo con la obligación número {i + 1} mediante este entregable.",
            fecha_realizacion=date(2024, 5, 10 + i),
        )
        db.add(act)
    await db.commit()
    await db.refresh(cc)
    return cc


def _fake_tercera_persona_response(n: int) -> MagicMock:
    resp = MagicMock()
    resp.content = "\n".join(f"{i + 1}| El contratista realizó la acción reescrita número {i + 1}." for i in range(n))
    return resp


# ── Integration: supervisión in third person, actividades unaffected ───────


async def test_informe_supervision_usa_tercera_persona(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    user = test_user["user"]

    mock_llm = AsyncMock()
    # 2 activities x 2 columns (descripcion + justificacion) = 4 texts in the batch.
    mock_llm.complete = AsyncMock(return_value=_fake_tercera_persona_response(4))

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        content, filename = await informe_service.generar_informe_supervision_docx(db, user.id, cuenta.id)

    assert filename.startswith("informe-supervision-")
    doc = Document(io.BytesIO(content))
    table_texts = [cell.text for t in doc.tables for r in t.rows for cell in r.cells]
    assert any("El contratista realizó la acción reescrita número" in tx for tx in table_texts)
    # Original first-person text must NOT leak into the supervisión report.
    assert not any("Elaboré y entregué el producto" in tx for tx in table_texts)


async def test_informe_actividades_no_afectado_por_tercera_persona(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    """Actividades report must stay first person, even with third-person conversion
    wired into the supervisión path — proving no cross-contamination."""
    user = test_user["user"]

    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(return_value=_fake_tercera_persona_response(4))

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        content, _filename = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta.id)

    doc = Document(io.BytesIO(content))
    table_texts = [cell.text for t in doc.tables for r in t.rows for cell in r.cells]
    assert any("Elaboré y entregué el producto" in tx for tx in table_texts)
    assert not any("El contratista realizó la acción reescrita número" in tx for tx in table_texts)
    # The LLM must not even be invoked for the actividades report.
    mock_llm.complete.assert_not_called()


async def test_informe_supervision_fail_open_llm_error(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    """When the LLM raises, supervisión still generates successfully using the
    original first-person text (fail-open contract)."""
    user = test_user["user"]

    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(side_effect=RuntimeError("ollama down"))

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        content, filename = await informe_service.generar_informe_supervision_docx(db, user.id, cuenta.id)

    assert filename.startswith("informe-supervision-")
    doc = Document(io.BytesIO(content))
    table_texts = [cell.text for t in doc.tables for r in t.rows for cell in r.cells]
    assert any("Elaboré y entregué el producto" in tx for tx in table_texts)
