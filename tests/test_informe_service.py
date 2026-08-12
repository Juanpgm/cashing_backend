"""Tests for informe_service (DOCX + ZIP generators).

The ZIP evidencias tests (billing-resilience-templates, slice #2) use a SYNTHETIC
real-leak-shaped corpus to exercise the mandatory secret scan — modeled on the
SHAPE of the real leak (a Postgres/Neon connection string + a Neon-style API key)
with entirely FAKE values. Never copy real credential values into this file.
"""

from __future__ import annotations

import io
import uuid
import zipfile
from datetime import date
from typing import Any
from unittest.mock import AsyncMock, MagicMock, patch

import pytest
from app.core.config import settings
from app.core.exceptions import ForbiddenError, ValidationError
from app.core.text_match import normalize
from app.models.actividad import Actividad
from app.models.contrato import Contrato
from app.models.cuenta_cobro import CuentaCobro, EstadoCuentaCobro, PosicionCuota
from app.models.evidencia import Evidencia
from app.models.evidencia_obligacion import EstadoEnlace, EvidenciaObligacion, FuenteEnlace
from app.models.obligacion import Obligacion, TipoObligacion
from app.models.plantilla_organismo import PlantillaOrganismo
from app.services import informe_service
from app.services.informe_constants import SENTINEL_SIN_EVIDENCIAS
from docx import Document
from httpx import AsyncClient
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

pytestmark = pytest.mark.asyncio

_LEAK_CORPUS_TEXTO = (
    "Notas internas\n"
    "DATABASE_URL=postgresql://fake_user:fake_pass_123@ep-fake-000000"
    ".us-east-1.aws.neon.tech/neondb\n"
    "NEON_API_KEY=napi_a8K3mXpQ9rZtL2wVbN7cJhF4sYdE1oIu\n"
)


def _fake_storage(download_map: dict[str, bytes] | None = None) -> AsyncMock:
    """AsyncMock StoragePort — `download` resolves by key from `download_map`,
    defaulting to an empty-but-nonempty placeholder for unmapped keys."""
    download_map = download_map or {}
    storage = AsyncMock()

    async def _download(key: str) -> bytes:
        return download_map.get(key, b"contenido de evidencia de prueba")

    storage.download = AsyncMock(side_effect=_download)
    storage.upload = AsyncMock(return_value="fake/key")
    return storage


# ── Fixtures ────────────────────────────────────────────────────────────────


@pytest.fixture
async def contrato(db: AsyncSession, test_user: dict[str, Any]) -> Contrato:
    user = test_user["user"]
    c = Contrato(
        usuario_id=user.id,
        numero_contrato="CTR-INF-001",
        objeto="Servicios profesionales de desarrollo de software",
        valor_total=36_000_000,
        valor_mensual=3_000_000,
        fecha_inicio=date(2024, 1, 1),
        fecha_fin=date(2024, 12, 31),
        entidad="Alcaldía",
        dependencia="TI",
        supervisor_nombre="Carlos Supervisor",
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


@pytest.fixture
async def obligaciones(db: AsyncSession, contrato: Contrato) -> list[Obligacion]:
    obs = [
        Obligacion(
            contrato_id=contrato.id,
            descripcion=f"Obligación contractual #{i + 1} con un texto razonablemente largo",
            tipo=TipoObligacion.GENERAL,
            orden=i,
        )
        for i in range(3)
    ]
    db.add_all(obs)
    await db.commit()
    for o in obs:
        await db.refresh(o)
    # The contrato was loaded with its `obligaciones` collection empty (lazy
    # selectin). Manually attach the new obligations so subsequent loads in
    # the same async session see them.
    contrato.obligaciones = list(obs)
    return obs


@pytest.fixture
async def cuenta(db: AsyncSession, contrato: Contrato, obligaciones: list[Obligacion]) -> CuentaCobro:
    cc = CuentaCobro(
        contrato_id=contrato.id,
        mes=5,
        anio=2024,
        estado=EstadoCuentaCobro.BORRADOR,
        valor=3_000_000,
    )
    db.add(cc)
    await db.commit()
    await db.refresh(cc)

    for i, ob in enumerate(obligaciones):
        act = Actividad(
            cuenta_cobro_id=cc.id,
            obligacion_id=ob.id,
            descripcion=f"Actividad realizada {i + 1}",
            justificacion=f"Justificación detallada {i + 1}",
            fecha_realizacion=date(2024, 5, 10 + i),
        )
        db.add(act)
    await db.commit()
    await db.refresh(cc)
    return cc


# ── Informe actividades ────────────────────────────────────────────────────


async def test_informe_actividades_genera_docx_valido(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    user = test_user["user"]
    content, filename = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta.id)
    assert filename.endswith(".docx")
    assert filename.startswith("informe-actividades-")
    assert len(content) > 1000  # not empty
    # Parse it back to verify it's a valid docx
    doc = Document(io.BytesIO(content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Informe de actividades" in full_text
    assert "CTR-INF-001" in full_text or any(
        "CTR-INF-001" in cell.text for t in doc.tables for r in t.rows for cell in r.cells
    )


# ── es_borrador machine-readable field (slice #7, task 7.5d) ────────────────


async def test_es_borrador_is_a_public_constant_true() -> None:
    """`ES_BORRADOR` is the machine-readable counterpart of `_BORRADOR_HEADER` —
    every generated informe is currently ALWAYS a draft (no approval workflow)."""
    assert informe_service.ES_BORRADOR is True


async def test_descargar_informe_actividades_incluye_header_es_borrador(
    client: AsyncClient, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    resp = await client.get(f"/api/v1/cuentas-cobro/{cuenta.id}/informe-actividades.docx", headers=test_user["headers"])
    assert resp.status_code == 200, resp.text
    assert resp.headers["X-Es-Borrador"] == "true"


async def test_descargar_informe_supervision_incluye_header_es_borrador(
    client: AsyncClient, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    resp = await client.get(f"/api/v1/cuentas-cobro/{cuenta.id}/informe-supervision.docx", headers=test_user["headers"])
    assert resp.status_code == 200, resp.text
    assert resp.headers["X-Es-Borrador"] == "true"


async def test_informe_actividades_sin_actividades_falla(
    db: AsyncSession, test_user: dict[str, Any], contrato: Contrato
) -> None:
    user = test_user["user"]
    cc = CuentaCobro(
        contrato_id=contrato.id,
        mes=6,
        anio=2024,
        valor=3_000_000,
        estado=EstadoCuentaCobro.BORRADOR,
    )
    db.add(cc)
    await db.commit()
    await db.refresh(cc)
    with pytest.raises(ValidationError):
        await informe_service.generar_informe_actividades_docx(db, user.id, cc.id)


# ── Informe supervisión ────────────────────────────────────────────────────


async def test_informe_supervision_genera_docx_valido(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    user = test_user["user"]
    content, filename = await informe_service.generar_informe_supervision_docx(db, user.id, cuenta.id)
    assert filename.endswith(".docx")
    assert filename.startswith("informe-supervision-")
    doc = Document(io.BytesIO(content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "supervisi" in full_text.lower()
    table_texts = [cell.text for t in doc.tables for r in t.rows for cell in r.cells]
    assert any("Carlos Supervisor" in tx for tx in table_texts)


# ── ZIP evidencias ─────────────────────────────────────────────────────────


async def test_zip_evidencias_estructura(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A3: an obligación folder carries ONLY real evidence files — no per-folder
    LEEME.txt anymore (the root LEEME.txt manifest is the only text file)."""
    user = test_user["user"]
    res = await db.execute(
        select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).order_by(Actividad.fecha_realizacion.asc())
    )
    act = res.scalars().first()
    assert act is not None
    ev = Evidencia(
        actividad_id=act.id,
        storage_key=f"evidencias/{act.id}/soporte.pdf",
        nombre_archivo="soporte.pdf",
        tipo_archivo="application/pdf",
        tamano_bytes=10,
    )
    db.add(ev)
    await db.commit()
    # `Actividad.evidencias` is `lazy="selectin"` — the query above already loaded
    # (and cached, empty) that collection on this same session's identity-mapped
    # `act`. Expire just that attribute so generar_zip_evidencias's own query
    # re-fetches it instead of reusing the stale pre-evidencia snapshot.
    db.expire(act, ["evidencias"])
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())

    content, filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)
    assert filename.endswith(".zip")
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        # Root readme
        assert "LEEME.txt" in names
        # No per-folder LEEME anymore (A3) — folders carry only the real files.
        leemes = [n for n in names if n.endswith("LEEME.txt") and "/" in n]
        assert leemes == []
        archivo = next(n for n in names if n.startswith("01_") and n.endswith("soporte.pdf"))
        assert zf.read(archivo)


async def test_ownership_otro_usuario_falla(db: AsyncSession, cuenta: CuentaCobro) -> None:
    fake_user_id = uuid.uuid4()
    with pytest.raises(ForbiddenError):
        await informe_service.generar_informe_actividades_docx(db, fake_user_id, cuenta.id)


# ── Packager hardening (billing-resilience-templates, slice #2) ────────────


@pytest.fixture
async def actividad_con_evidencia(db: AsyncSession, cuenta: CuentaCobro, obligaciones: list[Obligacion]) -> Actividad:
    """First actividad of `cuenta` (linked to the first obligación) plus one
    uploaded (storage-backed) evidencia."""
    res = await db.execute(
        select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).order_by(Actividad.fecha_realizacion.asc())
    )
    act = res.scalars().first()
    assert act is not None
    ev = Evidencia(
        actividad_id=act.id,
        storage_key=f"evidencias/{act.id}/foto.jpg",
        nombre_archivo="foto.jpg",
        tipo_archivo="image/jpeg",
        tamano_bytes=12,
    )
    db.add(ev)
    await db.commit()
    await db.refresh(act)
    return act


async def test_zip_evidencias_usa_bytes_reales_de_storage(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    actividad_con_evidencia: Actividad,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    user = test_user["user"]
    ev = actividad_con_evidencia.evidencias[0]
    contenido_real = b"contenido-binario-real-de-la-foto"
    storage = _fake_storage({ev.storage_key: contenido_real})
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        arcnames = [n for n in zf.namelist() if "evidencias/" in n]
        assert arcnames, "expected at least one real-bytes evidencia entry in the zip"
        assert zf.read(arcnames[0]) == contenido_real
    storage.download.assert_any_call(ev.storage_key)


async def test_zip_empaca_evidencia_link_confirmado_en_carpeta_de_obligacion(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    obligaciones: list[Obligacion],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Packaging counterpart of the H13 gate fix (live-reproduced 2026-08-12):
    an evidencia attached to the "sin clasificar" stub (obligacion_id=None) but
    CONFIRMED-linked to an obligación must ship its bytes under that obligation's
    folder — before the fix the folder carried only its LEEME."""
    user = test_user["user"]
    stub = Actividad(cuenta_cobro_id=cuenta.id, obligacion_id=None, descripcion="stub sin clasificar")
    db.add(stub)
    await db.flush()
    ev = Evidencia(
        actividad_id=stub.id,
        storage_key=f"evidencias/{stub.id}/soporte.pdf",
        nombre_archivo="soporte.pdf",
        tipo_archivo="application/pdf",
        tamano_bytes=10,
    )
    db.add(ev)
    await db.flush()
    db.add(
        EvidenciaObligacion(
            evidencia_id=ev.id,
            obligacion_id=obligaciones[0].id,
            status=EstadoEnlace.CONFIRMED.value,
            source=FuenteEnlace.USER.value,
        )
    )
    await db.commit()

    contenido_real = b"bytes-de-la-evidencia-linkeada"
    storage = _fake_storage({ev.storage_key: contenido_real})
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        matches = [n for n in zf.namelist() if n.startswith("01_") and n.endswith("soporte.pdf")]
        assert matches, "link-confirmed evidencia must land under its obligation folder"
        assert zf.read(matches[0]) == contenido_real
        # A3: no per-folder LEEME anymore — the folder is just the real file.
        leemes = [n for n in zf.namelist() if n.startswith("01_") and n.endswith("LEEME.txt")]
        assert leemes == []


async def test_zip_evidencias_incluye_bytes_subidos_por_el_seam_real(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    tmp_path: Any,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Regression for the evidencia-bytes bucket mismatch: uploads a real file
    through `evidencia_service.subir_evidencias` with the SAME storage seam the
    upload endpoints use (`get_storage(settings.S3_BUCKET_PDFS)`, mirroring
    `api.deps.get_pdf_storage`), then calls `generar_zip_evidencias` with NO
    storage mock at all — the real `get_evidencia_storage()` must resolve to
    the same bucket, or the file silently goes missing from the zip (the exact
    bug: the packager used to read `S3_BUCKET_EVIDENCIAS`, a bucket the upload
    endpoints never write to).

    `test_zip_evidencias_usa_bytes_reales_de_storage` above could NOT have
    caught this: it monkeypatches `_get_storage` with a bucket-agnostic fake
    that returns the same object regardless of which bucket string is passed.
    """
    from app.adapters.storage import get_storage
    from app.services import evidencia_service

    monkeypatch.setattr(settings, "STORAGE_PROVIDER", "local")
    monkeypatch.setattr(settings, "LOCAL_STORAGE_PATH", str(tmp_path))

    user = test_user["user"]
    res = await db.execute(
        select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).order_by(Actividad.fecha_realizacion.asc())
    )
    act = res.scalars().first()
    assert act is not None

    # `.bin` has no registered magic-byte signature (`_EXT_TO_MIME` in
    # `file_validation.py`) — evidencia accepts any extension outside the
    # blocklist, so arbitrary bytes are valid here without faking a real
    # JPEG header.
    contenido_real = b"contenido-binario-real-subido-por-el-endpoint"
    upload_storage = get_storage(settings.S3_BUCKET_PDFS)
    subidas = await evidencia_service.subir_evidencias(
        db=db,
        storage=upload_storage,
        usuario_id=user.id,
        actividad_id=act.id,
        archivos=[("foto.bin", "application/octet-stream", contenido_real)],
    )
    storage_key = subidas[0].storage_key

    # `Actividad.evidencias` is `lazy="selectin"` — the query above already
    # loaded (and cached, empty) that collection on this same session's
    # identity-mapped `act`. Expire just that attribute so
    # `generar_zip_evidencias`'s own query re-fetches it instead of reusing
    # the stale pre-upload snapshot (a same-session-reuse test artifact; a
    # real request always gets a fresh session per call). `expire_all()`
    # would be too broad here — it forces implicit lazy-loads outside an
    # awaited context elsewhere in the call chain (`MissingGreenlet`).
    db.expire(act, ["evidencias"])

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        arcnames = [n for n in zf.namelist() if "evidencias/" in n and "foto.bin" in n]
        assert arcnames, f"uploaded evidencia {storage_key} missing from zip — packager read the wrong bucket"
        assert zf.read(arcnames[0]) == contenido_real


async def test_resolver_estructura_organismo_returns_none_when_not_ingested(
    db: AsyncSession, contrato: Contrato
) -> None:
    """No `PlantillaOrganismo` has been ingested for this organism — the packager
    falls back to the default numbered folder structure (billing-resilience-
    templates, slice #5, task 5.15 — completes the slice #2 stub)."""
    resultado = await informe_service._resolver_estructura_organismo(db, contrato)
    assert resultado is None


async def test_resolver_estructura_organismo_returns_real_lookup_when_ingested(
    db: AsyncSession, contrato: Contrato
) -> None:
    """Once a template has been ingested for the contract's organism, the real
    lookup (`requisito_inference_service.obtener_plantilla_organismo`) replaces
    the slice #2 stub."""
    from app.core.text_match import normalize
    from app.models.plantilla_organismo import PlantillaOrganismo

    plantilla = PlantillaOrganismo(
        usuario_id=contrato.usuario_id,
        entidad=contrato.entidad,
        entidad_normalizada=normalize(contrato.entidad),
        tipo_documento="informe_actividades",
        formato="docx",
        estructura_json={
            "columnas": ["obligación", "avance del periodo", "evidencia"],
            "secciones": [],
            "anexo_refs": ["Ver Anexo: Carpeta /5. EVIDENCIAS/A1"],
            "notas": "",
        },
    )
    db.add(plantilla)
    await db.commit()

    resultado = await informe_service._resolver_estructura_organismo(db, contrato)
    assert resultado is not None
    assert resultado.estructura_json["anexo_refs"] == ["Ver Anexo: Carpeta /5. EVIDENCIAS/A1"]


async def test_zip_evidencias_uses_default_numbered_folders_when_no_organismo_template(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    # A folder only renders once it has a real evidencia file to carry (A3).
    res = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).limit(1))
    act = res.scalars().first()
    assert act is not None
    db.add(
        Evidencia(
            actividad_id=act.id,
            storage_key=f"evidencias/{act.id}/soporte.pdf",
            nombre_archivo="soporte.pdf",
            tipo_archivo="application/pdf",
            tamano_bytes=10,
        )
    )
    await db.commit()
    # See test_zip_evidencias_estructura for why this expire is required.
    db.expire(act, ["evidencias"])
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        folders = {n.split("/")[0] for n in zf.namelist() if "/" in n}
    assert any(f.startswith(("01_", "02_", "03_")) for f in folders)
    assert not any(f.startswith("A1_") for f in folders)


async def test_zip_evidencias_uses_anexo_style_folders_when_organismo_template_exists(
    db: AsyncSession,
    test_user: dict[str, Any],
    contrato: Contrato,
    cuenta: CuentaCobro,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """When the contract's organism has an ingested template whose structure
    includes anexo references (COEMPRESAR-style, 3-column + literal anexo
    refs), the packager numbers evidence folders "A{n}_..." to mirror the
    institutional convention instead of the plain "01_..." default."""
    from app.core.text_match import normalize
    from app.models.plantilla_organismo import PlantillaOrganismo

    user = test_user["user"]
    plantilla = PlantillaOrganismo(
        usuario_id=user.id,
        entidad=contrato.entidad,
        entidad_normalizada=normalize(contrato.entidad),
        tipo_documento="informe_actividades",
        formato="docx",
        estructura_json={
            "columnas": ["obligación", "avance del periodo", "evidencia"],
            "secciones": [],
            "anexo_refs": ["Ver Anexo: Carpeta /5. EVIDENCIAS/A1"],
            "notas": "",
        },
    )
    db.add(plantilla)
    # A folder only renders once it has a real evidencia file to carry (A3).
    res = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).limit(1))
    act = res.scalars().first()
    assert act is not None
    db.add(
        Evidencia(
            actividad_id=act.id,
            storage_key=f"evidencias/{act.id}/soporte.pdf",
            nombre_archivo="soporte.pdf",
            tipo_archivo="application/pdf",
            tamano_bytes=10,
        )
    )
    await db.commit()
    # See test_zip_evidencias_estructura for why this expire is required.
    db.expire(act, ["evidencias"])

    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        folders = {n.split("/")[0] for n in zf.namelist() if "/" in n}
    assert any(f.startswith("A1_") for f in folders)
    assert not any(f.startswith(("01_", "02_", "03_")) for f in folders)


async def test_zip_evidencias_estado_pendiente_en_manifest_modo_estandar(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """No evidencia files uploaded at all → every obligación is PENDIENTE. Standard
    mode still emits the package, listing PENDIENTE items in the manifest."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        root = zf.read("LEEME.txt").decode("utf-8")
        assert "PENDIENTE" in root


async def test_zip_evidencias_modo_final_raises_package_pendiente(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """No evidencia AND no justificación on any obligación → still PACKAGE_PENDIENTE
    (ground truth fix: a bare absence of evidencia alone is NOT blocking — see
    `test_zip_evidencias_modo_final_justificacion_sin_evidencia_emite_zip` below)."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    res = await db.execute(
        select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).order_by(Actividad.fecha_realizacion.asc())
    )
    for act in res.scalars().all():
        act.justificacion = ""
    await db.commit()

    with pytest.raises(ValidationError) as exc_info:
        await informe_service.generar_zip_evidencias(db, user.id, cuenta.id, modo="final")
    assert exc_info.value.code == "PACKAGE_PENDIENTE"


async def test_zip_evidencias_modo_final_justificacion_sin_evidencia_emite_zip(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    actividad_con_evidencia: Actividad,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Obligación 1 has evidencia; obligaciones 2/3 carry only the SENTINEL
    justificación — modo="final" must still emit the package: real approved
    cuentas radicate with partial evidencia coverage as long as every obligación
    has evidencia OR a justificación (the gate deliberately accepts the sentinel).
    H12 note: the base fixture's REAL justificaciones now make an obligación
    LISTO outright, so the sentinel is what keeps 2/3 in the "justificadas sin
    evidencia" LEEME split this test asserts."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    res = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id))
    for act in res.scalars().all():
        if act.id != actividad_con_evidencia.id:
            act.justificacion = SENTINEL_SIN_EVIDENCIAS
    await db.commit()

    content, filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id, modo="final")

    assert filename.endswith(".zip")
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        root = zf.read("LEEME.txt").decode("utf-8")
        assert "justificadas sin evidencia" in root


async def test_zip_evidencias_modo_final_obligacion_sin_nada_bloquea_aunque_otras_justificadas(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """Mixed cuenta: 2 obligaciones keep their justificación (base fixture), the
    3rd is blanked to neither evidencia nor justificación — modo="final" must
    still raise PACKAGE_PENDIENTE for that one obligación."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    res = await db.execute(
        select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).order_by(Actividad.fecha_realizacion.asc())
    )
    actos = res.scalars().all()
    actos[-1].justificacion = ""
    await db.commit()

    with pytest.raises(ValidationError) as exc_info:
        await informe_service.generar_zip_evidencias(db, user.id, cuenta.id, modo="final")
    assert exc_info.value.code == "PACKAGE_PENDIENTE"
    assert "1 obligación" in str(exc_info.value)


async def test_zip_evidencias_secreto_detectado_bloquea_paquete(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    actividad_con_evidencia: Actividad,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A real-leak-SHAPED (synthetic) text file among the evidencia bytes must halt
    packaging entirely — no zip emitted, SECRET_DETECTED_IN_PACKAGE raised."""
    user = test_user["user"]
    ev = actividad_con_evidencia.evidencias[0]
    storage = _fake_storage({ev.storage_key: _LEAK_CORPUS_TEXTO.encode("utf-8")})
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    with pytest.raises(ValidationError) as exc_info:
        await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)
    assert exc_info.value.code == "SECRET_DETECTED_IN_PACKAGE"


async def test_zip_evidencias_flag_desactivado_bypassa_scan(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    actividad_con_evidencia: Actividad,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """SECRET_SCAN_GATE_ENABLED=False bypasses the scan even with a leak-shaped
    payload present — emergency disable only."""
    user = test_user["user"]
    monkeypatch.setattr(settings, "SECRET_SCAN_GATE_ENABLED", False)
    ev = actividad_con_evidencia.evidencias[0]
    storage = _fake_storage({ev.storage_key: _LEAK_CORPUS_TEXTO.encode("utf-8")})
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    content, filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)
    assert filename.endswith(".zip")
    assert content


# ── DOCUMENTOS_CONTRATO/ en el paquete (clasificacion-documentos-secop, D5) ──


async def _setup_docs_contrato(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
) -> dict[str, Any]:
    """Checklist + links: an uploaded CONTRATO doc, a SECOP RPC doc and an
    unlinked OTROS SECOP doc (context). Returns the created objects."""
    from app.models.categoria_documento import CategoriaDocumento
    from app.models.documento_fuente import DocumentoFuente, TipoDocumentoFuente
    from app.models.secop import SecopDocumento
    from app.services import checklist_service

    contrato_id = cuenta.contrato_id
    user = test_user["user"]
    await checklist_service.asegurar_checklist(db, cuenta)

    doc_contrato = DocumentoFuente(
        usuario_id=user.id,
        contrato_id=contrato_id,
        storage_key=f"docs/{uuid.uuid4()}/contrato_firmado.pdf",
        nombre="contrato_firmado.pdf",
        tipo=TipoDocumentoFuente.CONTRATO,
        categoria=CategoriaDocumento.CONTRATO,
    )
    rpc_secop = SecopDocumento(
        id_documento_secop=f"DOC-{uuid.uuid4().hex[:10]}",
        numero_contrato="CTR-INF-001",
        nombre_archivo="rpc_secop.pdf",
        url_descarga="https://s/rpc_secop.pdf",
        categoria=CategoriaDocumento.REGISTRO_PRESUPUESTAL,
        datos_raw={},
    )
    contexto_secop = SecopDocumento(
        id_documento_secop=f"DOC-{uuid.uuid4().hex[:10]}",
        numero_contrato="CTR-INF-001",
        nombre_archivo="acta_interna.pdf",
        url_descarga="https://s/acta_interna.pdf",
        categoria=CategoriaDocumento.OTROS,
        datos_raw={},
    )
    db.add_all([doc_contrato, rpc_secop, contexto_secop])
    await db.commit()

    await checklist_service.vincular_documento_fuente(db, cuenta.id, "CONTRATO", doc_contrato.id)
    await checklist_service.vincular_secop_documento(db, cuenta.id, "RPC", rpc_secop.id)
    await db.commit()
    return {"doc_contrato": doc_contrato, "rpc_secop": rpc_secop, "contexto_secop": contexto_secop}


def _fake_descargas_secop(monkeypatch: pytest.MonkeyPatch, payloads: dict[str, bytes], errores: set[str] | None = None):
    from app.services import checklist_service

    errores = errores or set()

    async def _fake(url: str) -> bytes:
        if url in errores:
            raise RuntimeError("simulated secop download failure")
        return payloads[url]

    monkeypatch.setattr(checklist_service, "_descargar_secop_bytes", _fake)


async def test_zip_incluye_documentos_contrato_por_requisito_y_contexto(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Linked contract-level docs land in DOCUMENTOS_CONTRATO/{CODIGO}/ with real
    bytes (uploads via storage, SECOP via url_descarga); remaining context docs
    land in DOCUMENTOS_CONTRATO/CONTEXTO/."""
    user = test_user["user"]
    objetos = await _setup_docs_contrato(db, test_user, cuenta)

    bytes_contrato = b"bytes-reales-contrato-firmado"
    bytes_rpc = b"bytes-reales-rpc-secop"
    bytes_contexto = b"bytes-reales-acta-interna"
    storage = _fake_storage({objetos["doc_contrato"].storage_key: bytes_contrato})
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)
    _fake_descargas_secop(
        monkeypatch,
        {"https://s/rpc_secop.pdf": bytes_rpc, "https://s/acta_interna.pdf": bytes_contexto},
    )

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        assert "DOCUMENTOS_CONTRATO/CONTRATO/contrato_firmado.pdf" in names
        assert zf.read("DOCUMENTOS_CONTRATO/CONTRATO/contrato_firmado.pdf") == bytes_contrato
        assert "DOCUMENTOS_CONTRATO/RPC/rpc_secop.pdf" in names
        assert zf.read("DOCUMENTOS_CONTRATO/RPC/rpc_secop.pdf") == bytes_rpc
        assert "DOCUMENTOS_CONTRATO/CONTEXTO/acta_interna.pdf" in names
        assert zf.read("DOCUMENTOS_CONTRATO/CONTEXTO/acta_interna.pdf") == bytes_contexto


async def test_zip_documentos_contrato_skip_on_error(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """One failing download (storage or SECOP) skips ONLY that member — the
    package is still emitted with the rest."""
    user = test_user["user"]
    objetos = await _setup_docs_contrato(db, test_user, cuenta)

    storage = AsyncMock()
    storage.download = AsyncMock(side_effect=RuntimeError("storage down"))
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)
    _fake_descargas_secop(
        monkeypatch,
        {"https://s/rpc_secop.pdf": b"bytes-rpc"},
        errores={"https://s/acta_interna.pdf"},
    )

    content, filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    assert filename.endswith(".zip")
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        assert "DOCUMENTOS_CONTRATO/RPC/rpc_secop.pdf" in names  # the healthy member survives
        assert "DOCUMENTOS_CONTRATO/CONTRATO/contrato_firmado.pdf" not in names
        assert "DOCUMENTOS_CONTRATO/CONTEXTO/acta_interna.pdf" not in names
    assert objetos is not None


async def test_zip_documentos_contrato_entra_al_secret_scan(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A leak-shaped payload inside a DOCUMENTOS_CONTRATO member must halt the
    package — the new members join the mandatory secret scan."""
    user = test_user["user"]
    await _setup_docs_contrato(db, test_user, cuenta)

    storage = _fake_storage()
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)
    _fake_descargas_secop(
        monkeypatch,
        {
            "https://s/rpc_secop.pdf": b"bytes-rpc",
            "https://s/acta_interna.pdf": _LEAK_CORPUS_TEXTO.encode("utf-8"),
        },
    )

    with pytest.raises(ValidationError) as exc_info:
        await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)
    assert exc_info.value.code == "SECRET_DETECTED_IN_PACKAGE"


async def test_obtener_estado_listo_pendiente_sin_zip(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    actividad_con_evidencia: Actividad,
) -> None:
    # H12: a REAL justificación now counts as listo, and the base `cuenta`
    # fixture writes one on every actividad — downgrade the two non-evidencia
    # actividades to the sentinel (explicit absence of work) so this test keeps
    # exercising the evidencia-based split AND the sentinel-does-not-count rule.
    user = test_user["user"]
    res = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id))
    for act in res.scalars().all():
        if act.id != actividad_con_evidencia.id:
            act.justificacion = SENTINEL_SIN_EVIDENCIAS
    await db.commit()

    estado = await informe_service.obtener_estado_listo_pendiente(db, user.id, cuenta.id)

    assert estado.cuenta_cobro_id == cuenta.id
    assert len(estado.obligaciones) == 3
    listos = [o for o in estado.obligaciones if o.listo]
    assert len(listos) == 1
    assert estado.pendientes == 2
    assert estado.listo_para_radicar is False


async def test_estado_justificacion_real_sola_cuenta_como_listo(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
) -> None:
    """H12 parity: an obligación whose actividad carries a REAL justificación —
    no evidencia file, no confirmed link — is LISTO, matching the modo="final"
    package gate that has always accepted justificación-only coverage."""
    user = test_user["user"]
    # Workstream B (B5): coverage now keys off `justificacion_origen`, not text —
    # the base `cuenta` fixture's text defaults to `seed` (model default), so mark
    # it `llm` explicitly to exercise the "real justificación" branch under test.
    res = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id))
    for act in res.scalars().all():
        act.justificacion_origen = "llm"
    await db.commit()

    estado = await informe_service.obtener_estado_listo_pendiente(db, user.id, cuenta.id)

    # Base fixture: 3 actividades, all with real justificación, zero evidencias.
    assert estado.pendientes == 0
    assert estado.listo_para_radicar is True


async def test_estado_justificacion_sentinel_no_cuenta_como_listo(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
) -> None:
    """H12: the deterministic sentinel means "no work done" — it must NOT flip
    an obligación to LISTO."""
    user = test_user["user"]
    res = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id))
    for act in res.scalars().all():
        act.justificacion = SENTINEL_SIN_EVIDENCIAS
    await db.commit()

    estado = await informe_service.obtener_estado_listo_pendiente(db, user.id, cuenta.id)

    assert estado.pendientes == 3
    assert estado.listo_para_radicar is False


def test_tiene_justificacion_real_por_origen() -> None:
    """Workstream B (B5): coverage is decided by `justificacion_origen` alone —
    llm/manual count, seed/sin_labores/None never do."""
    assert informe_service._tiene_justificacion_real("llm") is True
    assert informe_service._tiene_justificacion_real("manual") is True
    assert informe_service._tiene_justificacion_real("seed") is False
    assert informe_service._tiene_justificacion_real("sin_labores") is False
    assert informe_service._tiene_justificacion_real(None) is False


async def test_estado_listo_cuenta_link_confirmado_en_stub_sin_clasificar(
    db: AsyncSession,
    test_user: dict[str, Any],
    cuenta: CuentaCobro,
    obligaciones: list[Obligacion],
) -> None:
    """Classify-after-upload path (live-reproduced 2026-08-11): an evidencia
    uploaded before classification hangs off the shared "sin clasificar" stub
    Actividad (obligacion_id=None); confirming its link later only writes
    `evidencia_obligacion`. The obligación must still count as LISTO — link
    table as source of truth, parity with cobertura/checklist."""
    user = test_user["user"]
    # H12: neutralize the base fixture's real justificaciones (they now count as
    # listo on their own) so this test isolates the link-table branch.
    res = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id))
    for act in res.scalars().all():
        act.justificacion = SENTINEL_SIN_EVIDENCIAS
    await db.flush()
    stub = Actividad(
        cuenta_cobro_id=cuenta.id,
        obligacion_id=None,
        descripcion="Pendiente de redactar — evidencia adjunta (stub)",
    )
    db.add(stub)
    await db.flush()
    ev = Evidencia(
        actividad_id=stub.id,
        storage_key=f"evidencias/{stub.id}/consolidado.pdf",
        nombre_archivo="consolidado.pdf",
        tipo_archivo="application/pdf",
        tamano_bytes=10,
    )
    db.add(ev)
    await db.flush()
    db.add(
        EvidenciaObligacion(
            evidencia_id=ev.id,
            obligacion_id=obligaciones[0].id,
            status=EstadoEnlace.CONFIRMED.value,
            source=FuenteEnlace.USER.value,
        )
    )
    await db.commit()

    estado = await informe_service.obtener_estado_listo_pendiente(db, user.id, cuenta.id)

    por_id = {o.obligacion_id: o.listo for o in estado.obligaciones}
    assert por_id[obligaciones[0].id] is True
    # A confirmed link on one obligación must not leak listo to the others.
    assert estado.pendientes == 2


# ── Paquete radicado completo (G1 docs generados, G3 enlaces, G5 diferencial) ──


def _mock_llm_offline(monkeypatch: pytest.MonkeyPatch) -> None:
    """`generar_informe_supervision_docx` (generated for every paquete) rewrites
    justificaciones to third person via the LLM — mock it to keep these tests
    offline and deterministic (fails open if unmocked, but must never hit the
    network)."""
    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(return_value=MagicMock(content="Narrativa de prueba."))
    monkeypatch.setattr(informe_service, "get_llm", lambda: mock_llm)


async def test_zip_incluye_informes_generados_en_raiz(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """The always-available generated docs (informe de actividades / supervisión)
    ship at the ZIP root numbered like the real radicated packages; the clone-only
    docs (cuenta de cobro / documento soporte) have no ingested plantilla here →
    skipped up front with a truthful "Sin plantilla" line (NOT the failure AVISO,
    whose re-download advice would hit the same missing-plantilla error)."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    _mock_llm_offline(monkeypatch)

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        assert "2. INFORME DE ACTIVIDADES.docx" in names
        assert "3. INFORME DE SUPERVISION.docx" in names
        # Valid DOCX bytes — the exact same generator the download endpoint uses
        Document(io.BytesIO(zf.read("2. INFORME DE ACTIVIDADES.docx")))
        assert "1. CUENTA DE COBRO.docx" not in names  # clone-only, sin plantilla
        assert "4. DOCUMENTO SOPORTE.xlsx" not in names
        root = zf.read("LEEME.txt").decode("utf-8")
        assert "2. INFORME DE ACTIVIDADES.docx" in root
        assert "Sin plantilla del organismo para 1. CUENTA DE COBRO.docx" in root
        assert "Sin plantilla del organismo para 4. DOCUMENTO SOPORTE.xlsx" in root
        assert "AVISO: No se pudo generar" not in root  # normal state, not a failure


async def test_zip_incluye_los_cuatro_documentos_cuando_hay_plantillas(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """With clonable plantillas for the organism (generators mocked at the same
    service seam) the package carries all 4 numbered docs and no AVISO."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    _mock_llm_offline(monkeypatch)
    monkeypatch.setattr(informe_service, "_hay_plantilla_organismo", AsyncMock(return_value=True))
    monkeypatch.setattr(
        informe_service, "generar_cuenta_cobro_docx", AsyncMock(return_value=(b"docx-cuenta-cobro", "cc.docx"))
    )
    monkeypatch.setattr(
        informe_service, "generar_documento_soporte_xlsx", AsyncMock(return_value=(b"xlsx-soporte", "ds.xlsx"))
    )

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        for esperado in (
            "1. CUENTA DE COBRO.docx",
            "2. INFORME DE ACTIVIDADES.docx",
            "3. INFORME DE SUPERVISION.docx",
            "4. DOCUMENTO SOPORTE.xlsx",
        ):
            assert esperado in names
        assert zf.read("1. CUENTA DE COBRO.docx") == b"docx-cuenta-cobro"
        assert zf.read("4. DOCUMENTO SOPORTE.xlsx") == b"xlsx-soporte"
        root = zf.read("LEEME.txt").decode("utf-8")
        assert "AVISO: No se pudo generar" not in root
        assert "Sin plantilla del organismo" not in root


async def test_zip_plantilla_no_clonable_se_omite_como_sin_plantilla(
    db: AsyncSession,
    test_user: dict[str, Any],
    contrato: Contrato,
    cuenta: CuentaCobro,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """FIX-4: a plantilla ROW that exists but isn't clonable must be treated the same
    as "no plantilla ingested" — the truthful "Sin plantilla del organismo..." skip,
    NOT the generic "No se pudo generar" failure AVISO (`generar_cuenta_cobro_docx`
    itself would raise `ValidationError` on a non-clonable plantilla; the pre-check
    must predict that instead of letting it fall through to the catch-all)."""
    from app.core.text_match import normalize
    from app.models.plantilla_organismo import PlantillaOrganismo

    user = test_user["user"]
    plantilla = PlantillaOrganismo(
        usuario_id=user.id,
        entidad=contrato.entidad,
        entidad_normalizada=normalize(contrato.entidad),
        tipo_documento="cuenta_cobro",
        formato="docx",
        estructura_json={"clonable": False},
    )
    db.add(plantilla)
    await db.commit()

    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    _mock_llm_offline(monkeypatch)
    spy = AsyncMock(side_effect=AssertionError("generator must not run for a non-clonable plantilla"))
    monkeypatch.setattr(informe_service, "generar_cuenta_cobro_docx", spy)

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    spy.assert_not_called()
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        assert "1. CUENTA DE COBRO.docx" not in names
        root = zf.read("LEEME.txt").decode("utf-8")
        assert "Sin plantilla del organismo para 1. CUENTA DE COBRO.docx" in root
        assert "AVISO: No se pudo generar 1. CUENTA DE COBRO.docx" not in root


async def test_hay_plantilla_organismo_documento_soporte_requiere_condiciones_completas(
    db: AsyncSession, test_user: dict[str, Any], contrato: Contrato
) -> None:
    """FIX-4: `_hay_plantilla_organismo("documento_soporte")` must mirror
    `generar_documento_soporte_xlsx`'s FULL precondition (clonable AND formato=xlsx
    AND non-empty campos AND fuente_documento_id) — a plantilla missing any one of
    these must read as "no usable plantilla", same as the generator would reject it."""
    from app.core.text_match import normalize
    from app.models.plantilla_organismo import PlantillaOrganismo

    user = test_user["user"]

    # clonable=True but no campos + no fuente_documento_id — real generator would raise.
    plantilla = PlantillaOrganismo(
        usuario_id=user.id,
        entidad=contrato.entidad,
        entidad_normalizada=normalize(contrato.entidad),
        tipo_documento="documento_soporte",
        formato="xlsx",
        estructura_json={"clonable": True},
    )
    db.add(plantilla)
    await db.commit()

    assert await informe_service._hay_plantilla_organismo(db, contrato, "documento_soporte") is False


async def test_zip_documento_generado_fallo_no_hunde_paquete(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """FAIL-OPEN: a generator raising (e.g. LLM failure) skips ONLY that doc —
    the package is still emitted and the root LEEME carries the genuine-failure
    AVISO (wording preserved: this is a real error, not a missing plantilla)."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    _mock_llm_offline(monkeypatch)
    monkeypatch.setattr(
        informe_service, "generar_informe_actividades_docx", AsyncMock(side_effect=RuntimeError("LLM caído"))
    )

    content, filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    assert filename.endswith(".zip")
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        assert "2. INFORME DE ACTIVIDADES.docx" not in names
        assert "3. INFORME DE SUPERVISION.docx" in names  # the healthy doc survives
        root = zf.read("LEEME.txt").decode("utf-8")
        assert (
            "AVISO: No se pudo generar 2. INFORME DE ACTIVIDADES.docx — "
            "descargalo por separado desde el paso Formato." in root
        )


async def test_zip_documentos_generados_entran_al_secret_scan(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """A leak-shaped payload inside a generated root doc must halt the package —
    the generated docs join the mandatory secret scan exactly like the
    DOCUMENTOS_CONTRATO members (`test_zip_documentos_contrato_entra_al_secret_scan`)."""
    user = test_user["user"]
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())
    _mock_llm_offline(monkeypatch)
    monkeypatch.setattr(
        informe_service,
        "generar_informe_actividades_docx",
        AsyncMock(return_value=(_LEAK_CORPUS_TEXTO.encode("utf-8"), "informe.docx")),
    )

    with pytest.raises(ValidationError) as exc_info:
        await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)
    assert exc_info.value.code == "SECRET_DETECTED_IN_PACKAGE"


async def test_zip_evidencia_link_only_no_genera_carpeta_ni_se_cuela_en_sin_clasificar(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Link-only evidencias (fuente/url, no storage_key) contribute no bytes —
    since A3 removed the per-folder LEEME (which used to document them under an
    "Enlaces:" section), they now simply produce no folder at all, and are NOT
    swept into sin_clasificar/ either (A2 only ships evidencias WITH storage_key).
    Covers both a usable url and a missing one — neither crashes packaging."""
    user = test_user["user"]
    res = await db.execute(
        select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id).order_by(Actividad.fecha_realizacion.asc())
    )
    act = res.scalars().first()
    assert act is not None
    db.add_all(
        [
            Evidencia(
                actividad_id=act.id,
                storage_key=None,
                nombre_archivo="correo de aprobación",
                fuente="gmail",
                url="https://mail.google.com/mail/u/0/#inbox/abc123",
            ),
            Evidencia(
                actividad_id=act.id,
                storage_key=None,
                nombre_archivo="evento sin enlace utilizable",
                fuente="calendar",
                url=None,
            ),
        ]
    )
    await db.commit()
    await db.refresh(act)
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
    assert not any(n.startswith("01_") for n in names)
    assert not any(n.startswith("sin_clasificar/") for n in names)


async def _vincular_cedula(db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro) -> None:
    """Checklist + an uploaded CEDULA doc linked to its (nivel-contrato,
    solo_primera_cuenta) checklist row."""
    from app.models.categoria_documento import CategoriaDocumento
    from app.models.documento_fuente import DocumentoFuente, TipoDocumentoFuente
    from app.services import checklist_service

    await checklist_service.asegurar_checklist(db, cuenta)
    doc = DocumentoFuente(
        usuario_id=test_user["user"].id,
        contrato_id=cuenta.contrato_id,
        storage_key=f"docs/{uuid.uuid4()}/cedula.pdf",
        nombre="cedula.pdf",
        tipo=TipoDocumentoFuente.CEDULA,
        categoria=CategoriaDocumento.CEDULA,
    )
    db.add(doc)
    await db.commit()
    await checklist_service.vincular_documento_fuente(db, cuenta.id, "CEDULA", doc.id)
    await db.commit()


async def test_zip_omite_docs_solo_primera_en_cuota_no_primera(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Differential package: the fixture cuenta is posicion=RECURRENTE (not first)
    — solo_primera_cuenta docs (CEDULA...) are omitted and the root LEEME says so."""
    user = test_user["user"]
    assert cuenta.posicion == PosicionCuota.RECURRENTE
    await _vincular_cedula(db, test_user, cuenta)
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        root = zf.read("LEEME.txt").decode("utf-8")
    assert not any(n.startswith("DOCUMENTOS_CONTRATO/CEDULA/") for n in names)
    assert "Documentos de primera cuota omitidos: ya radicados en la cuota 1" in root
    assert "CEDULA" in root


async def test_zip_incluye_docs_solo_primera_en_cuota_primera(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, monkeypatch: pytest.MonkeyPatch
) -> None:
    """On the first cuota the solo_primera docs ship normally — no omission note."""
    user = test_user["user"]
    cuenta.posicion = PosicionCuota.PRIMERA
    await db.commit()
    await _vincular_cedula(db, test_user, cuenta)
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: _fake_storage())

    content, _filename = await informe_service.generar_zip_evidencias(db, user.id, cuenta.id)

    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        root = zf.read("LEEME.txt").decode("utf-8")
    assert "DOCUMENTOS_CONTRATO/CEDULA/cedula.pdf" in names
    assert "Documentos de primera cuota omitidos" not in root


# ── Adaptive generation (billing-resilience-templates, slice #6) ───────────
#
# Per-organism layout selection, always-draft header, progressive narrative,
# and one-time obligation blanking. `generar_informe_actividades_docx`/
# `generar_informe_supervision_docx` keep their exact `(db, usuario_id,
# cuenta_id) -> (bytes, filename)` public signature — organismo/posicion/
# prior_context are resolved INTERNALLY from the already-loaded contrato/cuenta
# rather than added as new required params (design mentions adding params to
# the generators; every existing call site — `checklist_autogen_service.
# _GENERADORES`, `app/api/v1/cuentas_cobro.py`, the two informe tools — calls
# them positionally with exactly 3 args, so widening the signature would be a
# breaking, undeclared change to 3+ call sites for no behavioral benefit).


@pytest.fixture
async def contrato_dagma(db: AsyncSession, test_user: dict[str, Any]) -> Contrato:
    user = test_user["user"]
    c = Contrato(
        usuario_id=user.id,
        numero_contrato="CTR-DAGMA-001",
        objeto="Servicios profesionales ambientales",
        valor_total=24_000_000,
        valor_mensual=2_000_000,
        fecha_inicio=date(2024, 1, 1),
        fecha_fin=date(2024, 12, 31),
        entidad="DAGMA",
        dependencia="Ambiental",
        supervisor_nombre="Sup DAGMA",
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


@pytest.fixture
async def contrato_coempresar(db: AsyncSession, test_user: dict[str, Any]) -> Contrato:
    user = test_user["user"]
    c = Contrato(
        usuario_id=user.id,
        numero_contrato="CTR-COEMPRESAR-001",
        objeto="Servicios profesionales de saneamiento",
        valor_total=24_000_000,
        valor_mensual=2_000_000,
        fecha_inicio=date(2024, 1, 1),
        fecha_fin=date(2024, 12, 31),
        entidad="COEMPRESAR",
        dependencia="Operaciones",
        supervisor_nombre="Sup Coempresar",
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


async def _make_obligaciones(db: AsyncSession, contrato: Contrato, n: int = 2) -> list[Obligacion]:
    obs = [
        Obligacion(
            contrato_id=contrato.id,
            descripcion=f"Obligación contractual #{i + 1} con texto suficientemente largo",
            tipo=TipoObligacion.GENERAL,
            orden=i,
        )
        for i in range(n)
    ]
    db.add_all(obs)
    await db.commit()
    for o in obs:
        await db.refresh(o)
    contrato.obligaciones = list(obs)
    return obs


async def _make_cuenta_con_actividades(
    db: AsyncSession,
    contrato: Contrato,
    obligaciones: list[Obligacion],
    mes: int,
    anio: int = 2024,
    numero_cuota: int | None = None,
    posicion: PosicionCuota = PosicionCuota.RECURRENTE,
) -> CuentaCobro:
    cc = CuentaCobro(
        contrato_id=contrato.id,
        mes=mes,
        anio=anio,
        estado=EstadoCuentaCobro.BORRADOR,
        valor=2_000_000,
        numero_cuota=numero_cuota,
        posicion=posicion,
    )
    db.add(cc)
    await db.commit()
    await db.refresh(cc)
    for i, ob in enumerate(obligaciones):
        db.add(
            Actividad(
                cuenta_cobro_id=cc.id,
                obligacion_id=ob.id,
                descripcion=f"Actividad realizada {i + 1} en {mes}/{anio}",
                justificacion=f"Justificación detallada {i + 1} en {mes}/{anio}",
                fecha_realizacion=date(anio, mes, 10),
            )
        )
    await db.commit()
    await db.refresh(cc)
    return cc


async def _ingerir_plantilla(
    db: AsyncSession,
    contrato: Contrato,
    columnas: list[str],
    anexo_refs: list[str] | None = None,
    tipo_documento: str = "informe_actividades",
) -> PlantillaOrganismo:
    plantilla = PlantillaOrganismo(
        usuario_id=contrato.usuario_id,
        entidad=contrato.entidad,
        entidad_normalizada=normalize(contrato.entidad),
        tipo_documento=tipo_documento,
        formato="docx",
        estructura_json={
            "columnas": columnas,
            "secciones": [],
            "anexo_refs": anexo_refs or [],
            "notas": "",
        },
    )
    db.add(plantilla)
    await db.commit()
    await db.refresh(plantilla)
    return plantilla


# ── Per-organism layout selection (6.1-6.4) ─────────────────────────────────


async def test_informe_actividades_dagma_2_columnas(
    db: AsyncSession, test_user: dict[str, Any], contrato_dagma: Contrato
) -> None:
    """DAGMA-style organism: 2-column layout (obligación | avance), no evidence
    column — scenario spec `adaptive-informe-generation`."""
    user = test_user["user"]
    obligaciones = await _make_obligaciones(db, contrato_dagma)
    await _ingerir_plantilla(db, contrato_dagma, columnas=["Obligación", "Avance del periodo"])
    cuenta = await _make_cuenta_con_actividades(db, contrato_dagma, obligaciones, mes=1, numero_cuota=1)

    content, _filename = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta.id)
    doc = Document(io.BytesIO(content))
    tabla = doc.tables[1]
    headers = [c.text for c in tabla.rows[0].cells]
    assert headers == ["Obligación", "Avance del periodo"]
    assert len(tabla.columns) == 2
    assert not any("evidenc" in h.lower() for h in headers)


async def test_informe_actividades_coempresar_3_columnas_con_anexo_refs(
    db: AsyncSession, test_user: dict[str, Any], contrato_coempresar: Contrato
) -> None:
    """COEMPRESAR-style organism: 3-column layout (+ evidencia) and the literal
    anexo reference from the ingested template appears verbatim in the DOCX."""
    user = test_user["user"]
    obligaciones = await _make_obligaciones(db, contrato_coempresar)
    anexo_ref = "Ver Anexo: Carpeta /5. EVIDENCIAS/A1"
    await _ingerir_plantilla(
        db,
        contrato_coempresar,
        columnas=["Obligación", "Avance del periodo", "Evidencia"],
        anexo_refs=[anexo_ref],
    )
    cuenta = await _make_cuenta_con_actividades(db, contrato_coempresar, obligaciones, mes=1, numero_cuota=1)

    content, _filename = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta.id)
    doc = Document(io.BytesIO(content))
    tabla = doc.tables[1]
    headers = [c.text for c in tabla.rows[0].cells]
    assert headers == ["Obligación", "Avance del periodo", "Evidencia"]
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert anexo_ref in full_text


async def test_informe_actividades_default_4_columnas_sin_plantilla(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    """No organism template ingested — the current default 4-column layout is
    used unchanged (zero regression)."""
    user = test_user["user"]
    content, _filename = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta.id)
    doc = Document(io.BytesIO(content))
    tabla = doc.tables[1]
    headers = [c.text for c in tabla.rows[0].cells]
    assert headers == ["#", "Obligación", "Actividad realizada", "Justificación"]


# ── Always-draft header (6.7-6.8) ───────────────────────────────────────────


async def test_informe_actividades_siempre_carga_encabezado_borrador(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    user = test_user["user"]
    content, _filename = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta.id)
    doc = Document(io.BytesIO(content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "BORRADOR" in full_text
    assert "sujeto a revisión" in full_text


async def test_informe_supervision_siempre_carga_encabezado_borrador(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    user = test_user["user"]
    content, _filename = await informe_service.generar_informe_supervision_docx(db, user.id, cuenta.id)
    doc = Document(io.BytesIO(content))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "BORRADOR" in full_text


# ── Progressive narrative (6.5-6.6) ─────────────────────────────────────────


async def test_contexto_progresivo_ausente_en_primera_cuota(db: AsyncSession, contrato_dagma: Contrato) -> None:
    obligaciones = await _make_obligaciones(db, contrato_dagma)
    cuenta1 = await _make_cuenta_con_actividades(
        db, contrato_dagma, obligaciones, mes=1, numero_cuota=1, posicion=PosicionCuota.PRIMERA
    )
    contexto = await informe_service._construir_contexto_progresivo(db, contrato_dagma.id, cuenta1)
    assert contexto is None


async def test_contexto_progresivo_construido_desde_cuotas_previas(db: AsyncSession, contrato_dagma: Contrato) -> None:
    """Cuota 3's narrative is built from cuotas 1 and 2's recorded activity text."""
    obligaciones = await _make_obligaciones(db, contrato_dagma)
    await _make_cuenta_con_actividades(
        db, contrato_dagma, obligaciones, mes=1, numero_cuota=1, posicion=PosicionCuota.PRIMERA
    )
    await _make_cuenta_con_actividades(db, contrato_dagma, obligaciones, mes=2, numero_cuota=2)
    cuenta3 = await _make_cuenta_con_actividades(db, contrato_dagma, obligaciones, mes=3, numero_cuota=3)

    fake_resp = MagicMock()
    fake_resp.content = "Narrativa progresiva sintetizada."
    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(return_value=fake_resp)

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        contexto = await informe_service._construir_contexto_progresivo(db, contrato_dagma.id, cuenta3)

    assert contexto == "Narrativa progresiva sintetizada."
    mock_llm.complete.assert_awaited_once()
    prompt_sent = str(mock_llm.complete.call_args)
    assert "1/2024" in prompt_sent
    assert "2/2024" in prompt_sent


async def test_contexto_progresivo_llm_error_falla_abierto(db: AsyncSession, contrato_dagma: Contrato) -> None:
    """A narrative-synthesis LLM failure degrades to the raw joined summaries —
    it never blocks or raises (mirrors `_convertir_actividades_tercera_persona`)."""
    obligaciones = await _make_obligaciones(db, contrato_dagma)
    await _make_cuenta_con_actividades(
        db, contrato_dagma, obligaciones, mes=1, numero_cuota=1, posicion=PosicionCuota.PRIMERA
    )
    cuenta2 = await _make_cuenta_con_actividades(db, contrato_dagma, obligaciones, mes=2, numero_cuota=2)

    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(side_effect=RuntimeError("llm down"))

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        contexto = await informe_service._construir_contexto_progresivo(db, contrato_dagma.id, cuenta2)

    assert contexto is not None
    assert "Actividad realizada 1 en 1/2024" in contexto


async def test_contexto_progresivo_degrada_a_k3_cuando_excede_presupuesto(
    db: AsyncSession, contrato_dagma: Contrato, monkeypatch: pytest.MonkeyPatch
) -> None:
    """When the bounded char budget is exceeded, degrade to the K=3 most recent
    summaries + a note about how many were omitted."""
    monkeypatch.setattr(informe_service, "_MAX_TEXT_CHARS", 10)
    obligaciones = await _make_obligaciones(db, contrato_dagma, n=1)
    for mes in range(1, 5):
        await _make_cuenta_con_actividades(db, contrato_dagma, obligaciones, mes=mes, numero_cuota=mes)
    cuenta5 = await _make_cuenta_con_actividades(db, contrato_dagma, obligaciones, mes=5, numero_cuota=5)

    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(side_effect=RuntimeError("skip llm, inspect raw block"))

    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        contexto = await informe_service._construir_contexto_progresivo(db, contrato_dagma.id, cuenta5)

    assert contexto is not None
    assert "1/2024" not in contexto  # oldest cuota omitted
    assert "2/2024" in contexto
    assert "3/2024" in contexto
    assert "4/2024" in contexto
    assert "se omiten" in contexto.lower()


# ── One-time obligation blanking (6.9-6.10) ─────────────────────────────────


async def test_obligacion_una_vez_visible_en_cuota_primera_y_ausente_en_recurrente(
    db: AsyncSession, test_user: dict[str, Any], contrato_dagma: Contrato
) -> None:
    user = test_user["user"]
    obligaciones = await _make_obligaciones(db, contrato_dagma, n=2)
    obligacion_unica = obligaciones[0]
    obligacion_unica.una_vez = True
    db.add(obligacion_unica)
    await db.commit()

    cuenta1 = await _make_cuenta_con_actividades(
        db, contrato_dagma, obligaciones, mes=1, numero_cuota=1, posicion=PosicionCuota.PRIMERA
    )
    cuenta2 = await _make_cuenta_con_actividades(db, contrato_dagma, obligaciones, mes=2, numero_cuota=2)

    content1, _f1 = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta1.id)
    doc1 = Document(io.BytesIO(content1))
    tabla1 = doc1.tables[1]
    body1 = "\n".join(cell.text for row in tabla1.rows for cell in row.cells)
    assert obligacion_unica.descripcion in body1

    # cuenta2 has a prior cuota, so `_construir_contexto_progresivo` would call
    # the LLM to synthesize a narrative — mock it to keep this test offline and
    # deterministic (fails open if unmocked, but must never hit the network).
    mock_llm = AsyncMock()
    mock_llm.complete = AsyncMock(return_value=MagicMock(content="Narrativa de prueba."))
    with patch.object(informe_service, "get_llm", return_value=mock_llm):
        content2, _f2 = await informe_service.generar_informe_actividades_docx(db, user.id, cuenta2.id)
    doc2 = Document(io.BytesIO(content2))
    tabla2 = doc2.tables[1]
    body2 = "\n".join(cell.text for row in tabla2.rows for cell in row.cells)
    assert obligacion_unica.descripcion not in body2
    # The regular (not una_vez) obligación is still reported every cuota.
    assert obligaciones[1].descripcion in body2


# ── P4: consecutivo_ds non-fatal aviso (obtener_formato_valores) ────────────


async def test_formato_valores_organismo_sin_plantilla_ds_no_emite_aviso(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro
) -> None:
    """No DS plantilla ingested for the organismo → no consecutivo_ds aviso,
    generation proceeds unchanged (P4 scenario: non-DS organismo unaffected)."""
    user = test_user["user"]
    _valores, avisos = await informe_service.obtener_formato_valores(db, user.id, cuenta.id)
    assert not any("consecutivo DS" in a for a in avisos)


async def test_formato_valores_ds_requerido_consecutivo_nulo_emite_aviso_no_fatal(
    db: AsyncSession, test_user: dict[str, Any], cuenta: CuentaCobro, contrato: Contrato
) -> None:
    """DS plantilla exists for the organismo and `consecutivo_ds` is null →
    a clear non-fatal aviso is added, `valores` is still returned (fail-open,
    P4 scenario: missing consecutivo does not silently produce bad output)."""
    user = test_user["user"]
    plantilla = PlantillaOrganismo(
        usuario_id=user.id,
        entidad=contrato.entidad,
        entidad_normalizada=normalize(contrato.entidad),
        tipo_documento="documento_soporte",
        formato="xlsx",
        estructura_json={"clonable": True},
    )
    db.add(plantilla)
    await db.commit()
    assert cuenta.consecutivo_ds is None

    valores, avisos = await informe_service.obtener_formato_valores(db, user.id, cuenta.id)

    assert valores is not None
    assert "El Documento Soporte requiere el consecutivo DS; queda en blanco hasta que lo cargues." in avisos
