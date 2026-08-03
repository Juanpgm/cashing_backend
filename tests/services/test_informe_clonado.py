"""Tests for template-clone generation (docx clone, phase 3):
`informe_service._valores_plantilla` math, `_generar_docx_clonado` fail-open
fallback, and the cloned path wired into the public generators.
"""

from __future__ import annotations

import uuid
from datetime import date
from io import BytesIO
from typing import Any
from unittest.mock import AsyncMock

import pytest
import structlog
from app.models.actividad import Actividad
from app.models.contrato import Contrato
from app.models.cuenta_cobro import CuentaCobro, EstadoCuentaCobro, PosicionCuota
from app.models.documento_fuente import DocumentoFuente, TipoDocumentoFuente
from app.models.obligacion import Obligacion, TipoObligacion
from app.models.plantilla_organismo import PlantillaOrganismo
from app.services import informe_service
from app.services.informe_constants import SENTINEL_SIN_EVIDENCIAS
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

_CAMPOS = [
    {
        "direccion": "P0",
        "etiqueta": "Contrato No.",
        "campo": "contrato.numero_contrato",
        "valor_ejemplo": "XXX-000",
        "modo": "substring",
    }
]


def _build_template_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Contrato No. XXX-000")  # P0
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Fixtures ────────────────────────────────────────────────────────────────


@pytest.fixture
async def contrato(db: AsyncSession, test_user: dict[str, Any]) -> Contrato:
    user = test_user["user"]
    c = Contrato(
        usuario_id=user.id,
        numero_contrato="CTR-CLONE-001",
        objeto="Servicios profesionales de desarrollo",
        valor_total=12_000_000,
        valor_adicion=2_000_000,
        valor_mensual=1_000_000,
        fecha_inicio=date(2024, 1, 1),
        fecha_fin=date(2024, 12, 31),
        entidad="COEMPRESAR",
        supervisor_nombre="Carlos Supervisor",
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


@pytest.fixture
async def cuenta_tercera(db: AsyncSession, contrato: Contrato) -> CuentaCobro:
    """Cuota 3 with two prior cuotas (1M + 1.2M) for the acumulado/saldo math."""
    previas = [
        CuentaCobro(contrato_id=contrato.id, mes=1, anio=2024, valor=1_000_000, numero_cuota=1),
        CuentaCobro(contrato_id=contrato.id, mes=2, anio=2024, valor=1_200_000, numero_cuota=2),
    ]
    actual = CuentaCobro(
        contrato_id=contrato.id,
        mes=3,
        anio=2024,
        estado=EstadoCuentaCobro.BORRADOR,
        valor=1_000_000,
        numero_cuota=3,
    )
    db.add_all([*previas, actual])
    await db.commit()
    await db.refresh(actual)
    return actual


@pytest.fixture
async def plantilla_clonable(db: AsyncSession, contrato: Contrato) -> PlantillaOrganismo:
    doc = DocumentoFuente(
        usuario_id=contrato.usuario_id,
        contrato_id=contrato.id,
        storage_key="documentos/plantilla-coempresar.docx",
        nombre="plantilla-coempresar.docx",
        tipo=TipoDocumentoFuente.INFORME_ACTIVIDADES,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    p = PlantillaOrganismo(
        usuario_id=contrato.usuario_id,
        entidad="COEMPRESAR",
        entidad_normalizada="coempresar",
        tipo_documento="informe_actividades",
        formato="docx",
        estructura_json={
            "columnas": [],
            "secciones": [],
            "anexo_refs": [],
            "notas": "",
            "campos": _CAMPOS,
            "clonable": True,
            "avisos": [],
        },
        fuente_documento_id=doc.id,
    )
    db.add(p)
    await db.commit()
    await db.refresh(p)
    return p


async def _agregar_actividad(db: AsyncSession, contrato: Contrato, cuenta: CuentaCobro) -> None:
    ob = Obligacion(contrato_id=contrato.id, descripcion="Obligación única", tipo=TipoObligacion.GENERAL, orden=0)
    db.add(ob)
    await db.commit()
    await db.refresh(ob)
    contrato.obligaciones = [ob]
    db.add(
        Actividad(
            cuenta_cobro_id=cuenta.id,
            obligacion_id=ob.id,
            descripcion="Actividad realizada",
            justificacion="Justificación detallada",
            fecha_realizacion=date(2024, 3, 10),
        )
    )
    await db.commit()
    # Re-load the eagerly-cached (empty) actividades collection on the cuenta.
    await db.refresh(cuenta)


# ── _valores_plantilla math ─────────────────────────────────────────────────


async def test_valores_plantilla_acumulado_saldo_y_letras(
    db: AsyncSession, contrato: Contrato, cuenta_tercera: CuentaCobro
) -> None:
    valores = await informe_service._valores_plantilla(db, cuenta_tercera, contrato)

    assert valores["computed.acumulado"] == "$ 2.200.000"
    assert valores["computed.valor_total_con_adiciones"] == "$ 14.000.000"
    # 14.000.000 - 2.200.000 - 1.000.000
    assert valores["computed.saldo"] == "$ 10.800.000"
    assert valores["cuota.valor"] == "$ 1.000.000"
    assert valores["cuota.numero"] == "3"
    assert valores["cuota.periodo"] == "Marzo de 2024"
    assert valores["computed.cuota_ordinal_letras"] == "tercera"
    assert valores["computed.valor_letras"].endswith("PESOS M/CTE")
    assert valores["contratista.nombre"] == "Test User"
    assert valores["contratista.cedula"] == "123456789"
    # cargo_supervisor is None on this contrato → key omitted, not empty.
    assert "contrato.cargo_supervisor" not in valores
    # consecutivo_ds unset on this cuota (G4) → key omitted, not empty/None.
    assert "cuota.consecutivo_ds" not in valores


async def test_valores_plantilla_incluye_consecutivo_ds_cuando_esta_definido(
    db: AsyncSession, contrato: Contrato, cuenta_tercera: CuentaCobro
) -> None:
    """G4: cuota.consecutivo_ds resolves from CuentaCobro.consecutivo_ds when set
    (real SYJ packages radicate a "DS-4161-<consecutivo>" xlsx whose consecutive
    changes per cuota)."""
    cuenta_tercera.consecutivo_ds = "1768"
    db.add(cuenta_tercera)
    await db.commit()
    await db.refresh(cuenta_tercera)

    valores = await informe_service._valores_plantilla(db, cuenta_tercera, contrato)

    assert valores["cuota.consecutivo_ds"] == "1768"


# ── Fail-open fallback ──────────────────────────────────────────────────────


async def test_generador_cae_a_layout_legado_cuando_descarga_falla(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    test_user: dict[str, Any],
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
    plantilla_clonable: PlantillaOrganismo,
) -> None:
    await _agregar_actividad(db, contrato, cuenta_tercera)
    storage = AsyncMock()
    storage.download = AsyncMock(side_effect=RuntimeError("bucket down"))
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)
    # Keep the legacy path hermetic — no LLM for the progressive narrative.
    monkeypatch.setattr(informe_service, "_construir_contexto_progresivo", AsyncMock(return_value=None))

    content, filename = await informe_service.generar_informe_actividades_docx(
        db, test_user["user"].id, cuenta_tercera.id
    )

    assert filename.startswith("informe-actividades-")
    # Legacy built-from-scratch output — carries the BORRADOR header.
    doc = Document(BytesIO(content))
    assert doc.paragraphs[0].text == "BORRADOR — sujeto a revisión"


async def test_generar_docx_clonado_none_sin_fuente(
    db: AsyncSession, contrato: Contrato, cuenta_tercera: CuentaCobro, plantilla_clonable: PlantillaOrganismo
) -> None:
    plantilla_clonable.fuente_documento_id = None
    resultado = await informe_service._generar_docx_clonado(db, plantilla_clonable, cuenta_tercera, contrato, [], {})
    assert resultado is None


# ── Cloned happy path ───────────────────────────────────────────────────────


async def test_generador_devuelve_clon_sin_encabezado_borrador(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    test_user: dict[str, Any],
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
    plantilla_clonable: PlantillaOrganismo,
) -> None:
    await _agregar_actividad(db, contrato, cuenta_tercera)
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=_build_template_docx())
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    content, _filename = await informe_service.generar_informe_actividades_docx(
        db, test_user["user"].id, cuenta_tercera.id
    )

    doc = Document(BytesIO(content))
    assert doc.paragraphs[0].text == "Contrato No. CTR-CLONE-001"  # filled clone
    assert all("BORRADOR" not in p.text for p in doc.paragraphs)


# ── generar_cuenta_cobro_docx (clone-only) ──────────────────────────────────


async def test_generar_cuenta_cobro_docx_sin_plantilla_lanza_validation_error(
    db: AsyncSession, test_user: dict[str, Any], contrato: Contrato, cuenta_tercera: CuentaCobro
) -> None:
    from app.core.exceptions import ValidationError

    with pytest.raises(ValidationError):
        await informe_service.generar_cuenta_cobro_docx(db, test_user["user"].id, cuenta_tercera.id)


# ── _resolver_tipo_informe truth table (WU-1 / RELI-001) ────────────────────


def _cuenta_stub(final: bool) -> CuentaCobro:
    return CuentaCobro(
        informe_final=final,
        posicion=PosicionCuota.FINAL if final else PosicionCuota.RECURRENTE,
    )


class TestResolverTipoInforme:
    @pytest.mark.parametrize(
        ("marcado", "alternativo", "final", "esperado"),
        [
            # Refuter's failing scenario: template marked FINAL + cuenta final
            # must KEEP the marker where it is.
            ("INFORME FINAL", "INFORME PARCIAL", True, "actual"),
            ("INFORME FINAL", "INFORME PARCIAL", False, "alternativa"),
            ("INFORME PARCIAL", "INFORME FINAL", True, "alternativa"),
            ("INFORME PARCIAL", "INFORME FINAL", False, "actual"),
            # Only the alternative mentions final; missing marked text is not final.
            (None, "FINAL", False, "actual"),
            # Ambiguous — neither or both mention "final" → None (leave untouched).
            ("PARCIAL", "MENSUAL", True, None),
            ("FINAL", "INFORME FINAL", True, None),
            (None, None, False, None),
        ],
    )
    def test_truth_table(self, marcado: str | None, alternativo: str | None, final: bool, esperado: str | None) -> None:
        assert informe_service._resolver_tipo_informe(_cuenta_stub(final), marcado, alternativo) == esperado


# ── Cloned checkbox / blanking behavior (WU-1..WU-3) ────────────────────────


async def _crear_plantilla(db: AsyncSession, contrato: Contrato, campos: list[dict]) -> PlantillaOrganismo:
    doc = DocumentoFuente(
        usuario_id=contrato.usuario_id,
        contrato_id=contrato.id,
        storage_key="documentos/plantilla-campos.docx",
        nombre="plantilla-campos.docx",
        tipo=TipoDocumentoFuente.INFORME_ACTIVIDADES,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)
    p = PlantillaOrganismo(
        usuario_id=contrato.usuario_id,
        entidad="COEMPRESAR",
        entidad_normalizada="coempresar",
        tipo_documento="informe_actividades",
        formato="docx",
        estructura_json={
            "columnas": [],
            "secciones": [],
            "anexo_refs": [],
            "notas": "",
            "campos": campos,
            "clonable": True,
            "avisos": [],
        },
        fuente_documento_id=doc.id,
    )
    db.add(p)
    await db.commit()
    await db.refresh(p)
    return p


def _checkbox_docx(opcion_marcada: str, opcion_alternativa: str) -> bytes:
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = f"{opcion_marcada} X"
    t.cell(0, 1).text = opcion_alternativa
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


_CAMPO_CHECKBOX = {
    "direccion": "T0.R0.C0",
    "etiqueta": "Tipo de informe",
    "campo": "computed.tipo_informe",
    "valor_ejemplo": "X",
    "modo": "checkbox",
    "direccion_alternativa": "T0.R0.C1",
    "valor_alternativo": "X",
}


async def test_checkbox_final_marcado_y_cuenta_final_conserva_marcador(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
) -> None:
    """RELI-001 refuter scenario: template marked FINAL + cuenta final → "actual",
    the marker must NOT move to the PARCIAL cell."""
    cuenta_tercera.informe_final = True
    await db.commit()
    plantilla = await _crear_plantilla(db, contrato, [_CAMPO_CHECKBOX])
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=_checkbox_docx("INFORME FINAL", "INFORME PARCIAL"))
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    resultado = await informe_service._generar_docx_clonado(db, plantilla, cuenta_tercera, contrato, [], {})

    assert resultado is not None
    out = Document(BytesIO(resultado))
    assert out.tables[0].cell(0, 0).text == "INFORME FINAL X"  # marker untouched
    assert out.tables[0].cell(0, 1).text == "INFORME PARCIAL"  # no marker added


async def test_checkbox_parcial_marcado_y_cuenta_final_mueve_marcador(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
) -> None:
    cuenta_tercera.informe_final = True
    await db.commit()
    plantilla = await _crear_plantilla(db, contrato, [_CAMPO_CHECKBOX])
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=_checkbox_docx("INFORME PARCIAL", "INFORME FINAL"))
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    resultado = await informe_service._generar_docx_clonado(db, plantilla, cuenta_tercera, contrato, [], {})

    assert resultado is not None
    out = Document(BytesIO(resultado))
    assert "X" not in out.tables[0].cell(0, 0).text  # marker removed
    assert out.tables[0].cell(0, 1).text == "INFORME FINAL X"  # marker moved


async def test_checkbox_ambiguo_no_toca_marcador_y_avisa(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
) -> None:
    """Neither option mentions "final" → the campo is omitted from valores, the
    template marker stays untouched and a warning is logged."""
    plantilla = await _crear_plantilla(db, contrato, [_CAMPO_CHECKBOX])
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=_checkbox_docx("OPCION A", "OPCION B"))
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    with structlog.testing.capture_logs() as captured:
        resultado = await informe_service._generar_docx_clonado(db, plantilla, cuenta_tercera, contrato, [], {})

    assert resultado is not None
    out = Document(BytesIO(resultado))
    assert out.tables[0].cell(0, 0).text == "OPCION A X"
    assert out.tables[0].cell(0, 1).text == "OPCION B"
    eventos = [e for e in captured if e.get("event") == "informe_clonado_tipo_informe_ambiguo"]
    assert eventos and "marcar manualmente" in eventos[0]["aviso"]


def _obligaciones_con_actividades(contrato: Contrato, n: int) -> tuple[list[Actividad], dict[uuid.UUID, Obligacion]]:
    obligaciones_by_id: dict[uuid.UUID, Obligacion] = {}
    actividades: list[Actividad] = []
    for i in range(1, n + 1):
        ob = Obligacion(
            id=uuid.uuid4(),
            contrato_id=contrato.id,
            descripcion=f"Obligación {i}",
            tipo=TipoObligacion.GENERAL,
            orden=i,
        )
        obligaciones_by_id[ob.id] = ob
        actividades.append(
            Actividad(id=uuid.uuid4(), obligacion_id=ob.id, descripcion=f"Trabajo obligación {i}", justificacion=None)
        )
    return actividades, obligaciones_by_id


async def test_justificaciones_sobrantes_quedan_en_blanco(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
) -> None:
    """RELI-002: template with 3 justificacion campos + contract with 2
    obligaciones → the 3rd row's example narrative is blanked, not retained.

    Both actividades here have no evidencia and no justificación (billing-
    resilience-templates E2E finding A/B fix) — the sentinel renders instead
    of echoing `act.descripcion`."""
    doc = Document()
    t = doc.add_table(rows=3, cols=1)
    t.cell(0, 0).text = "Narrativa ejemplo uno"
    t.cell(1, 0).text = "Narrativa ejemplo dos"
    t.cell(2, 0).text = "Narrativa ejemplo tres de otra radicación"
    buf = BytesIO()
    doc.save(buf)
    campos = [
        {
            "direccion": f"T0.R{i}.C0",
            "etiqueta": f"Justificación {i + 1}",
            "campo": f"justificacion.{i + 1}",
            "valor_ejemplo": "Narrativa ejemplo",
            "modo": "justificacion",
        }
        for i in range(3)
    ]
    plantilla = await _crear_plantilla(db, contrato, campos)
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=buf.getvalue())
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)
    actividades, obligaciones_by_id = _obligaciones_con_actividades(contrato, 2)

    with structlog.testing.capture_logs() as captured:
        resultado = await informe_service._generar_docx_clonado(
            db, plantilla, cuenta_tercera, contrato, actividades, obligaciones_by_id
        )

    assert resultado is not None
    out = Document(BytesIO(resultado))
    assert out.tables[0].cell(0, 0).text == SENTINEL_SIN_EVIDENCIAS
    assert out.tables[0].cell(1, 0).text == SENTINEL_SIN_EVIDENCIAS
    assert out.tables[0].cell(2, 0).text == ""  # example narrative blanked
    eventos = [e for e in captured if e.get("event") == "informe_clonado_justificaciones_blanqueadas"]
    assert eventos and eventos[0]["n"] == 1


async def test_generar_docx_clonado_justificacion_evidencia_sentinel_prioridad(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
) -> None:
    """E2E finding A/B fix — per obligación, `justificacion.{n}` renders:
    real justificación text when present (ob 1); `act.descripcion` when
    justificación is empty but the obligación HAS evidencia (ob 2, unchanged
    prior behavior); the sentinel when there is NEITHER (ob 3)."""
    from app.models.evidencia import Evidencia

    doc = Document()
    doc.add_table(rows=3, cols=1)
    buf = BytesIO()
    doc.save(buf)
    campos = [
        {
            "direccion": f"T0.R{i}.C0",
            "etiqueta": f"Justificación {i + 1}",
            "campo": f"justificacion.{i + 1}",
            "valor_ejemplo": "",
            "modo": "justificacion",
        }
        for i in range(3)
    ]
    plantilla = await _crear_plantilla(db, contrato, campos)
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=buf.getvalue())
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    obs = [
        Obligacion(
            id=uuid.uuid4(),
            contrato_id=contrato.id,
            descripcion=f"Obligación {i}",
            tipo=TipoObligacion.GENERAL,
            orden=i,
        )
        for i in range(1, 4)
    ]
    obligaciones_by_id = {ob.id: ob for ob in obs}

    act_justificada = Actividad(
        id=uuid.uuid4(), obligacion_id=obs[0].id, descripcion="Trabajo 1", justificacion="Justificación real 1"
    )
    act_con_evidencia = Actividad(id=uuid.uuid4(), obligacion_id=obs[1].id, descripcion="Trabajo 2", justificacion=None)
    act_con_evidencia.evidencias = [
        Evidencia(id=uuid.uuid4(), storage_key="k", nombre_archivo="f.jpg", tipo_archivo="image/jpeg", tamano_bytes=1)
    ]
    act_sin_nada = Actividad(id=uuid.uuid4(), obligacion_id=obs[2].id, descripcion="Trabajo 3", justificacion=None)

    resultado = await informe_service._generar_docx_clonado(
        db, plantilla, cuenta_tercera, contrato, [act_justificada, act_con_evidencia, act_sin_nada], obligaciones_by_id
    )

    assert resultado is not None
    out = Document(BytesIO(resultado))
    assert out.tables[0].cell(0, 0).text == "Justificación real 1"
    assert out.tables[0].cell(1, 0).text == "Trabajo 2"
    assert out.tables[0].cell(2, 0).text == SENTINEL_SIN_EVIDENCIAS


async def test_campos_sin_dato_quedan_en_blanco(
    db: AsyncSession,
    monkeypatch: pytest.MonkeyPatch,
    contrato: Contrato,
    cuenta_tercera: CuentaCobro,
) -> None:
    """RESI-001: a nullable contrato field (supervisor_nombre=None) must blank
    the example's real data instead of silently retaining it."""
    contrato.supervisor_nombre = None
    await db.commit()
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "Supervisor"
    t.cell(0, 1).text = "Pepito Ejemplo Pérez"
    buf = BytesIO()
    doc.save(buf)
    campos = [
        {
            "direccion": "T0.R0.C1",
            "etiqueta": "Supervisor",
            "campo": "contrato.supervisor_nombre",
            "valor_ejemplo": "Pepito Ejemplo Pérez",
            "modo": "cell",
        }
    ]
    plantilla = await _crear_plantilla(db, contrato, campos)
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=buf.getvalue())
    monkeypatch.setattr(informe_service, "_get_storage", lambda *_a, **_k: storage)

    with structlog.testing.capture_logs() as captured:
        resultado = await informe_service._generar_docx_clonado(db, plantilla, cuenta_tercera, contrato, [], {})

    assert resultado is not None
    out = Document(BytesIO(resultado))
    assert out.tables[0].cell(0, 1).text == ""  # example supervisor blanked
    assert all("Pepito" not in c.text for row in out.tables[0].rows for c in row.cells)
    eventos = [e for e in captured if e.get("event") == "informe_clonado_campos_blanqueados"]
    assert eventos and eventos[0]["campos"] == ["Supervisor"]
