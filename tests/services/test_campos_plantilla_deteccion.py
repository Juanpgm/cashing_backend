"""Tests for clonable-template campo detection (docx clone, phase 2):
`requisito_inference_service.detectar_campos_plantilla` and its wiring into
`ingerir_plantilla_organismo` (persisted `estructura_json` keys, tipo
PLANTILLA → "cuenta_cobro" mapping).
"""

from __future__ import annotations

import json
from datetime import date
from io import BytesIO
from typing import Any
from unittest.mock import AsyncMock

import pytest
from app.models.contrato import Contrato
from app.models.documento_fuente import DocumentoFuente, TipoDocumentoFuente
from app.schemas.agent import LLMResponse
from app.schemas.plantilla_organismo import CamposPlantillaLLM, EstructuraPlantillaLLM
from app.services import document_service
from app.services import requisito_inference_service as svc
from docx import Document
from sqlalchemy.ext.asyncio import AsyncSession

_CAMPO_VALIDO = {
    "direccion": "P0",
    "etiqueta": "Contrato No.",
    "campo": "contrato.numero_contrato",
    "valor_ejemplo": "ABC-123",
    "modo": "substring",
}

_CAMPO_BOGUS = {
    "direccion": "T9.R0.C0",
    "etiqueta": "Fantasma",
    "campo": "cuota.valor",
    "valor_ejemplo": "$ 9.999.999",
    "modo": "cell",
}

_CAMPOS_JSON = json.dumps({"campos": [_CAMPO_VALIDO, _CAMPO_BOGUS]})

_ESTRUCTURA_JSON = json.dumps(
    {"columnas": [], "secciones": [], "anexo_refs": [], "notas": "Carta de cuenta de cobro COEMPRESAR."}
)


def _build_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Contrato No. ABC-123")  # P0
    doc.add_paragraph("Valor Cuota a cancelar: $ 4.800.000")  # P1
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class _FakeLLMPorSchema:
    """Return schema-appropriate content: ingestion issues TWO structured calls
    (estructura + campos) against the same patched `get_llm`."""

    def __init__(self, por_schema: dict[type, str]) -> None:
        self._por_schema = por_schema

    async def complete(self, messages, temperature=0.0, max_tokens=4096, response_format=None, **kwargs) -> LLMResponse:
        return LLMResponse(content=self._por_schema[response_format], model="fake/test-model", total_tokens=30)


def _patch_llm(monkeypatch: pytest.MonkeyPatch, por_schema: dict[type, str]) -> None:
    import app.adapters.llm as llm_pkg

    monkeypatch.setattr(llm_pkg, "get_llm", lambda model=None: _FakeLLMPorSchema(por_schema), raising=True)


def _fake_storage(content: bytes) -> AsyncMock:
    storage = AsyncMock()
    storage.download = AsyncMock(return_value=content)
    return storage


# ── detectar_campos_plantilla ─────────────────────────────────────────────────


async def test_detectar_campos_filtra_deterministicamente(monkeypatch: pytest.MonkeyPatch) -> None:
    _patch_llm(monkeypatch, {CamposPlantillaLLM: _CAMPOS_JSON})

    campos, avisos = await svc.detectar_campos_plantilla(_build_docx())

    assert [c["campo"] for c in campos] == ["contrato.numero_contrato"]
    assert len(avisos) == 1
    assert "Fantasma" in avisos[0]


async def test_detectar_campos_usa_max_tokens_amplio(monkeypatch: pytest.MonkeyPatch) -> None:
    """Regression: real skeletons overflow 4096 completion tokens (valor_ejemplo is
    echoed verbatim), truncating the JSON. The call must request a wider budget."""
    captured: dict[str, Any] = {}

    class _CapturingLLM:
        async def complete(self, messages, **kwargs: Any) -> LLMResponse:
            captured.update(kwargs)
            return LLMResponse(content=_CAMPOS_JSON, model="fake/test-model", total_tokens=30)

    import app.adapters.llm as llm_pkg

    monkeypatch.setattr(llm_pkg, "get_llm", lambda model=None: _CapturingLLM(), raising=True)

    await svc.detectar_campos_plantilla(_build_docx())
    assert captured["max_tokens"] >= 16384


async def test_detectar_campos_degrada_en_error_llm(monkeypatch: pytest.MonkeyPatch) -> None:
    class _BoomLLM:
        async def complete(self, *a: Any, **kw: Any) -> LLMResponse:
            raise RuntimeError("llm down")

    import app.adapters.llm as llm_pkg

    monkeypatch.setattr(llm_pkg, "get_llm", lambda model=None: _BoomLLM(), raising=True)

    campos, avisos = await svc.detectar_campos_plantilla(_build_docx())
    assert campos == []
    assert len(avisos) == 1


async def test_detectar_campos_degrada_con_bytes_invalidos() -> None:
    campos, avisos = await svc.detectar_campos_plantilla(b"not-a-docx")
    assert campos == []
    assert len(avisos) == 1


# ── ingestion wiring: persisted keys + tipo PLANTILLA mapping ────────────────


@pytest.fixture
async def contrato_coempresar(db: AsyncSession, test_user: dict[str, Any]) -> Contrato:
    user = test_user["user"]
    c = Contrato(
        usuario_id=user.id,
        numero_contrato="ABC-123",
        objeto="Servicios profesionales — organismo COEMPRESAR",
        valor_total=12_000_000,
        valor_mensual=1_000_000,
        fecha_inicio=date(2024, 1, 1),
        fecha_fin=date(2024, 12, 31),
        entidad="COEMPRESAR",
    )
    db.add(c)
    await db.commit()
    await db.refresh(c)
    return c


async def test_ingestion_docx_persists_campos_clonable_avisos(
    db: AsyncSession, monkeypatch: pytest.MonkeyPatch, contrato_coempresar: Contrato
) -> None:
    doc = DocumentoFuente(
        usuario_id=contrato_coempresar.usuario_id,
        contrato_id=contrato_coempresar.id,
        storage_key="documentos/cuenta-cobro-coempresar.docx",
        nombre="cuenta-cobro-coempresar.docx",
        tipo=TipoDocumentoFuente.PLANTILLA,
    )
    db.add(doc)
    await db.commit()
    await db.refresh(doc)

    monkeypatch.setattr(document_service, "extraer_texto_documento", AsyncMock(return_value=("texto carta", [])))
    monkeypatch.setattr("app.adapters.storage.get_storage", lambda *_a, **_k: _fake_storage(_build_docx()))
    _patch_llm(monkeypatch, {EstructuraPlantillaLLM: _ESTRUCTURA_JSON, CamposPlantillaLLM: _CAMPOS_JSON})

    plantilla, avisos = await svc.ingerir_plantilla_organismo(
        db, contrato_coempresar.usuario_id, contrato_coempresar.id, doc.id
    )
    await db.commit()

    assert plantilla is not None
    assert avisos == []
    # tipo PLANTILLA maps to the cuenta-de-cobro letter family.
    assert plantilla.tipo_documento == "cuenta_cobro"
    assert plantilla.formato == "docx"
    estructura = plantilla.estructura_json
    assert estructura["clonable"] is True
    assert [c["campo"] for c in estructura["campos"]] == ["contrato.numero_contrato"]
    assert len(estructura["avisos"]) == 1  # the bogus campo was discarded with an aviso
    assert estructura["notas"] == "Carta de cuenta de cobro COEMPRESAR."
