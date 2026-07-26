"""Tests for docx_clone_service pure functions (skeleton, address resolution, filling).

Builds a small in-memory docx with python-docx — including a paragraph whose
text is split across multiple runs mid-value, the classic Word fragmentation
case — and round-trips it through extraer_esqueleto / validar_campos / rellenar.
"""

from __future__ import annotations

from io import BytesIO

from app.services.docx_clone_service import (
    extraer_esqueleto,
    rellenar,
    resolver_direccion,
    validar_campos,
)
from docx import Document

CAMPOS = [
    {
        "direccion": "P0",
        "etiqueta": "Número de contrato",
        "campo": "contrato.numero",
        "valor_ejemplo": "ABC-123",
        "modo": "substring",
    },
    {
        "direccion": "T0.R0.C1",
        "etiqueta": "Valor cuota",
        "campo": "cuota.valor",
        "valor_ejemplo": "$ 4.800.000",
        "modo": "cell",
    },
    {
        "direccion": "P1",
        "etiqueta": "Valor en texto",
        "campo": "cuota.valor_texto",
        "valor_ejemplo": "$ 4.800.000",
        "modo": "substring",
    },
    {
        "direccion": "T0.R1.C0",
        "etiqueta": "Tipo de cuota",
        "campo": "cuota.tipo",
        "valor_ejemplo": "X",
        "modo": "checkbox",
        "direccion_alternativa": "T0.R1.C1",
        "valor_alternativo": "X",
    },
    {
        "direccion": "T1.R0.C0",
        "etiqueta": "Justificación",
        "campo": "cuota.justificacion",
        "valor_ejemplo": "Justificación pendiente",
        "modo": "justificacion",
    },
]

CAMPO_BOGUS = {
    "direccion": "T9.R0.C0",
    "etiqueta": "Fantasma",
    "campo": "fantasma",
    "valor_ejemplo": "nada",
    "modo": "cell",
}


def _build_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Contrato No. ABC-123")  # P0
    # P1 — value fragmented across three runs, as Word often does.
    p = doc.add_paragraph()
    p.add_run("Valor: $ 4.8")
    p.add_run("00").bold = True
    p.add_run(".000")
    # T0 — label/value pair + checkbox pair, with a nested table in cell (0,0).
    t0 = doc.add_table(rows=2, cols=2)
    t0.style = "Table Grid"
    t0.cell(0, 0).text = "Cuota"
    t0.cell(0, 1).text = "$ 4.800.000"
    t0.cell(1, 0).text = "PARCIAL X"
    t0.cell(1, 1).text = "FINAL"
    nested = t0.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "Anidada"
    # T1 — justification cell.
    t1 = doc.add_table(rows=1, cols=1)
    t1.cell(0, 0).text = "Justificación pendiente"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestExtraerEsqueleto:
    def test_emits_addressed_lines_in_body_order(self) -> None:
        esq = extraer_esqueleto(_build_docx())
        lines = esq.splitlines()
        assert "P0: Contrato No. ABC-123" in lines
        assert "P1: Valor: $ 4.800.000" in lines  # fragmented runs concatenated
        assert "T0.R0.C1: $ 4.800.000" in lines
        assert "T0.R1.C0: PARCIAL X" in lines
        assert "T0.R0.C0.T0.R0.C0: Anidada" in lines  # nested table recursion
        assert "T1.R0.C0: Justificación pendiente" in lines
        assert lines.index("P0: Contrato No. ABC-123") < lines.index("T0.R0.C1: $ 4.800.000")

    def test_truncates_at_line_boundary_with_marker(self) -> None:
        esq = extraer_esqueleto(_build_docx(), max_chars=30)
        lines = esq.splitlines()
        assert lines[-1] == "... [esqueleto truncado]"
        assert all(len(line) <= 30 for line in lines[:-1])


class TestResolverDireccion:
    def test_resolves_paragraphs_cells_and_nested_cells(self) -> None:
        doc = Document(BytesIO(_build_docx()))
        assert resolver_direccion(doc, "P0").text == "Contrato No. ABC-123"
        assert resolver_direccion(doc, "T0.R1.C1").text == "FINAL"
        assert resolver_direccion(doc, "T0.R0.C0.T0.R0.C0").text == "Anidada"

    def test_returns_none_on_out_of_range_or_malformed(self) -> None:
        doc = Document(BytesIO(_build_docx()))
        for direccion in ("P99", "T9.R0.C0", "T0.R9.C0", "T0.R0.C9", "T0.R0", "garbage", "", "P-1"):
            assert resolver_direccion(doc, direccion) is None


class TestValidarCampos:
    def test_keeps_resolvable_campos_and_drops_bogus_with_aviso(self) -> None:
        validos, avisos = validar_campos(_build_docx(), [*CAMPOS, CAMPO_BOGUS])
        assert validos == CAMPOS
        assert len(avisos) == 1
        assert "Fantasma" in avisos[0]
        assert "T9.R0.C0" in avisos[0]

    def test_drops_campo_whose_valor_ejemplo_is_absent(self) -> None:
        campo = {**CAMPOS[1], "valor_ejemplo": "$ 9.999.999"}
        validos, avisos = validar_campos(_build_docx(), [campo])
        assert validos == []
        assert len(avisos) == 1

    def test_drops_campo_contenido_en_sibling_misma_direccion(self) -> None:
        """Dedup: at one direccion, a valor_ejemplo contained in a sibling's would
        be consumed by the sibling's fill — the shorter campo is dropped with an
        aviso; an equal campo at ANOTHER direccion is untouched."""
        doc = Document()
        doc.add_paragraph("Periodo: Abril 28 de 2026")  # P0
        doc.add_paragraph("Año: 2026")  # P1
        buf = BytesIO()
        doc.save(buf)
        periodo = {
            "direccion": "P0",
            "etiqueta": "Periodo",
            "campo": "cuota.periodo",
            "valor_ejemplo": "Abril 28 de 2026",
            "modo": "substring",
        }
        anio_p0 = {
            "direccion": "P0",
            "etiqueta": "Año",
            "campo": "cuota.anio",
            "valor_ejemplo": "2026",
            "modo": "substring",
        }
        anio_p1 = {**anio_p0, "direccion": "P1"}

        validos, avisos = validar_campos(buf.getvalue(), [anio_p0, periodo, anio_p1])

        assert validos == [periodo, anio_p1]
        assert len(avisos) == 1
        assert "'Año'" in avisos[0] and "'Periodo'" in avisos[0]

    def test_checkbox_requires_resolvable_direccion_alternativa(self) -> None:
        campo = {**CAMPOS[3], "direccion_alternativa": "T9.R9.C9"}
        validos, avisos = validar_campos(_build_docx(), [campo])
        assert validos == []
        assert len(avisos) == 1


class TestRellenar:
    def test_fills_all_modos_and_preserves_structure(self) -> None:
        valores = {
            "cuota.valor": "$ 5.200.000",
            "cuota.valor_texto": "$ 5.200.000",
            "cuota.tipo": "alternativa",
            "cuota.justificacion": "Primer párrafo de la justificación.\n\nSegundo párrafo.\nTercera línea.",
            # "contrato.numero" intentionally absent → skipped silently
        }
        out = Document(BytesIO(rellenar(_build_docx(), CAMPOS, valores)))

        # substring across fragmented runs
        assert out.paragraphs[1].text == "Valor: $ 5.200.000"
        # campo without a value is untouched
        assert out.paragraphs[0].text == "Contrato No. ABC-123"
        # cell modo replaces the whole cell text
        assert out.tables[0].cell(0, 1).text == "$ 5.200.000"
        # checkbox swapped from PARCIAL to FINAL
        assert "X" not in out.tables[0].cell(1, 0).text
        assert out.tables[0].cell(1, 1).text == "FINAL X"
        # justificacion written as multiple paragraphs
        just = [p.text for p in out.tables[1].cell(0, 0).paragraphs if p.text]
        assert just == ["Primer párrafo de la justificación.", "Segundo párrafo.", "Tercera línea."]
        # structure/styles survived the round trip
        assert len(out.tables) == 2
        assert out.tables[0].style.name == "Table Grid"
        assert out.tables[0].cell(0, 0).tables[0].cell(0, 0).text == "Anidada"

    def test_substring_tolerates_whitespace_runs(self) -> None:
        """Regression: skeleton lines are whitespace-collapsed, so the LLM's
        valor_ejemplo has single spaces where the document may have runs of
        spaces/tabs (real case: DAGMA 'GESTION DEL MEDIO<17 spaces>AMBIENTE').
        validar_campos accepts those campos (it normalizes); rellenar must too."""
        doc = Document()
        doc.add_paragraph("Organismo: DEPARTAMENTO DE GESTION DEL MEDIO                 AMBIENTE")
        buf = BytesIO()
        doc.save(buf)
        campo = {
            "direccion": "P0",
            "etiqueta": "Organismo",
            "campo": "contrato.entidad",
            "valor_ejemplo": "DEPARTAMENTO DE GESTION DEL MEDIO AMBIENTE",
            "modo": "substring",
        }
        validos, avisos = validar_campos(buf.getvalue(), [campo])
        assert validos == [campo] and avisos == []
        out = Document(BytesIO(rellenar(buf.getvalue(), [campo], {"contrato.entidad": "ENTIDAD DE PRUEBA"})))
        assert out.paragraphs[0].text == "Organismo: ENTIDAD DE PRUEBA"

    def test_substring_fills_longest_valor_ejemplo_first(self) -> None:
        """Safety net: even listed short-first, the longer overlapping substring
        campo fills first so the short one still finds its own occurrence."""
        doc = Document()
        doc.add_paragraph("Periodo: Abril 28 de 2026 — vigencia 2026")
        buf = BytesIO()
        doc.save(buf)
        campos = [
            {"direccion": "P0", "etiqueta": "Año", "campo": "cuota.anio", "valor_ejemplo": "2026", "modo": "substring"},
            {
                "direccion": "P0",
                "etiqueta": "Periodo",
                "campo": "cuota.periodo",
                "valor_ejemplo": "Abril 28 de 2026",
                "modo": "substring",
            },
        ]
        out = Document(BytesIO(rellenar(buf.getvalue(), campos, {"cuota.anio": "2099", "cuota.periodo": "PERIODO-X"})))
        assert out.paragraphs[0].text == "Periodo: PERIODO-X — vigencia 2099"

    def test_checkbox_actual_is_a_noop(self) -> None:
        out = Document(BytesIO(rellenar(_build_docx(), CAMPOS, {"cuota.tipo": "actual"})))
        assert out.tables[0].cell(1, 0).text == "PARCIAL X"
        assert out.tables[0].cell(1, 1).text == "FINAL"
