"""Docx clone service — pure functions to clone-and-fill entity Word templates.

Edits the original document in place (via python-docx) so letterhead, logos and
styles survive. No DB, no LLM — callers pass docx bytes plus a list of campo
dicts and get bytes back.

Address scheme ("direccion") over the document BODY (headers/footers excluded):
- ``P{i}``            — i-th body paragraph (0-based, body order)
- ``T{t}.R{r}.C{c}``  — cell c of row r of body table t (0-based)
- nested tables chain: ``T{t}.R{r}.C{c}.T{t2}.R{r2}.C{c2}`` (tables indexed
  within their containing cell)

Campo dict contract (produced upstream, consumed here):
    {"direccion": "T2.R1.C3", "etiqueta": "Valor Cuota a cancelar",
     "campo": "cuota.valor", "valor_ejemplo": "$ 4.800.000",
     "modo": "cell" | "substring" | "checkbox" | "justificacion",
     # checkbox only:
     "direccion_alternativa": "T0.R1.C2", "valor_alternativo": "X"}
"""

from __future__ import annotations

import re
from collections.abc import Iterator
from io import BytesIO

from docx import Document
from docx.document import Document as DocxDocument  # real class; `Document` above is a factory
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.core.text_match import normalize

_TRUNCATION_MARKER = "... [esqueleto truncado]"

_P_RE = re.compile(r"P(\d+)")
_T_RE = re.compile(r"T(\d+)")
_R_RE = re.compile(r"R(\d+)")
_C_RE = re.compile(r"C(\d+)")


def _norm(text: str | None) -> str:
    """Accent/case-insensitive, whitespace-collapsed form for tolerant matching."""
    return " ".join(normalize(text).split())


def _clean(text: str) -> str:
    """Collapse internal whitespace for skeleton lines."""
    return " ".join(text.split())


def _iter_table(table: Table, prefix: str) -> Iterator[tuple[str, str]]:
    """Yield ``(direccion, text)`` for every non-empty cell, recursing into nested tables.

    ponytail: merged cells repeat the same _Cell across the row span, so a merged
    cell can emit duplicate lines — addresses still resolve, dedupe when it hurts.
    """
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            direccion = f"{prefix}.R{r}.C{c}"
            text = _clean(" ".join(p.text for p in cell.paragraphs))
            if text:
                yield direccion, text
            for t2, nested in enumerate(cell.tables):
                yield from _iter_table(nested, f"{direccion}.T{t2}")


def _iter_nodes(doc: DocxDocument) -> Iterator[tuple[str, str]]:
    """Yield ``(direccion, text)`` for every non-empty body node in document order."""
    p_idx = t_idx = 0
    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            text = _clean(item.text)
            if text:
                yield f"P{p_idx}", text
            p_idx += 1
        elif isinstance(item, Table):
            yield from _iter_table(item, f"T{t_idx}")
            t_idx += 1


def extraer_esqueleto(docx_bytes: bytes, max_chars: int = 14_000) -> str:
    """Render the document body as one ``direccion: text`` line per non-empty node.

    Headers/footers are skipped. Output is cut at a line boundary once it would
    exceed ``max_chars``, appending a truncation marker line.
    """
    doc = Document(BytesIO(docx_bytes))
    lines: list[str] = []
    total = 0
    for direccion, text in _iter_nodes(doc):
        line = f"{direccion}: {text}"
        if total + len(line) + 1 > max_chars:
            lines.append(_TRUNCATION_MARKER)
            break
        lines.append(line)
        total += len(line) + 1
    return "\n".join(lines)


def resolver_direccion(doc: DocxDocument, direccion: str) -> Paragraph | _Cell | None:
    """Resolve an address to its Paragraph or _Cell; None on malformed/out-of-range."""
    try:
        direccion = (direccion or "").strip()
        m = _P_RE.fullmatch(direccion)
        if m:
            paragraphs = doc.paragraphs
            i = int(m.group(1))
            return paragraphs[i] if i < len(paragraphs) else None
        parts = direccion.split(".")
        if not parts or len(parts) % 3 != 0:
            return None
        container: DocxDocument | _Cell = doc
        cell: _Cell | None = None
        for j in range(0, len(parts), 3):
            tm = _T_RE.fullmatch(parts[j])
            rm = _R_RE.fullmatch(parts[j + 1])
            cm = _C_RE.fullmatch(parts[j + 2])
            if not (tm and rm and cm):
                return None
            tables = container.tables
            t = int(tm.group(1))
            if t >= len(tables):
                return None
            rows = tables[t].rows
            r = int(rm.group(1))
            if r >= len(rows):
                return None
            cells = rows[r].cells
            c = int(cm.group(1))
            if c >= len(cells):
                return None
            cell = cells[c]
            container = cell
        return cell
    except Exception:  # defensive: address input is untrusted (LLM-produced)
        return None


def texto_direccion(doc: DocxDocument, direccion: str) -> str | None:
    """Actual text of the node at ``direccion``, or None when it doesn't resolve.

    Used by checkbox resolution: the campo contract's ``valor_ejemplo`` is only
    the marker glyph ("X"), so callers that need the OPTION TEXTS (e.g. to tell
    the PARCIAL cell from the FINAL cell) must read them from the document.
    """
    node = resolver_direccion(doc, direccion)
    return node.text if node is not None else None


def validar_campos(docx_bytes: bytes, campos: list[dict]) -> tuple[list[dict], list[str]]:
    """Keep only campos whose direccion resolves and whose valor_ejemplo appears there.

    Checkbox campos additionally require a resolvable ``direccion_alternativa``.
    Returns ``(campos_validos, avisos)`` — avisos are user-facing Spanish strings.
    """
    doc = Document(BytesIO(docx_bytes))
    validos: list[dict] = []
    avisos: list[str] = []
    for campo in campos:
        etiqueta = campo.get("etiqueta") or campo.get("campo") or "?"
        direccion = campo.get("direccion", "")
        node = resolver_direccion(doc, direccion)
        if node is None:
            avisos.append(f"Se descartó el campo '{etiqueta}': la dirección '{direccion}' no existe en la plantilla.")
            continue
        valor_ejemplo = campo.get("valor_ejemplo", "")
        if _norm(valor_ejemplo) not in _norm(node.text):
            avisos.append(
                f"Se descartó el campo '{etiqueta}': el valor de ejemplo '{valor_ejemplo}' no aparece en '{direccion}'."
            )
            continue
        if campo.get("modo") == "checkbox":
            alternativa = campo.get("direccion_alternativa", "")
            if resolver_direccion(doc, alternativa) is None:
                avisos.append(
                    f"Se descartó el campo '{etiqueta}': la dirección alternativa "
                    f"'{alternativa}' no existe en la plantilla."
                )
                continue
        validos.append(campo)
    # Dedup: within the same direccion, a campo whose valor_ejemplo is contained
    # in a sibling's would be consumed by the sibling's fill (rellenar replaces
    # the first match) — keep the longer campo deterministically.
    # ponytail: O(n²) pairwise scan, fine for tens of campos per template.
    finales: list[dict] = []
    for campo in validos:
        contenedor = None
        if campo.get("modo", "cell") in ("cell", "substring"):
            na = _norm(campo.get("valor_ejemplo", ""))
            for otro in validos:
                if otro is campo or otro.get("modo", "cell") not in ("cell", "substring"):
                    continue
                nb = _norm(otro.get("valor_ejemplo", ""))
                if otro.get("direccion") == campo.get("direccion") and na != nb and na in nb:
                    contenedor = otro
                    break
        if contenedor is not None:
            avisos.append(
                f"Campo '{campo.get('etiqueta') or campo.get('campo') or '?'}' descartado: su valor de ejemplo "
                f"está contenido en el de '{contenedor.get('etiqueta') or contenedor.get('campo') or '?'}' "
                "en la misma ubicación."
            )
            continue
        finales.append(campo)
    return finales, avisos


def _replace_text_preserving_runs(paragraph: Paragraph, old: str, new: str) -> bool:
    """Replace the first occurrence of ``old`` across run boundaries, keeping styles.

    Runs are concatenated to locate ``old`` (exact first, then case-insensitive,
    then whitespace-run tolerant);
    the inserted text takes the style of the run where the match starts, middle
    runs are blanked, boundary runs keep their surrounding text. Returns False
    when ``old`` is not found.
    """
    runs = paragraph.runs
    full = "".join(run.text for run in runs)
    idx = full.find(old)
    if idx < 0:
        idx = full.lower().find(old.lower())
    if idx >= 0:
        end = idx + len(old)
    else:
        # Skeleton lines are whitespace-collapsed, so an LLM valor_ejemplo uses
        # single spaces where the document may have runs of spaces/tabs.
        tokens = old.split()
        m = re.search(r"\s+".join(re.escape(t) for t in tokens), full, re.IGNORECASE) if tokens else None
        if m is None:
            return False
        idx, end = m.span()
    pos = 0
    started = False
    for run in runs:
        run_len = len(run.text)
        if end <= pos:
            break
        if idx >= pos + run_len:
            pos += run_len
            continue
        local_start = max(idx - pos, 0)
        local_end = min(end - pos, run_len)
        if not started:
            run.text = run.text[:local_start] + new + run.text[local_end:]
            started = True
        else:
            run.text = run.text[local_end:]
        pos += run_len
    return True


def _node_paragraphs(node: Paragraph | _Cell) -> list[Paragraph]:
    return node.paragraphs if isinstance(node, _Cell) else [node]


def _set_paragraph_text(paragraph: Paragraph, new: str) -> None:
    """Set the paragraph's whole text keeping the first run's style; blank the rest."""
    runs = paragraph.runs
    if runs:
        runs[0].text = new
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new)


def _blank_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.text = ""


def rellenar(docx_bytes: bytes, campos: list[dict], valores: dict[str, str]) -> bytes:
    """Fill the template: apply each campo whose ``campo`` key has an entry in ``valores``.

    Campos without a value (or whose direccion no longer resolves) are skipped
    silently. Modo semantics:
    - ``cell``: the node's entire text becomes the value (first paragraph keeps
      its first run's style; other runs/paragraphs are blanked).
    - ``substring``: replace ``valor_ejemplo`` in place across fragmented runs.
    - ``checkbox``: the value is the literal string ``"actual"`` (leave the
      marker where it is — no-op) or ``"alternativa"`` (remove the marker
      ``valor_ejemplo`` from ``direccion`` and write ``valor_alternativo`` into
      ``direccion_alternativa``).
    - ``justificacion``: the node (a cell) gets multi-paragraph text — value is
      split on blank lines then newlines; first line keeps the cell style,
      remaining lines become new paragraphs.
    """
    doc = Document(BytesIO(docx_bytes))
    # Safety net for overlapping campos the validator didn't see together: fill
    # substring campos longest-valor_ejemplo-first (permuted within their own
    # slots so order relative to other modos is unchanged), so a short value
    # ("2026") can't be consumed by a longer sibling's replacement first.
    orden = list(campos)
    sub_idx = [i for i, c in enumerate(orden) if c.get("modo") == "substring"]
    subs = sorted((orden[i] for i in sub_idx), key=lambda c: -len(c.get("valor_ejemplo") or ""))
    for i, c in zip(sub_idx, subs, strict=True):
        orden[i] = c
    for campo in orden:
        clave = campo.get("campo")
        if clave is None or clave not in valores:
            continue
        valor = valores[clave]
        node = resolver_direccion(doc, campo.get("direccion", ""))
        if node is None:
            continue
        modo = campo.get("modo", "cell")
        if modo == "cell":
            paragraphs = _node_paragraphs(node)
            _set_paragraph_text(paragraphs[0], valor)
            for p in paragraphs[1:]:
                _blank_paragraph(p)
        elif modo == "substring":
            for p in _node_paragraphs(node):
                if _replace_text_preserving_runs(p, campo.get("valor_ejemplo", ""), valor):
                    break
        elif modo == "checkbox":
            if valor != "alternativa":
                continue  # "actual" (or anything else) → marker already in place
            marcador = campo.get("valor_ejemplo", "X")
            for p in _node_paragraphs(node):
                if _replace_text_preserving_runs(p, marcador, ""):
                    break
            alt = resolver_direccion(doc, campo.get("direccion_alternativa", ""))
            if alt is not None:
                alt_paragraphs = _node_paragraphs(alt)
                alt_marcador = campo.get("valor_alternativo") or marcador
                runs = alt_paragraphs[0].runs
                if runs:
                    runs[-1].text = f"{runs[-1].text} {alt_marcador}".strip()
                else:
                    alt_paragraphs[0].add_run(alt_marcador)
        elif modo == "justificacion":
            lines = [ln for chunk in valor.split("\n\n") for ln in chunk.split("\n") if ln.strip()] or [""]
            paragraphs = _node_paragraphs(node)
            _set_paragraph_text(paragraphs[0], lines[0])
            for p in paragraphs[1:]:
                _blank_paragraph(p)
            if isinstance(node, _Cell):
                for line in lines[1:]:
                    node.add_paragraph(line)
    out = BytesIO()
    doc.save(out)
    return out.getvalue()
