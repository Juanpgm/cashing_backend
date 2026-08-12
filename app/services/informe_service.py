"""Informe service — generate DOCX reports and ZIP evidence folder for a CuentaCobro.

Produces three artifacts on demand for the contractor's billing package:
1. Informe de actividades del contratista (DOCX)
2. Informe de supervisión (DOCX)
3. Carpeta de evidencias (ZIP) — folder structure with placeholders per obligation.
"""

from __future__ import annotations

import calendar
import io
import uuid
import zipfile
from dataclasses import dataclass
from datetime import date
from decimal import Decimal
from typing import Literal

import structlog
from docx import Document
from docx.document import Document as DocxDocument  # real class for type hints; `Document` above is a factory
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.adapters.llm import get_llm
from app.adapters.storage import get_storage as _get_storage
from app.agent.prompts.supervision_tercera_persona import (
    TERCERA_PERSONA_SYSTEM_PROMPT,
    build_tercera_persona_prompt,
    parse_tercera_persona,
)
from app.core import text_match
from app.core.config import settings
from app.core.exceptions import (
    ACTIVIDADES_MISSING,
    PACKAGE_PENDIENTE,
    SECRET_DETECTED_IN_PACKAGE,
    ForbiddenError,
    NotFoundError,
    ValidationError,
)
from app.core.numero_letras import ordinal_letras, valor_en_letras
from app.models.actividad import Actividad
from app.models.contrato import Contrato
from app.models.cuenta_cobro import CuentaCobro, PosicionCuota
from app.models.documento_fuente import DocumentoFuente
from app.models.evidencia import Evidencia
from app.models.evidencia_obligacion import EstadoEnlace, EvidenciaObligacion
from app.models.obligacion import Obligacion
from app.models.plantilla_organismo import PlantillaOrganismo
from app.models.usuario import Usuario
from app.schemas.agent import LLMMessage
from app.services import document_service, secret_scan_service
from app.services.informe_constants import SENTINEL_SIN_EVIDENCIAS, TEXTO_SIN_LABORES

logger = structlog.get_logger("service.informe")

_MESES = [
    "Enero",
    "Febrero",
    "Marzo",
    "Abril",
    "Mayo",
    "Junio",
    "Julio",
    "Agosto",
    "Septiembre",
    "Octubre",
    "Noviembre",
    "Diciembre",
]

# Every generated informe carries this label — a CONSTANT, NEVER an LLM decision
# (billing-resilience-templates, slice #6, design D6, tasks 6.7-6.8).
_BORRADOR_HEADER = "BORRADOR — sujeto a revisión"

# Machine-readable counterpart of `_BORRADOR_HEADER` (billing-resilience-templates,
# slice #7, task 7.5d — carry-over from slice #6 verify-report WARNING 1): design
# D6's `es_borrador: bool = True` existed ONLY as a DOCX text header, so no API/tool
# caller could detect draft status without parsing the DOCX. Every informe this
# service generates is currently ALWAYS a draft (no review/approval workflow exists
# yet), so this is a constant — never derived, never an LLM decision — exposed for
# callers that need a programmatic flag: the direct-download endpoints surface it as
# an `X-Es-Borrador` response header, and `radicacion_prep_service.preparar_radicacion`
# surfaces it on its result.
ES_BORRADOR = True

# Progressive narrative bounds (slice #6, design D6) — mirrors the char-budget
# guard already proven in `requisito_inference_service._MAX_TEXT_CHARS`.
_MAX_TEXT_CHARS = 14_000
_K_RESUMENES_RECIENTES = 3

_NARRATIVA_SYSTEM_PROMPT = (
    "Eres un asistente que redacta un párrafo breve (máximo 3-4 frases) resumiendo, "
    "en tercera persona y en español neutro, el avance reportado en cuotas anteriores "
    "de un contrato de prestación de servicios. Usa SOLO la información entregada, sin "
    "inventar datos nuevos."
)


# ── Internal helpers ────────────────────────────────────────────────────────


async def _load_context(
    db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID
) -> tuple[CuentaCobro, Contrato, Usuario, list[Obligacion], list[Actividad]]:
    """Load all data needed to render the informes, validating ownership."""
    result = await db.execute(
        select(CuentaCobro)
        .options(
            selectinload(CuentaCobro.contrato).selectinload(Contrato.obligaciones),
            selectinload(CuentaCobro.actividades).selectinload(Actividad.evidencias),
        )
        .where(CuentaCobro.id == cuenta_id, CuentaCobro.deleted_at.is_(None))
    )
    cuenta = result.scalar_one_or_none()
    if cuenta is None:
        raise NotFoundError("CuentaCobro", str(cuenta_id))
    if cuenta.contrato.usuario_id != usuario_id:
        raise ForbiddenError()

    user_result = await db.execute(select(Usuario).where(Usuario.id == usuario_id))
    usuario = user_result.scalar_one_or_none()
    if usuario is None:
        raise NotFoundError("Usuario", str(usuario_id))

    obligaciones = sorted(cuenta.contrato.obligaciones, key=lambda o: o.orden)
    actividades = sorted(cuenta.actividades, key=lambda a: a.created_at)
    return cuenta, cuenta.contrato, usuario, obligaciones, actividades


def _periodo_str(cuenta: CuentaCobro) -> str:
    return f"{_MESES[cuenta.mes - 1]} de {cuenta.anio}"


def _formato_fecha(d: date | None) -> str:
    return d.strftime("%d/%m/%Y") if d else "—"


def _formato_valor(v: float) -> str:
    return f"$ {float(v):,.2f}"


def _dec(v: float | Decimal | None) -> Decimal:
    """Lossless float/Numeric → Decimal (money must never be float math)."""
    return Decimal(str(v)) if v is not None else Decimal(0)


def _formato_valor_pesos(v: float | Decimal) -> str:
    """Colombian whole-peso format used by entity templates: "$ 4.800.000"."""
    return f"$ {int(_dec(v)):,}".replace(",", ".")


def _add_kv_table(doc: DocxDocument, rows: list[tuple[str, str]]) -> None:
    """Add a 2-column key/value table for header info."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        # Bold the key column
        for paragraph in table.cell(i, 0).paragraphs:
            for run in paragraph.runs:
                run.bold = True


def _prefijo_borrador(origen: str | None) -> str:
    """Visible `[BORRADOR] ` prefix for `seed`-origin justificación text — Gemini
    text was pending/failed but there was evidence, so the seed fallback text
    must never be mistaken for a final, LLM-authored justificación (B6)."""
    return "[BORRADOR] " if origen == "seed" else ""


def _es_sin_labores(act: Actividad) -> bool:
    """True when this actividad's justificación is the tactful "sin labores" text
    (new `origen="sin_labores"`, or the historical sentinel persisted before this
    migration — backward-compat, B4)."""
    return act.justificacion_origen == "sin_labores" or (act.justificacion or "").strip() == SENTINEL_SIN_EVIDENCIAS


def _add_actividades_table(
    doc: DocxDocument,
    actividades: list[Actividad],
    obligaciones_by_id: dict[uuid.UUID, Obligacion],
    overrides: dict[uuid.UUID, tuple[str, str]] | None = None,
) -> None:
    """Render activities as a bordered table: # | Obligación | Actividad | Justificación.

    ``overrides`` optionally maps an activity id to a ``(actividad, justificacion)``
    pair to render instead of ``act.descripcion``/``act.justificacion`` (used by the
    supervisión report to render third-person text). The obligación column is never
    affected by overrides. Defaults to None, preserving the exact original behavior
    for the actividades report.
    """
    headers = ["#", "Obligación", "Actividad realizada", "Justificación"]
    table = doc.add_table(rows=1 + len(actividades), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for idx, act in enumerate(actividades, start=1):
        row = table.rows[idx]
        ob = obligaciones_by_id.get(act.obligacion_id) if act.obligacion_id else None
        ob_text = ob.descripcion if ob else "—"
        override = overrides.get(act.id) if overrides else None
        actividad_text = override[0] if override else act.descripcion
        justificacion_text = override[1] if override else (act.justificacion or "—")
        if not override and justificacion_text != "—" and not _es_sin_labores(act):
            justificacion_text = _prefijo_borrador(act.justificacion_origen) + justificacion_text
        row.cells[0].text = str(idx)
        row.cells[1].text = ob_text
        row.cells[2].text = actividad_text
        row.cells[3].text = justificacion_text


# ── Adaptive generation (billing-resilience-templates, slice #6) ────────────


def _mapear_columna(nombre_columna: str) -> str:
    """Classify an ORGANISM'S OWN ingested column name into a semantic render
    slot. Unknown/unrecognized columns render blank rather than guessing —
    zero risk of putting the wrong content in a column we can't confidently
    identify."""
    norm = text_match.normalize(nombre_columna)
    if "oblig" in norm:
        return "obligacion"
    if "evidenc" in norm:
        return "evidencia"
    if "justific" in norm:
        return "justificacion"
    if "avance" in norm or "activ" in norm or "cumplim" in norm:
        return "avance"
    if norm in {"#", "no", "n", "num", "numero"}:
        return "numero"
    return "otro"


def _add_actividades_table_adaptativa(
    doc: DocxDocument,
    columnas: list[str],
    actividades: list[Actividad],
    obligaciones_by_id: dict[uuid.UUID, Obligacion],
    overrides: dict[uuid.UUID, tuple[str, str]] | None = None,
) -> None:
    """Render the activities table using the organism's OWN ingested column
    layout instead of the fixed 4-column default (per-organism layout
    selection, slice #6, tasks 6.1-6.4). Column CONTENT is classified by
    `_mapear_columna`; a column this can't confidently classify renders blank
    rather than risk misplacing content in the wrong slot.
    """
    table = doc.add_table(rows=1 + len(actividades), cols=len(columnas))
    table.style = "Table Grid"
    for i, h in enumerate(columnas):
        cell = table.cell(0, i)
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    slots = [_mapear_columna(c) for c in columnas]
    for idx, act in enumerate(actividades, start=1):
        row = table.rows[idx]
        ob = obligaciones_by_id.get(act.obligacion_id) if act.obligacion_id else None
        override = overrides.get(act.id) if overrides else None
        actividad_text = override[0] if override else act.descripcion
        justificacion_text = override[1] if override else (act.justificacion or "—")
        if not override and justificacion_text != "—" and not _es_sin_labores(act):
            justificacion_text = _prefijo_borrador(act.justificacion_origen) + justificacion_text
        for col_idx, slot in enumerate(slots):
            if slot == "numero":
                texto = str(idx)
            elif slot == "obligacion":
                texto = ob.descripcion if ob else "—"
            elif slot == "avance":
                texto = actividad_text
            elif slot == "justificacion":
                texto = justificacion_text
            elif slot == "evidencia":
                count = len(act.evidencias) if act.evidencias else 0
                texto = f"{count} archivo(s)" if count else "—"
            else:
                texto = "—"
            row.cells[col_idx].text = texto


def _add_anexo_refs(doc: DocxDocument, anexo_refs: list[str]) -> None:
    """Append the ORGANISM'S OWN literal anexo-reference strings verbatim
    (spec: "Anexo reference preserved verbatim") — never paraphrased."""
    if not anexo_refs:
        return
    doc.add_paragraph()
    doc.add_heading("Referencias a anexos", level=2)
    for ref in anexo_refs:
        doc.add_paragraph(ref)


def _add_borrador_header(doc: DocxDocument) -> None:
    """Every generated informe carries this label as its very first paragraph
    (design D6, tasks 6.7-6.8) — a CONSTANT, never an LLM decision."""
    p = doc.add_paragraph(_BORRADOR_HEADER)
    p.alignment = 1  # center
    for run in p.runs:
        run.bold = True


async def _resolver_layout_organismo(
    db: AsyncSession, contrato: Contrato, tipo_documento: str
) -> PlantillaOrganismo | None:
    """Look up the ingested per-organism template structure for THIS document
    type (informe_actividades / informe_supervision) to select the informe's
    column layout (slice #6). Independent of `_resolver_estructura_organismo`
    (slice #5), which always resolves the "informe_actividades" structure for
    the evidence ZIP's folder naming regardless of which DOCX is generated.
    Fails open (returns None) on any lookup error — a layout-selection failure
    must never block informe generation; the caller falls back to the default
    fixed layout exactly as when no template was ever ingested.
    """
    from app.services import requisito_inference_service

    try:
        return await requisito_inference_service.obtener_plantilla_organismo(
            db, contrato.usuario_id, contrato.id, tipo_documento=tipo_documento
        )
    except Exception as exc:  # a lookup failure must not block generation
        await logger.awarning("informe_layout_organismo_lookup_failed", contrato_id=str(contrato.id), error=str(exc))
        return None


def _filtrar_actividades_una_vez(
    actividades: list[Actividad],
    cuenta: CuentaCobro,
    obligaciones_by_id: dict[uuid.UUID, Obligacion],
) -> list[Actividad]:
    """Omit content for one-time (`Obligacion.una_vez`) obligations in every
    cuota AFTER the contrato's first one (spec: "One-time obligations blank
    after cuota 1", tasks 6.9-6.10). The cuota's own `posicion` (slice #3,
    persisted — never re-derived here) decides, NOT `numero_cuota`, since a
    legacy/backfilled row's `numero_cuota` may be absent."""
    if cuenta.posicion == PosicionCuota.PRIMERA:
        return actividades
    return [
        act
        for act in actividades
        if not (
            act.obligacion_id is not None
            and (ob := obligaciones_by_id.get(act.obligacion_id)) is not None
            and ob.una_vez
        )
    ]


async def _resumen_cuota(db: AsyncSession, cuenta: CuentaCobro) -> str:
    """Best-effort text summary of a prior cuota's reported activities, used ONLY
    as progressive-narrative context (slice #6) — never persisted, never shown
    to the caller directly (it's synthesized further by `_generar_narrativa_progresiva`)."""
    result = await db.execute(select(Actividad).where(Actividad.cuenta_cobro_id == cuenta.id))
    partes: list[str] = []
    for act in result.scalars().all():
        texto = (act.descripcion or "").strip()
        if act.justificacion:
            texto = f"{texto} — {act.justificacion.strip()}" if texto else act.justificacion.strip()
        if texto:
            partes.append(texto)
    return "; ".join(partes)


async def _cargar_cuotas_previas(
    db: AsyncSession, contrato_id: uuid.UUID, cuenta_actual: CuentaCobro
) -> list[CuentaCobro]:
    """Every cuota of the SAME contrato with `numero_cuota < cuenta_actual.numero_cuota`,
    chronologically ordered. Returns an empty list for the first cuota (or any cuota
    missing `numero_cuota`) — no prior context can exist."""
    if cuenta_actual.numero_cuota is None or cuenta_actual.numero_cuota <= 1:
        return []
    result = await db.execute(
        select(CuentaCobro)
        .where(
            CuentaCobro.contrato_id == contrato_id,
            CuentaCobro.deleted_at.is_(None),
            CuentaCobro.numero_cuota.is_not(None),
            CuentaCobro.numero_cuota < cuenta_actual.numero_cuota,
        )
        .order_by(CuentaCobro.numero_cuota.asc())
    )
    return list(result.scalars().all())


async def _generar_narrativa_progresiva(bloque_contexto: str) -> str:
    """LLM-synthesize a short narrative from the bounded prior-cuota context
    block. Fails OPEN exactly like `_convertir_actividades_tercera_persona`:
    any LLM error or empty response falls back to the raw `bloque_contexto`
    itself, so a narrative-synthesis failure never blocks informe generation.
    """
    try:
        llm = get_llm()
        resp = await llm.complete(
            [
                LLMMessage(role="system", content=_NARRATIVA_SYSTEM_PROMPT),
                LLMMessage(role="user", content=bloque_contexto),
            ],
            temperature=0.2,
            max_tokens=400,
        )
        texto = (resp.content or "").strip()
        return texto or bloque_contexto
    except Exception as exc:
        await logger.awarning("informe_narrativa_progresiva_llm_error", error=str(exc))
        return bloque_contexto


async def _construir_contexto_progresivo(
    db: AsyncSession, contrato_id: uuid.UUID, cuenta_actual: CuentaCobro
) -> str | None:
    """Build the progressive-narrative context for cuota N from cuotas 1..N-1 of
    the SAME contrato (billing-resilience-templates, slice #6, design D6).
    Returns `None` for the first cuota (no prior context exists). Bounded by
    `_MAX_TEXT_CHARS`; when the joined summaries exceed it, degrades to the
    `_K_RESUMENES_RECIENTES` most-recent summaries + a note of how many were
    omitted (design D6, mirrors `requisito_inference_service`'s char-budget
    guard). Fails open — any loading error returns `None` rather than raising,
    since a missing narrative must never block informe generation.
    """
    try:
        previas = await _cargar_cuotas_previas(db, contrato_id, cuenta_actual)
        if not previas:
            return None
        resumenes = [(pc.numero_cuota, _periodo_str(pc), await _resumen_cuota(db, pc)) for pc in previas]
    except Exception as exc:
        await logger.awarning("informe_contexto_progresivo_load_failed", error=str(exc))
        return None

    total_chars = sum(len(texto) for _n, _p, texto in resumenes)
    if total_chars > _MAX_TEXT_CHARS:
        recientes = resumenes[-_K_RESUMENES_RECIENTES:]
        omitidas = len(resumenes) - len(recientes)
        partes = [
            f"(se omiten {omitidas} cuota(s) anteriores por longitud; se muestran las {len(recientes)} más recientes)"
        ]
        partes += [f"Cuota {n} ({periodo}): {texto}" for n, periodo, texto in recientes]
    else:
        partes = [f"Cuota {n} ({periodo}): {texto}" for n, periodo, texto in resumenes]

    bloque = "\n".join(partes)
    return await _generar_narrativa_progresiva(bloque)


async def _convertir_actividades_tercera_persona(
    actividades: list[Actividad],
) -> dict[uuid.UUID, tuple[str, str]]:
    """Batch-rewrite each activity's descripcion/justificacion into third person.

    Makes a SINGLE LLM call over all non-empty texts (both columns, flattened across
    all activities) to keep the supervisión report generation cheap and fast. Fails
    OPEN on any error, empty batch, or unparseable/mismatched response: returns an
    empty dict so the caller falls back to the original first-person text and report
    generation never breaks because of this conversion step.
    """
    empty_markers = {"", "—"}

    # (actividad_id, column) -> flattened text index, built only for non-empty texts.
    entries: list[tuple[uuid.UUID, str]] = []
    textos: list[str] = []
    for act in actividades:
        descripcion = (act.descripcion or "").strip()
        if descripcion and descripcion not in empty_markers:
            entries.append((act.id, "descripcion"))
            textos.append(descripcion)
        justificacion = (act.justificacion or "").strip()
        if justificacion and justificacion not in empty_markers:
            entries.append((act.id, "justificacion"))
            textos.append(justificacion)

    if not textos:
        return {}

    try:
        llm = get_llm()
        resp = await llm.complete(
            [
                LLMMessage(role="system", content=TERCERA_PERSONA_SYSTEM_PROMPT),
                LLMMessage(role="user", content=build_tercera_persona_prompt(textos)),
            ],
            temperature=0.2,
            max_tokens=min(2048, 120 * len(textos)),
        )
        reescritos = parse_tercera_persona(resp.content, expected=len(textos))
        if reescritos is None:
            await logger.awarning("informe_supervision.tercera_persona_parse_failed", batch_size=len(textos))
            return {}
    except Exception as exc:
        await logger.awarning("informe_supervision.tercera_persona_llm_error", error=str(exc), batch_size=len(textos))
        return {}

    by_actividad: dict[uuid.UUID, dict[str, str]] = {}
    for (act_id, column), texto in zip(entries, reescritos, strict=True):
        by_actividad.setdefault(act_id, {})[column] = texto

    overrides: dict[uuid.UUID, tuple[str, str]] = {}
    for act in actividades:
        converted = by_actividad.get(act.id)
        if not converted:
            continue
        actividad_text = converted.get("descripcion", act.descripcion)
        justificacion_text = converted.get("justificacion", act.justificacion or "—")
        overrides[act.id] = (actividad_text, justificacion_text)
    return overrides


# ── Template cloning (docx clone, phase 3) ──────────────────────────────────


async def _valores_plantilla(db: AsyncSession, cuenta: CuentaCobro, contrato: Contrato) -> dict[str, str]:
    """Resolve the canonical campo vocabulary to concrete string values.

    Money is Decimal all the way and rendered in the entity's whole-peso format
    ("$ 4.800.000"). Keys whose source data is missing are omitted — `rellenar`
    skips campos without a value.
    """
    from sqlalchemy import and_, or_

    usuario = await db.get(Usuario, contrato.usuario_id)

    result = await db.execute(
        select(CuentaCobro)
        .where(
            CuentaCobro.contrato_id == contrato.id,
            CuentaCobro.deleted_at.is_(None),
            CuentaCobro.id != cuenta.id,
            or_(
                CuentaCobro.anio < cuenta.anio,
                and_(CuentaCobro.anio == cuenta.anio, CuentaCobro.mes < cuenta.mes),
            ),
        )
        .order_by(CuentaCobro.anio.asc(), CuentaCobro.mes.asc())
    )
    previas = list(result.scalars().all())
    acumulado = sum((_dec(p.valor) for p in previas), Decimal(0))
    total_con_adiciones = _dec(contrato.valor_total) + _dec(contrato.valor_adicion)
    valor_cuota = _dec(cuenta.valor)
    saldo = total_con_adiciones - acumulado - valor_cuota
    numero_cuota = cuenta.numero_cuota or len(previas) + 1
    fecha_cuota = cuenta.fecha_transaccion or date(
        cuenta.anio, cuenta.mes, calendar.monthrange(cuenta.anio, cuenta.mes)[1]
    )

    valores: dict[str, str | None] = {
        "contrato.numero_contrato": contrato.numero_contrato,
        "contrato.objeto": contrato.objeto,
        "contrato.entidad": contrato.entidad,
        "contrato.dependencia": contrato.dependencia,
        "contrato.valor_total": _formato_valor_pesos(contrato.valor_total),
        "contrato.valor_mensual": _formato_valor_pesos(contrato.valor_mensual),
        "contrato.fecha_inicio": _formato_fecha(contrato.fecha_inicio),
        "contrato.fecha_fin": _formato_fecha(contrato.fecha_fin),
        "contrato.supervisor_nombre": contrato.supervisor_nombre,
        "contrato.cargo_supervisor": contrato.cargo_supervisor,
        "contratista.nombre": usuario.nombre if usuario else None,
        "contratista.cedula": (usuario.cedula if usuario else None) or contrato.documento_proveedor,
        "cuota.valor": _formato_valor_pesos(valor_cuota),
        "cuota.numero": str(numero_cuota),
        "cuota.mes": _MESES[cuenta.mes - 1],
        "cuota.anio": str(cuenta.anio),
        "cuota.periodo": _periodo_str(cuenta),
        "cuota.fecha": _formato_fecha(fecha_cuota),
        "computed.acumulado": _formato_valor_pesos(acumulado),
        "computed.saldo": _formato_valor_pesos(saldo),
        "computed.valor_total_con_adiciones": _formato_valor_pesos(total_con_adiciones),
        "computed.valor_letras": valor_en_letras(valor_cuota),
        "computed.cuota_ordinal_letras": ordinal_letras(numero_cuota),
        # Documento Soporte consecutive (xlsx clone, phase 4; G4 migration 034):
        # user-supplied per cuota via CuentaCobro.consecutivo_ds — omitted (None)
        # when not set, so fill-time skips the campo (existing blank-field behavior).
        "cuota.consecutivo_ds": cuenta.consecutivo_ds,
    }
    return {k: v for k, v in valores.items() if v}


_AVISO_TIPO_INFORME_AMBIGUO = "No se pudo determinar la opción marcada del tipo de informe; marcar manualmente."


def _resolver_tipo_informe(cuenta: CuentaCobro, texto_marcado: str | None, texto_alternativo: str | None) -> str | None:
    """Checkbox resolution from the ACTUAL text of both option nodes in the
    template (the campo contract's valor_ejemplo is only the marker glyph "X",
    and the etiqueta is generic — neither identifies which option is which).

    Returns "actual" when the currently-marked option matches the cuenta
    (marked option mentions FINAL ⇔ the cuenta is final), "alternativa" when
    the marker must move, and None when the options are ambiguous (neither or
    both mention "final") — the campo must then be left untouched.
    """
    marcado_es_final = "final" in text_match.normalize(texto_marcado or "")
    alt_es_final = "final" in text_match.normalize(texto_alternativo or "")
    if marcado_es_final == alt_es_final:
        return None
    es_final = bool(cuenta.informe_final) or cuenta.posicion == PosicionCuota.FINAL
    return "actual" if marcado_es_final == es_final else "alternativa"


async def _tipo_informe_desde_plantilla(
    db: AsyncSession, plantilla: PlantillaOrganismo, campo: dict, cuenta: CuentaCobro
) -> str | None:
    """Resolve the tipo-informe checkbox for the formato-valores preview by
    downloading the original template bytes and reading both option texts
    (minimal correct option: the preview endpoint doesn't otherwise have the
    bytes, and resolving from campo metadata alone is impossible — see
    `_resolver_tipo_informe`). Any failure returns None so the caller emits
    the manual-marking aviso instead of guessing.
    """
    from app.services.docx_clone_service import texto_direccion

    try:
        if not plantilla.fuente_documento_id:
            return None
        doc_fuente = await db.get(DocumentoFuente, plantilla.fuente_documento_id)
        if doc_fuente is None:
            return None
        storage = _get_storage(settings.S3_BUCKET_DOCUMENTOS)
        original = await storage.download(doc_fuente.storage_key)
        doc = Document(io.BytesIO(original))
        return _resolver_tipo_informe(
            cuenta,
            texto_direccion(doc, campo.get("direccion", "")),
            texto_direccion(doc, campo.get("direccion_alternativa", "")),
        )
    except Exception as exc:
        await logger.awarning("formato_valores_tipo_informe_fallo", cuenta_id=str(cuenta.id), error=str(exc))
        return None


def _justificaciones_tercera_por_orden(
    actividades_visibles: list[Actividad],
    obligaciones_by_id: dict[uuid.UUID, Obligacion],
    overrides: dict[uuid.UUID, tuple[str, str]],
) -> dict[int, str]:
    """Third-person justification text per obligación orden (1-based), from the
    `_convertir_actividades_tercera_persona` overrides — for the cloned
    supervisión path.

    An actividad with no justificación AND whose obligación has no evidencia
    is exempted from the LLM tercera-persona conversion — the sentinel is a
    fixed deterministic string, not user prose, and the LLM could mangle its
    exact wording. It renders as `TEXTO_SIN_LABORES` directly instead (B4: a
    single tactful text serves both 1st and 3rd person).

    The `[BORRADOR] ` prefix (B6) is applied HERE, once, for `seed`-origin text
    — not in the generic `_add_actividades_table` consumer, so it doesn't
    interfere with the override's editorial third-person rewrite.
    """
    actividades_por_ob: dict[uuid.UUID | None, list[Actividad]] = {}
    for act in actividades_visibles:
        actividades_por_ob.setdefault(act.obligacion_id, []).append(act)

    por_orden: dict[int, str] = {}
    for n, ob in enumerate(sorted(obligaciones_by_id.values(), key=lambda o: o.orden), start=1):
        acts_ob = actividades_por_ob.get(ob.id, [])
        tiene_evidencia = any(act.evidencias for act in acts_ob)
        partes: list[str] = []
        for act in acts_ob:
            justificacion = (act.justificacion or "").strip()
            if _es_sin_labores(act) or (not justificacion and not tiene_evidencia):
                texto = TEXTO_SIN_LABORES
            else:
                ov = overrides.get(act.id)
                texto = (ov[1] if ov[1] and ov[1] != "—" else ov[0]) if ov else (act.justificacion or act.descripcion)
                texto = (texto or "").strip()
                if texto:
                    texto = _prefijo_borrador(act.justificacion_origen) + texto
            if texto:
                partes.append(texto)
        if partes:
            por_orden[n] = "\n\n".join(partes)
    return por_orden


async def _generar_docx_clonado(
    db: AsyncSession,
    plantilla: PlantillaOrganismo | None,
    cuenta: CuentaCobro,
    contrato: Contrato,
    actividades_visibles: list[Actividad],
    obligaciones_by_id: dict[uuid.UUID, Obligacion],
    justificaciones_override: dict[int, str] | None = None,
) -> bytes | None:
    """Clone-and-fill the organism's original DOCX template. Fails open.

    Guards: DOCX format, `clonable` flag, persisted campos, and a resolvable
    source documento. ANY failure (missing fuente, storage error, fill error)
    logs `informe_clonado_fallback` and returns None so the caller falls back
    to the built-from-scratch legacy layout.
    """
    from app.services.docx_clone_service import rellenar, texto_direccion

    try:
        if plantilla is None or plantilla.formato != "docx":
            return None
        estructura = plantilla.estructura_json or {}
        campos = estructura.get("campos") or []
        if not (estructura.get("clonable") and campos and plantilla.fuente_documento_id):
            return None
        doc_fuente = await db.get(DocumentoFuente, plantilla.fuente_documento_id)
        if doc_fuente is None:
            return None

        storage = _get_storage(settings.S3_BUCKET_DOCUMENTOS)
        original = await storage.download(doc_fuente.storage_key)

        valores = await _valores_plantilla(db, cuenta, contrato)
        doc_original = Document(io.BytesIO(original))
        for campo in campos:
            if campo.get("campo") == "computed.tipo_informe" and campo.get("modo") == "checkbox":
                tipo = _resolver_tipo_informe(
                    cuenta,
                    texto_direccion(doc_original, campo.get("direccion", "")),
                    texto_direccion(doc_original, campo.get("direccion_alternativa", "")),
                )
                if tipo is not None:
                    valores["computed.tipo_informe"] = tipo
                else:
                    # Ambiguous options — leave the template's marker untouched.
                    await logger.awarning(
                        "informe_clonado_tipo_informe_ambiguo",
                        cuenta_id=str(cuenta.id),
                        aviso=_AVISO_TIPO_INFORME_AMBIGUO,
                    )

        actividades_por_ob: dict[uuid.UUID | None, list[Actividad]] = {}
        for act in actividades_visibles:
            actividades_por_ob.setdefault(act.obligacion_id, []).append(act)
        for n, ob in enumerate(sorted(obligaciones_by_id.values(), key=lambda o: o.orden), start=1):
            if justificaciones_override is not None and n in justificaciones_override:
                texto = justificaciones_override[n]
            else:
                acts_ob = actividades_por_ob.get(ob.id, [])
                # An obligación with no evidencia anywhere and no written
                # justificación would otherwise echo `act.descripcion` here —
                # meaningless in a real informe (it just repeats the
                # obligación's own seeded activity text). Render the tactful
                # "sin labores" text instead; keep the descripcion fallback
                # ONLY when the obligación has real evidencia (unchanged prior
                # behavior). `seed`-origin text gets a visible [BORRADOR] prefix
                # (B6) — Gemini text was pending/failed, never pass it off as final.
                tiene_evidencia = any(act.evidencias for act in acts_ob)
                partes: list[str] = []
                for act in acts_ob:
                    justificacion = (act.justificacion or "").strip()
                    if _es_sin_labores(act) or (not justificacion and not tiene_evidencia):
                        partes.append(TEXTO_SIN_LABORES)
                    elif justificacion:
                        partes.append(_prefijo_borrador(act.justificacion_origen) + justificacion)
                    else:
                        descripcion = (act.descripcion or "").strip()
                        if descripcion:
                            partes.append(descripcion)
                texto = "\n\n".join(partes)
            if texto:
                valores[f"justificacion.{n}"] = texto

        # Blank example data for every campo without a resolved value: `rellenar`
        # skips missing keys, which would silently keep the ORIGINAL example's
        # real data (another submission's names, cédulas, narratives) in the
        # generated official document. Checkbox campos are excluded — blanking a
        # marker makes no sense; ambiguity is already logged above.
        justificaciones_blanqueadas = 0
        campos_blanqueados: list[str] = []
        for campo in campos:
            clave = campo.get("campo")
            if not clave or clave in valores or campo.get("modo") == "checkbox":
                continue
            valores[clave] = ""
            if campo.get("modo") == "justificacion":
                justificaciones_blanqueadas += 1
            else:
                campos_blanqueados.append(campo.get("etiqueta") or clave)
        if justificaciones_blanqueadas:
            await logger.awarning(
                "informe_clonado_justificaciones_blanqueadas",
                cuenta_id=str(cuenta.id),
                n=justificaciones_blanqueadas,
            )
        if campos_blanqueados:
            await logger.awarning(
                "informe_clonado_campos_blanqueados",
                cuenta_id=str(cuenta.id),
                campos=campos_blanqueados,
            )

        return rellenar(original, campos, valores)
    except Exception:
        await logger.awarning(
            "informe_clonado_fallback",
            cuenta_id=str(cuenta.id),
            plantilla_id=str(plantilla.id) if plantilla is not None else None,
            exc_info=True,
        )
        return None


# ── Public API ──────────────────────────────────────────────────────────────


async def generar_informe_actividades_docx(
    db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID
) -> tuple[bytes, str]:
    """Generate the contractor's activities report as DOCX. Returns (bytes, filename).

    Adaptive generation (billing-resilience-templates, slice #6): organismo,
    posicion, and prior-cuota context are resolved INTERNALLY from `contrato`/
    `cuenta` (already loaded by `_load_context`) rather than accepted as extra
    params — every existing call site keeps calling this with exactly 3
    positional args.

    Clone path (phase 3): when the organism's plantilla is clonable, the
    original entity DOCX is cloned and filled instead of built from scratch.
    The BORRADOR header is deliberately NOT injected into cloned bytes (it
    would corrupt the entity's format) — draft status stays on the
    `X-Es-Borrador` response header.
    """
    cuenta, contrato, usuario, obligaciones, actividades = await _load_context(db, usuario_id, cuenta_id)
    if not actividades:
        raise ValidationError(
            "No hay actividades registradas. Genera o ingresa actividades antes de descargar el informe.",
            code=ACTIVIDADES_MISSING,
        )

    obligaciones_by_id = {ob.id: ob for ob in obligaciones}
    actividades_visibles = _filtrar_actividades_una_vez(actividades, cuenta, obligaciones_by_id)
    layout = await _resolver_layout_organismo(db, contrato, "informe_actividades")
    filename = f"informe-actividades-{contrato.numero_contrato}-{cuenta.anio}-{cuenta.mes:02d}.docx"

    clonado = await _generar_docx_clonado(db, layout, cuenta, contrato, actividades_visibles, obligaciones_by_id)
    if clonado is not None:
        await logger.ainfo(
            "informe_actividades_clonado", cuenta_id=str(cuenta_id), usuario_id=str(usuario_id), size=len(clonado)
        )
        return clonado, filename

    contexto_progresivo = await _construir_contexto_progresivo(db, contrato.id, cuenta)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_borrador_header(doc)

    title = doc.add_heading("Informe de actividades del contratista", level=1)
    title.alignment = 1  # center
    doc.add_paragraph(f"Período reportado: {_periodo_str(cuenta)}").alignment = 1
    doc.add_paragraph()

    _add_kv_table(
        doc,
        [
            ("Entidad", contrato.entidad or "—"),
            ("Dependencia", contrato.dependencia or "—"),
            ("N° Contrato", contrato.numero_contrato),
            ("Contratista", usuario.nombre),
            ("C.C.", usuario.cedula or "—"),
            ("Supervisor", contrato.supervisor_nombre or "—"),
            ("Objeto", contrato.objeto),
            ("Período", _periodo_str(cuenta)),
            ("Valor a cobrar", _formato_valor(cuenta.valor)),
        ],
    )

    if contexto_progresivo:
        doc.add_paragraph()
        doc.add_heading("Contexto de cuotas anteriores", level=2)
        doc.add_paragraph(contexto_progresivo)

    doc.add_paragraph()
    doc.add_heading("Actividades realizadas", level=2)
    columnas = layout.estructura_json.get("columnas") if layout else None
    if columnas:
        _add_actividades_table_adaptativa(doc, columnas, actividades_visibles, obligaciones_by_id)
        _add_anexo_refs(doc, layout.estructura_json.get("anexo_refs") or [])  # type: ignore[union-attr]
    else:
        _add_actividades_table(doc, actividades_visibles, obligaciones_by_id)

    doc.add_paragraph()
    doc.add_paragraph(
        "Certifico que las actividades aquí relacionadas fueron ejecutadas en cumplimiento de las "
        "obligaciones contractuales pactadas durante el período reportado."
    )
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_______________________________")
    doc.add_paragraph(usuario.nombre)
    doc.add_paragraph(f"C.C. {usuario.cedula or '—'}")
    doc.add_paragraph("Contratista")

    buf = io.BytesIO()
    doc.save(buf)

    await logger.ainfo(
        "informe_actividades_generado",
        cuenta_id=str(cuenta_id),
        usuario_id=str(usuario_id),
        size=len(buf.getvalue()),
    )
    return buf.getvalue(), filename


async def generar_informe_supervision_docx(
    db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID
) -> tuple[bytes, str]:
    """Generate the supervisor's report as DOCX. Returns (bytes, filename).

    Adaptive generation (billing-resilience-templates, slice #6): see
    `generar_informe_actividades_docx`'s docstring — organismo/posicion/prior
    context are resolved internally, the public signature is unchanged.

    Clone path (phase 3): same as the actividades report, with justificaciones
    pre-converted to third person. The BORRADOR header is NOT injected into
    cloned bytes (it would corrupt the entity's format) — draft status stays on
    the `X-Es-Borrador` response header.
    """
    cuenta, contrato, usuario, obligaciones, actividades = await _load_context(db, usuario_id, cuenta_id)
    if not actividades:
        raise ValidationError(
            "No hay actividades registradas. Genera o ingresa actividades antes de descargar el informe.",
            code=ACTIVIDADES_MISSING,
        )

    obligaciones_by_id = {ob.id: ob for ob in obligaciones}
    actividades_visibles = _filtrar_actividades_una_vez(actividades, cuenta, obligaciones_by_id)
    overrides = await _convertir_actividades_tercera_persona(actividades_visibles)
    layout = await _resolver_layout_organismo(db, contrato, "informe_supervision")
    filename = f"informe-supervision-{contrato.numero_contrato}-{cuenta.anio}-{cuenta.mes:02d}.docx"

    clonado = await _generar_docx_clonado(
        db,
        layout,
        cuenta,
        contrato,
        actividades_visibles,
        obligaciones_by_id,
        justificaciones_override=_justificaciones_tercera_por_orden(actividades_visibles, obligaciones_by_id, overrides)
        or None,
    )
    if clonado is not None:
        await logger.ainfo(
            "informe_supervision_clonado", cuenta_id=str(cuenta_id), usuario_id=str(usuario_id), size=len(clonado)
        )
        return clonado, filename

    contexto_progresivo = await _construir_contexto_progresivo(db, contrato.id, cuenta)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_borrador_header(doc)

    title = doc.add_heading("Informe de supervisión", level=1)
    title.alignment = 1
    doc.add_paragraph(f"Período supervisado: {_periodo_str(cuenta)}").alignment = 1
    doc.add_paragraph()

    _add_kv_table(
        doc,
        [
            ("Entidad", contrato.entidad or "—"),
            ("Dependencia", contrato.dependencia or "—"),
            ("N° Contrato", contrato.numero_contrato),
            ("Contratista", usuario.nombre),
            ("C.C.", usuario.cedula or "—"),
            ("Supervisor", contrato.supervisor_nombre or "—"),
            ("Objeto", contrato.objeto),
            ("Período", _periodo_str(cuenta)),
            ("Valor autorizado", _formato_valor(cuenta.valor)),
        ],
    )

    if contexto_progresivo:
        doc.add_paragraph()
        doc.add_heading("Contexto de cuotas anteriores", level=2)
        doc.add_paragraph(contexto_progresivo)

    doc.add_paragraph()
    doc.add_heading("Verificación del cumplimiento", level=2)
    doc.add_paragraph(
        "El supervisor del contrato certifica que ha verificado el cumplimiento de las obligaciones "
        "contractuales por parte del contratista durante el período reportado, con base en las "
        "actividades y evidencias relacionadas a continuación."
    )

    doc.add_heading("Actividades verificadas", level=2)
    columnas = layout.estructura_json.get("columnas") if layout else None
    if columnas:
        _add_actividades_table_adaptativa(doc, columnas, actividades_visibles, obligaciones_by_id, overrides=overrides)
        _add_anexo_refs(doc, layout.estructura_json.get("anexo_refs") or [])  # type: ignore[union-attr]
    else:
        _add_actividades_table(doc, actividades_visibles, obligaciones_by_id, overrides=overrides)

    doc.add_paragraph()
    doc.add_heading("Concepto del supervisor", level=2)
    doc.add_paragraph(
        "Las actividades ejecutadas se encuentran a satisfacción y corresponden al objeto y "
        "obligaciones del contrato. Se recomienda autorizar el pago correspondiente al período."
    )

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_______________________________")
    doc.add_paragraph(contrato.supervisor_nombre or "—")
    doc.add_paragraph("Supervisor del contrato")

    buf = io.BytesIO()
    doc.save(buf)

    await logger.ainfo(
        "informe_supervision_generado",
        cuenta_id=str(cuenta_id),
        usuario_id=str(usuario_id),
        size=len(buf.getvalue()),
    )
    return buf.getvalue(), filename


async def generar_cuenta_cobro_docx(db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID) -> tuple[bytes, str]:
    """Generate the entity's cuenta-de-cobro letter as DOCX. Returns (bytes, filename).

    Clone-only: unlike the informes there is NO built-from-scratch fallback —
    this letter family only exists as the organism's own ingested template
    (tipo_documento="cuenta_cobro"). No BORRADOR header is injected (it would
    corrupt the entity format); draft status stays on the `X-Es-Borrador`
    response header.
    """
    cuenta, contrato, _usuario, obligaciones, actividades = await _load_context(db, usuario_id, cuenta_id)
    plantilla = await _resolver_layout_organismo(db, contrato, "cuenta_cobro")
    if plantilla is None or not (plantilla.estructura_json or {}).get("clonable"):
        raise ValidationError("No hay plantilla de cuenta de cobro ingerida para este organismo.")

    obligaciones_by_id = {ob.id: ob for ob in obligaciones}
    actividades_visibles = _filtrar_actividades_una_vez(actividades, cuenta, obligaciones_by_id)
    contenido = await _generar_docx_clonado(db, plantilla, cuenta, contrato, actividades_visibles, obligaciones_by_id)
    if contenido is None:
        raise ValidationError("No se pudo generar la cuenta de cobro desde la plantilla del organismo.")

    filename = f"cuenta-cobro-{contrato.numero_contrato}-{cuenta.anio}-{cuenta.mes:02d}.docx"
    await logger.ainfo(
        "cuenta_cobro_docx_generado", cuenta_id=str(cuenta_id), usuario_id=str(usuario_id), size=len(contenido)
    )
    return contenido, filename


async def generar_documento_soporte_xlsx(
    db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID
) -> tuple[bytes, str]:
    """Generate the entity's Documento Soporte spreadsheet as XLSX. Returns (bytes, filename).

    Clone-only, like `generar_cuenta_cobro_docx`: this document only exists as
    the organism's own ingested xlsx template (tipo_documento=
    "documento_soporte"). It is OPTIONAL everywhere — nothing in the billing
    flow requires it. Values come from `_valores_plantilla` (no justificaciones
    for xlsx); draft status stays on the `X-Es-Borrador` response header.
    """
    from app.services.xlsx_clone_service import rellenar_xlsx

    cuenta, contrato, _usuario, _obligaciones, _actividades = await _load_context(db, usuario_id, cuenta_id)
    plantilla = await _resolver_layout_organismo(db, contrato, "documento_soporte")
    estructura = (plantilla.estructura_json or {}) if plantilla is not None else {}
    campos = estructura.get("campos") or []
    if (
        plantilla is None
        or plantilla.formato != "xlsx"
        or not (estructura.get("clonable") and campos and plantilla.fuente_documento_id)
    ):
        raise ValidationError("No hay plantilla de documento soporte ingerida para este organismo.")

    try:
        doc_fuente = await db.get(DocumentoFuente, plantilla.fuente_documento_id)
        if doc_fuente is None:
            raise NotFoundError("Documento", str(plantilla.fuente_documento_id))
        storage = _get_storage(settings.S3_BUCKET_DOCUMENTOS)
        original = await storage.download(doc_fuente.storage_key)
        valores = await _valores_plantilla(db, cuenta, contrato)
        contenido = rellenar_xlsx(original, campos, valores)
    except Exception as exc:
        await logger.awarning(
            "documento_soporte_xlsx_fallo",
            cuenta_id=str(cuenta_id),
            plantilla_id=str(plantilla.id),
            exc_info=True,
        )
        raise ValidationError("No se pudo generar el documento soporte desde la plantilla del organismo.") from exc

    filename = f"documento-soporte-{contrato.numero_contrato}-{cuenta.anio}-{cuenta.mes:02d}.xlsx"
    await logger.ainfo(
        "documento_soporte_xlsx_generado", cuenta_id=str(cuenta_id), usuario_id=str(usuario_id), size=len(contenido)
    )
    return contenido, filename


async def obtener_formato_valores(
    db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID
) -> tuple[dict[str, str], list[str]]:
    """Resolved campo values + computed financials for the UI preview.

    Returns ``(valores, avisos)``. When a clonable cuenta-de-cobro plantilla
    exists, its checkbox campo also resolves `computed.tipo_informe`.
    """
    cuenta, contrato, _usuario, _obligaciones, _actividades = await _load_context(db, usuario_id, cuenta_id)
    valores = await _valores_plantilla(db, cuenta, contrato)
    avisos: list[str] = []

    plantilla_ds = await _resolver_layout_organismo(db, contrato, "documento_soporte")
    if plantilla_ds is not None and cuenta.consecutivo_ds is None:
        avisos.append("El Documento Soporte requiere el consecutivo DS; queda en blanco hasta que lo cargues.")

    plantilla = await _resolver_layout_organismo(db, contrato, "cuenta_cobro")
    if plantilla is not None:
        estructura = plantilla.estructura_json or {}
        campos = estructura.get("campos") or []
        for campo in campos:
            if campo.get("campo") == "computed.tipo_informe" and campo.get("modo") == "checkbox":
                tipo = await _tipo_informe_desde_plantilla(db, plantilla, campo, cuenta)
                if tipo is not None:
                    valores["computed.tipo_informe"] = tipo
                else:
                    avisos.append(_AVISO_TIPO_INFORME_AMBIGUO)
        if estructura.get("clonable"):
            # Mirror `_generar_docx_clonado`'s blanking: warn the user which
            # detected campos have no contrato/cuota data and will render blank.
            # Justificaciones (filled from actividades at generation time) and
            # checkboxes are not data-backed campos, so they are excluded.
            vistos: set[str] = set()
            sin_dato: list[str] = []
            for campo in campos:
                clave = campo.get("campo")
                if (
                    not clave
                    or clave in valores
                    or clave in vistos
                    or campo.get("modo") in ("checkbox", "justificacion")
                ):
                    continue
                vistos.add(clave)
                sin_dato.append(campo.get("etiqueta") or clave)
            if sin_dato:
                avisos.append("Campos sin dato del contrato/cuota (quedarán en blanco): " + ", ".join(sin_dato) + ".")
    else:
        avisos.append(
            "No hay plantilla de cuenta de cobro ingerida para este organismo; "
            "se muestran únicamente los valores calculados."
        )
    return valores, avisos


def _safe_dirname(text: str, max_len: int = 80) -> str:
    """Sanitize text to a filesystem-safe folder name."""
    bad = '<>:"/\\|?*\n\r\t'
    cleaned = "".join(c if c not in bad else "-" for c in text).strip()
    cleaned = " ".join(cleaned.split())  # collapse whitespace
    if len(cleaned) > max_len:
        cleaned = cleaned[:max_len].rstrip()
    return cleaned or "obligacion"


async def _resolver_estructura_organismo(db: AsyncSession, contrato: Contrato) -> PlantillaOrganismo | None:
    """Look up a per-organism ingested template structure to number/name folders.

    Real lookup (billing-resilience-templates, slice #5, task 5.15 — completes
    the slice #2 stub): delegates to `requisito_inference_service.
    obtener_plantilla_organismo_por_contrato` (normalized `Contrato.entidad`
    match), passing the `contrato` this function already received — already
    loaded+ownership-validated by `_load_context` in `generar_zip_evidencias`,
    so no redundant ownership round-trip (slice #7, task 7.5c; carry-over from
    slice #5 verify-report SUGGESTION a). Returns `None` when no template has
    been ingested for this organism — the packager then falls back to the
    default numbered folder structure below, exactly as the slice #2 stub
    always did (zero regression when nothing was ingested).
    """
    from app.services import requisito_inference_service

    return await requisito_inference_service.obtener_plantilla_organismo_por_contrato(db, contrato.usuario_id, contrato)


@dataclass(frozen=True)
class ObligacionEstado:
    """LISTO/PENDIENTE state of a single obligación within the evidence package."""

    obligacion_id: uuid.UUID
    descripcion: str
    listo: bool


@dataclass(frozen=True)
class EstadoListoPendiente:
    cuenta_cobro_id: uuid.UUID
    obligaciones: list[ObligacionEstado]
    pendientes: int
    listo_para_radicar: bool


def _calcular_estado_obligaciones(
    obligaciones: list[Obligacion],
    actividades_por_ob: dict[uuid.UUID | None, list[Actividad]],
    cuenta_id: uuid.UUID | None = None,
    obligaciones_con_link: frozenset[uuid.UUID] = frozenset(),
) -> list[ObligacionEstado]:
    """An obligación is LISTO when at least one of its actividades in this cuenta
    carries at least one evidencia (uploaded file or external link), OR when it has
    a CONFIRMED `evidencia_obligacion` link (`obligaciones_con_link`), OR when at
    least one of its actividades carries a REAL (non-sentinel) justificación;
    PENDIENTE otherwise. Mirrors the human LEEME workflow this packager has always
    shown.

    The link-table branch is the root-cause fix for the classify-after-upload path
    (live-reproduced 2026-08-11): an evidencia uploaded before classification lands
    on the shared "sin clasificar" stub Actividad (obligacion_id=None); confirming
    the link later only writes `evidencia_obligacion` and never re-attaches the
    Evidencia row, so counting only `act.evidencias` left the obligación PENDIENTE
    forever — steps 4/7 and the modo="final" gate disagreed with cobertura and the
    checklist EVIDENCIAS row, which already count confirmed links
    (`cobertura_service.calcular_cobertura`, «link table as source of truth»).

    The justificación branch (H12, 2026-08-11) aligns this split with the
    modo="final" package gate, which already accepts justificación-only coverage
    (ground truth: real approved cuotas radicate with partial evidencia). The
    deterministic sentinels ("No he desarrollado esta obligación…") do NOT count —
    they mean explicit absence of work, not justificación-only coverage."""

    estados: list[ObligacionEstado] = []
    for ob in obligaciones:
        acts = actividades_por_ob.get(ob.id, [])
        listo = (
            any(act.evidencias for act in acts)
            or ob.id in obligaciones_con_link
            or any(_tiene_justificacion_real(act.justificacion_origen) for act in acts)
        )
        estados.append(ObligacionEstado(obligacion_id=ob.id, descripcion=ob.descripcion, listo=listo))
    return estados


def _tiene_justificacion_real(justificacion_origen: str | None) -> bool:
    """A justificación counts as real coverage only when it comes from the LLM or a
    manual edit — seed/sin_labores/unset never count. Shared by
    `_calcular_estado_obligaciones` and `checklist_service.construir_checklist_completo`
    so the two gates never diverge. Works for both the ORM `Actividad`
    (`act.justificacion_origen`) and the dict shape from `checklist_service`
    (`act.get("justificacion_origen")`)."""
    return justificacion_origen in ("llm", "manual")


async def _obligaciones_con_link_confirmado(db: AsyncSession, cuenta_id: uuid.UUID) -> frozenset[uuid.UUID]:
    """Obligación ids covered by a CONFIRMED evidencia link within this cuenta —
    same rule `cobertura_service.calcular_cobertura` applies (scoped through the
    cuenta's actividades so cross-cuenta evidencias never leak in)."""
    return frozenset(await _evidencias_por_link_confirmado(db, cuenta_id))


async def _evidencias_por_link_confirmado(db: AsyncSession, cuenta_id: uuid.UUID) -> dict[uuid.UUID, list[Evidencia]]:
    """CONFIRMED-linked Evidencia rows per obligación for this cuenta.

    Live-reproduced 2026-08-12 (second E2E run): an evidencia confirmed to an
    obligación after upload stays attached to the shared "sin clasificar" stub
    Actividad, so the packager — which places bytes per `act.evidencias` — shipped
    the obligation folder with only its LEEME and no evidence files, while the
    gates (H13) already counted the obligación as covered. This map lets the zip
    place link-confirmed files under their obligation folders. An evidencia linked
    to several obligaciones is deliberately packaged into each folder (mirrors the
    human bundle)."""
    result = await db.execute(
        select(EvidenciaObligacion.obligacion_id, Evidencia)
        .join(Evidencia, Evidencia.id == EvidenciaObligacion.evidencia_id)
        .join(Actividad, Actividad.id == Evidencia.actividad_id)
        .where(
            Actividad.cuenta_cobro_id == cuenta_id,
            EvidenciaObligacion.status == EstadoEnlace.CONFIRMED.value,
            EvidenciaObligacion.obligacion_id.is_not(None),
        )
    )
    por_ob: dict[uuid.UUID, list[Evidencia]] = {}
    for ob_id, ev in result.all():
        if ob_id is not None:
            por_ob.setdefault(ob_id, []).append(ev)
    return por_ob


async def _evidencias_con_archivo(db: AsyncSession, cuenta_id: uuid.UUID) -> list[Evidencia]:
    """Every Evidencia with real bytes (`storage_key` set) attached to this
    cuenta's actividades — regardless of obligación classification. Feeds the
    `sin_clasificar/` root folder (A2): the caller filters out whatever the
    per-obligación loop already packaged, so orphan uploads (the "sin
    clasificar" stub, `obligacion_id=None`) and evidencias whose only link is
    still PROPOSED (not CONFIRMED) still ship instead of being silently
    dropped from the zip."""
    result = await db.execute(
        select(Evidencia)
        .join(Actividad, Actividad.id == Evidencia.actividad_id)
        .where(Actividad.cuenta_cobro_id == cuenta_id, Evidencia.storage_key.is_not(None))
        .order_by(Evidencia.created_at.asc())
    )
    return list(result.scalars().all())


async def obtener_estado_listo_pendiente(
    db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID
) -> EstadoListoPendiente:
    """Read-only LISTO/PENDIENTE split — the same computation `generar_zip_evidencias`
    uses internally, without producing a package."""
    _cuenta, _contrato, _usuario, obligaciones, actividades = await _load_context(db, usuario_id, cuenta_id)
    actividades_por_ob: dict[uuid.UUID | None, list[Actividad]] = {}
    for act in actividades:
        actividades_por_ob.setdefault(act.obligacion_id, []).append(act)

    con_link = await _obligaciones_con_link_confirmado(db, cuenta_id)
    estados = _calcular_estado_obligaciones(obligaciones, actividades_por_ob, cuenta_id, con_link)
    pendientes = sum(1 for e in estados if not e.listo)
    return EstadoListoPendiente(
        cuenta_cobro_id=cuenta_id,
        obligaciones=estados,
        pendientes=pendientes,
        listo_para_radicar=pendientes == 0,
    )


async def _texto_para_scan(contenido: bytes, nombre_archivo: str) -> str | None:
    """Best-effort text recovery for the secret scan ONLY — the real bytes still go
    into the zip untouched. Tries a direct UTF-8 decode first (covers the common
    case: text documents, exported emails), then falls back to the same extraction
    ladder used elsewhere (`document_service.extraer_texto_documento`) for binary
    formats (PDF/DOCX/images), per design decision D2. Fails open on extraction
    error or when nothing readable is recovered — a single unreadable file must
    never sink the whole package; only ITS scan coverage is skipped, its real bytes
    are still packaged."""
    try:
        return contenido.decode("utf-8")
    except UnicodeDecodeError:
        pass
    try:
        texto, _avisos = await document_service.extraer_texto_documento(contenido, nombre_archivo)
    except Exception as exc:  # one bad file must not sink the whole scan
        await logger.awarning("zip_evidencias_scan_extract_failed", nombre_archivo=nombre_archivo, error=str(exc))
        return None
    return texto


async def _agregar_documentos_contrato(
    db: AsyncSession,
    cuenta: CuentaCobro,
    miembros_zip: list[tuple[str, bytes]],
    miembros_scan: list[tuple[str, bytes]],
    omitir_codigos: set[str] | frozenset[str] = frozenset(),
) -> None:
    """DOCUMENTOS_CONTRATO/ section of the paquete (clasificacion-documentos-secop, D5).

    - ``DOCUMENTOS_CONTRATO/{REQUISITO_CODIGO}/``: real bytes of docs linked via
      ``DocumentoRequisitoVinculo`` to the cuenta's contract-level checklist rows
      (uploads via storage, SECOP via ``url_descarga``).
    - ``DOCUMENTOS_CONTRATO/CONTEXTO/``: remaining context docs
      (``listar_documentos_contexto`` — OTROS/unmapped, unlinked).
    - ``omitir_codigos``: requisito codes excluded from the package — the caller
      passes the catalog's ``solo_primera_cuenta`` codes on a non-first cuota
      (differential package: those docs were already radicated with cuota 1).

    Same skip-on-error semantics as the evidencias section: one failing download
    never sinks the package. Every added member also joins ``miembros_scan`` so
    the mandatory secret scan covers it. Read-only: zero link-table writers.
    """
    from app.models.documento_cuenta_cobro import DocumentoCuentaCobro, DocumentoRequisitoVinculo
    from app.services import checklist_service

    storage_docs = _get_storage(settings.S3_BUCKET_DOCUMENTOS)

    async def _agregar_miembro(arcname: str, contenido: bytes, nombre: str) -> None:
        miembros_zip.append((arcname, contenido))
        texto_scan = await _texto_para_scan(contenido, nombre)
        if texto_scan:
            miembros_scan.append((arcname, texto_scan.encode("utf-8")))

    async def _bytes_upload(storage_key: str, nombre: str) -> bytes | None:
        try:
            return await storage_docs.download(storage_key)
        except Exception as exc:  # one missing file must not sink the whole package
            await logger.awarning(
                "zip_documentos_contrato_download_failed", storage_key=storage_key, nombre=nombre, error=str(exc)
            )
            return None

    async def _bytes_secop(url: str | None, nombre: str) -> bytes | None:
        if not url:
            return None
        try:
            return await checklist_service._descargar_secop_bytes(url)
        except Exception as exc:
            await logger.awarning("zip_documentos_contrato_secop_failed", url=url, nombre=nombre, error=str(exc))
            return None

    rows = (
        (
            await db.execute(
                select(DocumentoCuentaCobro)
                .options(
                    selectinload(DocumentoCuentaCobro.vinculos).selectinload(
                        DocumentoRequisitoVinculo.documento_fuente
                    ),
                    selectinload(DocumentoCuentaCobro.vinculos).selectinload(DocumentoRequisitoVinculo.secop_documento),
                )
                .where(DocumentoCuentaCobro.cuenta_cobro_id == cuenta.id)
            )
        )
        .scalars()
        .all()
    )
    # ponytail: a doc linked to several contract-level rows (e.g. RPC+CDP alias)
    # is fetched once per link; cache per-scan if checklist sizes ever grow.
    for fila in rows:
        codigo = fila.requisito_codigo
        if codigo is None or not checklist_service.es_nivel_contrato(codigo):
            continue
        if codigo in omitir_codigos:
            continue  # first-cuota-only doc on a later cuota — already radicated in cuota 1
        for v in fila.vinculos:
            if v.documento_fuente is not None:
                nombre = v.documento_fuente.nombre
                contenido = await _bytes_upload(v.documento_fuente.storage_key, nombre)
            elif v.secop_documento is not None:
                nombre = v.secop_documento.nombre_archivo or v.secop_documento.id_documento_secop
                contenido = await _bytes_secop(v.secop_documento.url_descarga, nombre)
            else:
                continue
            if contenido is None:
                continue
            await _agregar_miembro(
                f"DOCUMENTOS_CONTRATO/{codigo}/{_safe_dirname(nombre, max_len=150)}", contenido, nombre
            )

    for item in await checklist_service.listar_documentos_contexto(db, cuenta):
        nombre = item["nombre"]
        if item["fuente"] == "upload" and item.get("storage_key"):
            contenido = await _bytes_upload(item["storage_key"], nombre)
        else:
            contenido = await _bytes_secop(item.get("url_descarga"), nombre)
        if contenido is None:
            continue
        await _agregar_miembro(f"DOCUMENTOS_CONTRATO/CONTEXTO/{_safe_dirname(nombre, max_len=150)}", contenido, nombre)


async def _hay_plantilla_organismo(db: AsyncSession, contrato: Contrato, tipo_documento: str) -> bool:
    """Whether this organism's ingested plantilla actually satisfies its generator's
    real preconditions — not just "a plantilla row exists". Mirrors:
    - `generar_cuenta_cobro_docx` (tipo_documento="cuenta_cobro"): plantilla + clonable.
    - `generar_documento_soporte_xlsx` (tipo_documento="documento_soporte"): plantilla +
      clonable + formato xlsx + non-empty campos + fuente_documento_id.

    A plantilla that exists but fails these must be pre-check-skipped with the
    truthful "Sin plantilla..." aviso — not fall through to the generic
    "No se pudo generar" failure AVISO in `_generar_documentos_paquete`.
    """
    plantilla = await _resolver_layout_organismo(db, contrato, tipo_documento)
    if plantilla is None:
        return False
    estructura = plantilla.estructura_json or {}
    if not estructura.get("clonable"):
        return False
    if tipo_documento == "documento_soporte":
        return bool(plantilla.formato == "xlsx" and estructura.get("campos") and plantilla.fuente_documento_id)
    return True


async def _generar_documentos_paquete(
    db: AsyncSession, usuario_id: uuid.UUID, cuenta_id: uuid.UUID, contrato: Contrato
) -> tuple[list[tuple[str, bytes]], list[str]]:
    """Root-level generated documents of the radicated package, numbered like the
    real approved packages (1. CUENTA DE COBRO ... 4. DOCUMENTO SOPORTE).

    Reuses EXACTLY the generators behind the download endpoints. Clone-only
    documents (cuenta de cobro / documento soporte) only exist as the organism's
    ingested plantilla — when none is ingested (a NORMAL state, not a failure)
    they are skipped up front with a truthful "Sin plantilla" line pointing at
    the plantilla ingestion, not at a re-download that would hit the same error.
    Genuine generation errors (clone error, LLM failure, ...) stay FAIL-OPEN per
    document: they skip ONLY that file and yield the "No se pudo generar" AVISO.
    Packaging never fails because a generated document failed. Returns
    ``(miembros, avisos)``.
    """
    # (arcname, generator, clone-only plantilla tipo — None for docs with a built-from-scratch fallback)
    generadores = [
        ("1. CUENTA DE COBRO.docx", generar_cuenta_cobro_docx, "cuenta_cobro"),
        ("2. INFORME DE ACTIVIDADES.docx", generar_informe_actividades_docx, None),
        ("3. INFORME DE SUPERVISION.docx", generar_informe_supervision_docx, None),
        ("4. DOCUMENTO SOPORTE.xlsx", generar_documento_soporte_xlsx, "documento_soporte"),
    ]
    miembros: list[tuple[str, bytes]] = []
    avisos: list[str] = []
    for arcname, generador, tipo_clone_only in generadores:
        if tipo_clone_only is not None and not await _hay_plantilla_organismo(db, contrato, tipo_clone_only):
            avisos.append(
                f"Sin plantilla del organismo para {arcname} — "
                "ingresá la plantilla en el paso Formato para incluirlo en el paquete."
            )
            continue
        try:
            contenido, _filename = await generador(db, usuario_id, cuenta_id)
        except Exception as exc:  # one failed generated doc must not sink the package
            await logger.awarning(
                "zip_documento_generado_fallo", cuenta_id=str(cuenta_id), documento=arcname, error=str(exc)
            )
            avisos.append(f"AVISO: No se pudo generar {arcname} — descargalo por separado desde el paso Formato.")
            continue
        miembros.append((arcname, contenido))
    return miembros, avisos


def get_evidencia_storage() -> object:
    """Storage adapter for evidencia BYTES (photos, PDFs, etc. attached to an
    actividad) — same bucket the upload endpoints write to (`S3_BUCKET_PDFS`,
    see `api.deps.get_pdf_storage`). Every evidencia-bytes READ must go through
    this helper so the read bucket can never silently drift from the upload
    bucket again: it previously pointed at `S3_BUCKET_EVIDENCIAS`, which the
    upload endpoints never write to, so every uploaded evidencia was fetched
    with a wrong key prefix, failed with `FileNotFoundError`, and was dropped
    from the generated ZIP by the fail-open `except` below — silently, since
    that fallback only logs a warning.

    `S3_BUCKET_EVIDENCIAS` remains correct for the PACKAGE ARTIFACT itself
    (the generated zip stored under `paquetes/...` — see
    `radicacion_prep_service.preparar_radicacion`, `paquete_service`,
    `tools/catalog/paquete.py`). Do not repoint those; only evidencia-bytes
    reads belong behind this helper.
    """
    return _get_storage(settings.S3_BUCKET_PDFS)


async def generar_zip_evidencias(
    db: AsyncSession,
    usuario_id: uuid.UUID,
    cuenta_id: uuid.UUID,
    modo: Literal["standard", "final"] = "standard",
) -> tuple[bytes, str]:
    """Build a ZIP with real evidence bytes in numbered per-obligación folders (the
    default fallback structure — per-organism structures arrive in slice #5).

    Runs a mandatory fail-closed secret scan before emitting the zip
    (`SECRET_SCAN_GATE_ENABLED`, default True) and computes a LISTO/PENDIENTE split
    mirroring the human LEEME workflow:

    - ``modo="standard"`` (default, manual/tool-driven packaging): emits the
      package even with PENDIENTE obligaciones, listing them in the manifest.
    - ``modo="final"`` (used by the slice #7 radicación orchestrator): raises
      `PACKAGE_PENDIENTE` instead of emitting a partial package.
    """
    cuenta, contrato, _usuario, obligaciones, actividades = await _load_context(db, usuario_id, cuenta_id)

    actividades_por_ob: dict[uuid.UUID | None, list[Actividad]] = {}
    for act in actividades:
        actividades_por_ob.setdefault(act.obligacion_id, []).append(act)

    estructura_organismo = await _resolver_estructura_organismo(db, contrato)
    # COEMPRESAR-style organisms number their evidence folders "A{n}_..." (their
    # institutional template literally references "Carpeta ... A1" anexos) — when
    # an ingested template for this organism carries anexo references, mirror that
    # convention instead of the plain "01_..." default. Column layout / section
    # order from `estructura_json` are consumed by `adaptive-informe-generation`
    # (slice #6) for the DOCX informes themselves, not by this folder builder.
    usa_folders_anexo = bool(estructura_organismo and estructura_organismo.estructura_json.get("anexo_refs"))

    evidencias_link_por_ob = await _evidencias_por_link_confirmado(db, cuenta_id)
    con_link = frozenset(evidencias_link_por_ob)
    estados = _calcular_estado_obligaciones(obligaciones, actividades_por_ob, cuenta_id, con_link)
    pendientes_desc = [e.descripcion for e in estados if not e.listo]

    # Ground truth (real approved SYJ cuotas): partial evidencia coverage is the
    # NORM, not a defect — an obligación with no evidencia files still radicates
    # fine as long as SOME actividad carries a real justificación (the sentinel
    # or user-written text). The LEEME LISTO/PENDIENTE split above stays
    # evidence-based (unchanged); this separate split decides the modo="final"
    # gate below and only blocks obligaciones with NEITHER evidencia NOR
    # justificación.
    def _tiene_justificacion(ob_id: uuid.UUID | None) -> bool:
        return any((act.justificacion or "").strip() for act in actividades_por_ob.get(ob_id, []))

    justificadas_sin_evidencia_desc = [
        e.descripcion for e in estados if not e.listo and _tiene_justificacion(e.obligacion_id)
    ]
    sin_evidencia_ni_justificacion_desc = [
        e.descripcion for e in estados if not e.listo and not _tiene_justificacion(e.obligacion_id)
    ]

    # Generated documents at the package root (ground truth: the real radicated
    # bundle carries them alongside evidencias + contract docs). Generated BEFORE
    # the root LEEME so it can report which are included / pending.
    docs_generados, avisos_generados = await _generar_documentos_paquete(db, usuario_id, cuenta_id, contrato)

    # Differential package: on a non-first cuota, first-cuota-only contract docs
    # (catalog rows with solo_primera_cuenta) were already radicated with cuota 1.
    # "First" reuses the checklist's own persisted-posicion predicate.
    from app.services import checklist_service

    codigos_solo_primera: set[str] = set()
    if not checklist_service._is_first_cuenta(cuenta):
        catalogo = await checklist_service.listar_catalogo(db)
        codigos_solo_primera = {req.codigo for req in catalogo if req.solo_primera_cuenta}

    storage = get_evidencia_storage()

    miembros_zip: list[tuple[str, bytes]] = []
    miembros_scan: list[tuple[str, bytes]] = []

    # A4: visibility counters for the fail-open evidence downloads below — surfaced
    # in the zip_evidencias_generado log and, when non-zero, as an AVISO line in
    # the root LEEME (assembled further down, once these are known).
    evidencias_empacadas = 0
    evidencias_fallidas = 0
    # Evidencia ids already placed inside an obligación folder — the sin_clasificar
    # pass (A2) below skips these so nothing ships twice.
    ids_empacadas: set[uuid.UUID] = set()

    def _agregar_texto(nombre: str, texto: str) -> None:
        contenido = texto.encode("utf-8")
        miembros_zip.append((nombre, contenido))
        miembros_scan.append((nombre, contenido))

    root_readme_lines = [
        "Carpeta de evidencias",
        "=====================",
        "",
        f"Contrato: {contrato.numero_contrato}",
        f"Objeto: {contrato.objeto}",
        f"Período: {_periodo_str(cuenta)}",
        f"Obligaciones: {len(obligaciones)}",
        f"Actividades reportadas: {len(actividades)}",
        "",
        "Estructura:",
        "  - Una subcarpeta por obligación contractual, con los archivos de",
        "    evidencia ya clasificados para esa obligación (sin README individual).",
        "  - Una obligación sin evidencia clasificada no genera subcarpeta.",
        "  - sin_clasificar/: evidencias con archivo que aún no quedaron",
        "    confirmadas contra ninguna obligación específica.",
        "",
        "Estado de completitud (LISTO/PENDIENTE)",
        "========================================",
        "",
        f"Obligaciones con evidencia (LISTO): {len(obligaciones) - len(pendientes_desc)}/{len(obligaciones)}",
    ]
    if pendientes_desc:
        root_readme_lines.append("Obligaciones SIN evidencia (PENDIENTE):")
        root_readme_lines += [f"  - {d}" for d in pendientes_desc]
    else:
        root_readme_lines.append("Obligaciones SIN evidencia (PENDIENTE): (ninguna)")
    if justificadas_sin_evidencia_desc:
        # Clarifies the split above: these PENDIENTE-by-evidencia obligaciones
        # are still radicables — they carry a justificación (sentinel or texto
        # propio) that satisfies the modo="final" gate on its own.
        root_readme_lines.append("  (de las anteriores, justificadas sin evidencia — no bloquean la radicación):")
        root_readme_lines += [f"    - {d}" for d in justificadas_sin_evidencia_desc]
    root_readme_lines += [
        "",
        "Documentos generados (raíz del paquete)",
        "========================================",
        "",
    ]
    if docs_generados:
        root_readme_lines += [f"  - {arcname}" for arcname, _contenido in docs_generados]
    else:
        root_readme_lines.append("  (ninguno)")
    root_readme_lines += avisos_generados
    if codigos_solo_primera:
        root_readme_lines += [
            "",
            "Documentos de primera cuota omitidos: ya radicados en la cuota 1 ("
            + ", ".join(sorted(codigos_solo_primera))
            + ").",
        ]
    # NOTE: root LEEME.txt is written further down (after the obligación +
    # sin_clasificar packaging loops), once evidencias_fallidas is known — see the
    # AVISO line appended there (A4).

    for arcname, contenido in docs_generados:
        miembros_zip.append((arcname, contenido))
        texto_scan = await _texto_para_scan(contenido, arcname)
        if texto_scan:
            miembros_scan.append((arcname, texto_scan.encode("utf-8")))

    if not obligaciones:
        _agregar_texto(
            "00_sin_obligaciones/LEEME.txt",
            "El contrato no tiene obligaciones registradas. Cargue el contrato y extraiga obligaciones primero.\n",
        )

    for idx, ob in enumerate(obligaciones, start=1):
        folder_prefix = f"A{idx}" if usa_folders_anexo else f"{idx:02d}"
        folder = f"{folder_prefix}_{_safe_dirname(ob.descripcion)}"
        acts = actividades_por_ob.get(ob.id, [])
        # Evidencias confirmed to this obligación via the link table but attached
        # to another actividad (typically the "sin clasificar" upload stub) —
        # without this branch their bytes never shipped, leaving the obligation
        # folder LEEME-only while the gates reported the obligación covered
        # (live-reproduced 2026-08-12, H13 packaging counterpart).
        ids_act = {ev.id for act in acts for ev in act.evidencias}
        evs_link = [ev for ev in evidencias_link_por_ob.get(ob.id, []) if ev.id not in ids_act]

        # Actividad-attached evidencias plus the link-confirmed ones (dedup'd
        # above) — both classes ship real bytes into this obligation's folder.
        # No LEEME here anymore (A3): a folder is ONLY the real files, and an
        # obligación with none of them simply produces no folder at all.
        for ev in [ev for act in acts for ev in act.evidencias] + evs_link:
            ids_empacadas.add(ev.id)
            if ev.storage_key is None:
                continue  # external-link evidence — no real bytes to fetch (no new StoragePort methods this change)
            try:
                contenido = await storage.download(ev.storage_key)
            except Exception as exc:  # one missing file must not sink the whole package
                await logger.awarning(
                    "zip_evidencias_download_failed",
                    evidencia_id=str(ev.id),
                    storage_key=ev.storage_key,
                    error=str(exc),
                )
                # Diagnostic instrumentation (paquete-regenerar-evidencia-nueva,
                # Slice 1) — discriminates "read-but-dropped" (this fail-open
                # branch, storage-durability race) from "never loaded" (pt1
                # shows no such evidencia at all).
                await logger.awarning(
                    "zip_evidencia_download",
                    cuenta_id=str(cuenta_id),
                    evidencia_id=str(ev.id),
                    storage_key=ev.storage_key,
                    download_ok=False,
                    bytes_len=0,
                )
                evidencias_fallidas += 1
                continue
            await logger.ainfo(
                "zip_evidencia_download",
                cuenta_id=str(cuenta_id),
                evidencia_id=str(ev.id),
                storage_key=ev.storage_key,
                download_ok=True,
                bytes_len=len(contenido),
            )
            evidencias_empacadas += 1
            arcname = f"{folder}/evidencias/{_safe_dirname(ev.nombre_archivo, max_len=150)}"
            miembros_zip.append((arcname, contenido))
            texto_scan = await _texto_para_scan(contenido, ev.nombre_archivo)
            if texto_scan:
                miembros_scan.append((arcname, texto_scan.encode("utf-8")))

    # A2: every evidencia with real bytes not already placed in an obligación
    # folder above — orphan uploads ("sin clasificar" stub, obligacion_id=None)
    # and evidencias whose only link is still PROPOSED (not CONFIRMED) — ships
    # under a root sin_clasificar/ folder instead of being silently dropped.
    # Supersedes the old 99_otras_actividades/LEEME.txt block (A3): that block
    # never shipped bytes anyway, only a filename-less activity listing.
    for ev in await _evidencias_con_archivo(db, cuenta_id):
        if ev.id in ids_empacadas:
            continue
        try:
            contenido = await storage.download(ev.storage_key)
        except Exception as exc:  # one missing file must not sink the whole package
            await logger.awarning(
                "zip_evidencia_sin_clasificar_download_failed",
                evidencia_id=str(ev.id),
                storage_key=ev.storage_key,
                error=str(exc),
            )
            evidencias_fallidas += 1
            continue
        evidencias_empacadas += 1
        arcname = f"sin_clasificar/{_safe_dirname(ev.nombre_archivo, max_len=150)}"
        miembros_zip.append((arcname, contenido))
        texto_scan = await _texto_para_scan(contenido, ev.nombre_archivo)
        if texto_scan:
            miembros_scan.append((arcname, texto_scan.encode("utf-8")))

    # A4: root LEEME AVISO when some evidencia couldn't be packaged, then write it —
    # deferred from earlier in this function so evidencias_fallidas is known.
    if evidencias_fallidas:
        root_readme_lines += [
            "",
            f"AVISO: {evidencias_fallidas} evidencia(s) no se pudieron incluir en este paquete "
            "(fallo al descargar el archivo desde almacenamiento) — reintente más tarde o "
            "contacte soporte si persiste.",
        ]
    root_readme_lines.append("")
    _agregar_texto("LEEME.txt", "\n".join(root_readme_lines))

    # Contract-level docs ship classified in the paquete (D5) — before the scan
    # gate so every new member is covered by it.
    await _agregar_documentos_contrato(db, cuenta, miembros_zip, miembros_scan, omitir_codigos=codigos_solo_primera)

    if settings.SECRET_SCAN_GATE_ENABLED:
        hallazgos = await secret_scan_service.escanear_paquete(miembros_scan)
        if hallazgos:
            archivos = sorted({h.archivo for h in hallazgos})
            await logger.aerror(
                "zip_evidencias_secreto_detectado",
                cuenta_id=str(cuenta_id),
                archivos=archivos,
                tipos=sorted({h.tipo for h in hallazgos}),
            )
            raise ValidationError(
                "Se detectaron posibles credenciales/secretos dentro del paquete de evidencias — "
                "no se generó el zip. Revise los archivos: " + ", ".join(archivos),
                code=SECRET_DETECTED_IN_PACKAGE,
            )

    if modo == "final" and sin_evidencia_ni_justificacion_desc:
        raise ValidationError(
            f"No se puede finalizar la radicación: {len(sin_evidencia_ni_justificacion_desc)} "
            "obligación(es) sin evidencia ni justificación.",
            code=PACKAGE_PENDIENTE,
        )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for arcname, contenido in miembros_zip:
            zf.writestr(arcname, contenido)

    filename = f"evidencias-{contrato.numero_contrato}-{cuenta.anio}-{cuenta.mes:02d}.zip"
    await logger.ainfo(
        "zip_evidencias_generado",
        cuenta_id=str(cuenta_id),
        usuario_id=str(usuario_id),
        size=len(buf.getvalue()),
        modo=modo,
        pendientes=len(pendientes_desc),
        evidencias_empacadas=evidencias_empacadas,
        evidencias_fallidas=evidencias_fallidas,
    )
    return buf.getvalue(), filename
