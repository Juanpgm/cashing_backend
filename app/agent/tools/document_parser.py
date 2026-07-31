"""Document parser tool — extract text from PDF, DOCX, XLSX, archives, and plain text.

The agent chat lets users drop arbitrary files (including compressed archives). The
parser therefore tries, in order:

1. A format-specific extractor for known rich formats (PDF/DOCX/XLSX).
2. Archive expansion for `.zip`/`.tar`/`.tar.gz`/`.tgz` — each member is extracted
   recursively (guarded against zip bombs by member count / per-member size / total
   size caps and by skipping executables).
3. A best-effort UTF-8/Latin-1 text decode for anything else that looks like text
   (`.txt`, `.csv`, `.md`, `.json`, source code, etc.).

`parse_document` raises `ValueError` only when the bytes are genuinely non-textual and
not a supported archive, so callers that rely on that signal (the upload/extraction
pipeline) keep their contract.
"""

from __future__ import annotations

import io
import tarfile
import zipfile
from collections.abc import Iterator
from pathlib import Path

import structlog

logger = structlog.get_logger("agent.tools.document_parser")

# Archive-expansion safety caps (defence against zip bombs and runaway extraction).
_ARCHIVE_MAX_MEMBERS = 50
_ARCHIVE_MAX_MEMBER_BYTES = 25 * 1024 * 1024  # 25 MB uncompressed per member
_ARCHIVE_MAX_TOTAL_BYTES = 100 * 1024 * 1024  # 100 MB uncompressed total
_ARCHIVE_MAX_TEXT_CHARS = 200_000  # cap the concatenated text a single archive yields

# Extensions we never try to read as text (binaries / executables / scripts).
_NON_TEXT_EXTENSIONS = {
    ".exe",
    ".dll",
    ".so",
    ".dylib",
    ".bin",
    ".msi",
    ".apk",
    ".jar",
    ".bat",
    ".cmd",
    ".com",
    ".scr",
    ".ps1",
    ".sh",
    ".app",
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".bmp",
    ".tiff",
    ".webp",
    ".ico",
    ".mp3",
    ".mp4",
    ".mov",
    ".avi",
    ".mkv",
    ".wav",
    ".flac",
}

_ARCHIVE_EXTENSIONS = {".zip", ".tar", ".gz", ".tgz", ".rar"}


def parse_pdf(content: bytes) -> str:
    """Extract text from PDF bytes using pdfplumber."""
    import pdfplumber

    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _docx_table_text(table) -> list[str]:
    """One line per table row (cells joined with ' | '), recursing into nested tables."""
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
        for cell in row.cells:
            for nested in cell.tables:
                lines.extend(_docx_table_text(nested))
    return lines


def parse_docx(content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx (body paragraphs + tables)."""
    import docx

    doc = docx.Document(io.BytesIO(content))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        parts.extend(_docx_table_text(table))
    return "\n".join(parts)


def parse_xlsx(content: bytes) -> str:
    """Extract text from XLSX bytes using openpyxl."""
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    text_parts: list[str] = []
    for ws in wb.worksheets:
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            text_parts.append(f"[{ws.title}]\n" + "\n".join(rows))
    wb.close()
    return "\n\n".join(text_parts)


def parse_text(content: bytes) -> str:
    """Best-effort decode of arbitrary bytes as text.

    Raises ValueError if the bytes don't look like text (too many undecodable bytes),
    so archives/binaries don't masquerade as empty documents.
    """
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            decoded = content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
        # Reject if the decode produced too many replacement/null chars — a sign the
        # bytes were binary, not text.
        if decoded and decoded.count("\x00") / max(len(decoded), 1) < 0.01:
            return decoded
    msg = "Content does not appear to be decodable text"
    raise ValueError(msg)


def _extract_member(name: str, data: bytes) -> str:
    """Extract text from a single archive member, dispatching by its extension.

    Returns "" for members that can't be read (skipped binaries, parse failures) so a
    single bad file never aborts the whole archive.
    """
    ext = Path(name).suffix.lower()
    if ext in _NON_TEXT_EXTENSIONS:
        return ""
    try:
        if ext == ".pdf":
            return parse_pdf(data)
        if ext == ".docx":
            return parse_docx(data)
        if ext in (".xlsx", ".xls"):
            return parse_xlsx(data)
        if ext in _ARCHIVE_EXTENSIONS:
            # Nested archive: expand one level; deeper nesting is uncommon and the
            # total-bytes cap still bounds it.
            return parse_archive(data, name)
        return parse_text(data)
    except Exception as exc:  # noqa: BLE001 — one bad member must not sink the archive
        logger.warning("archive_member_parse_failed", member=name, error=str(exc))
        return ""


def _iter_zip_members(content: bytes):
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            if info.file_size > _ARCHIVE_MAX_MEMBER_BYTES:
                logger.warning("archive_member_too_large", member=info.filename, size=info.file_size)
                continue
            yield info.filename, zf.read(info)


def _iter_rar_members(content: bytes):
    """Yield members of a .rar archive via `rarfile`.

    `rarfile` needs an external unrar/bsdtar binary at runtime (bsdtar ships in the
    Docker image via libarchive-tools). Without one — e.g. a bare local dev setup —
    extraction degrades gracefully: we log and yield nothing, the upload still
    succeeds with only the filename as signal.
    """
    import rarfile

    try:
        rf = rarfile.RarFile(io.BytesIO(content))
    except Exception as exc:  # corrupt rar / missing backend
        logger.warning("rar_open_failed", error=str(exc))
        return
    with rf:
        for info in rf.infolist():
            if info.is_dir():
                continue
            if info.file_size > _ARCHIVE_MAX_MEMBER_BYTES:
                logger.warning("archive_member_too_large", member=info.filename, size=info.file_size)
                continue
            try:
                yield info.filename, rf.read(info)
            except Exception as exc:  # missing unrar tool, bad member
                logger.warning("rar_member_read_failed", member=info.filename, error=str(exc))
                return


def _iter_tar_members(content: bytes):
    with tarfile.open(fileobj=io.BytesIO(content)) as tf:
        for member in tf.getmembers():
            if not member.isfile():
                continue
            if member.size > _ARCHIVE_MAX_MEMBER_BYTES:
                logger.warning("archive_member_too_large", member=member.name, size=member.size)
                continue
            extracted = tf.extractfile(member)
            if extracted is None:
                continue
            yield member.name, extracted.read()


def _archive_member_iter(content: bytes, filename: str):
    """Pick the right member iterator (rar/tar/zip) for an archive filename."""
    ext = Path(filename).suffix.lower()
    if ext == ".rar":
        return _iter_rar_members(content)
    if ext in (".tar", ".gz", ".tgz") or filename.lower().endswith(".tar.gz"):
        return _iter_tar_members(content)
    return _iter_zip_members(content)


def is_archive_filename(filename: str) -> bool:
    """True if `filename`'s extension marks it as a supported archive (zip/tar/gz/tgz/rar).

    Shared by `parse_document` (text-preview expansion) and
    `agent_chat_service._expand_attachments_for_tools` (making archive members
    individually importable), so both agree on exactly which uploads are archives.
    """
    ext = Path(filename).suffix.lower()
    return ext in _ARCHIVE_EXTENSIONS or filename.lower().endswith(".tar.gz")


def iter_archive_members(content: bytes, filename: str) -> Iterator[tuple[str, bytes]]:
    """Yield `(member_path, member_bytes)` for each member of a zip/tar/gz archive.

    Reuses `_iter_zip_members`/`_iter_tar_members` — directories and members over
    `_ARCHIVE_MAX_MEMBER_BYTES` are skipped there already. Bounded here by BOTH
    `_ARCHIVE_MAX_MEMBERS` total members AND `_ARCHIVE_MAX_TOTAL_BYTES` total
    decompressed bytes — the same two caps `parse_archive` enforces (mirrored
    exactly, same accumulate-then-check order) — so a caller (e.g. the agent
    chat's attachment expansion, or the plantillas-organismo archive endpoint)
    can't be flooded by a zip-bomb-style archive with a huge member count OR a
    small archive that expands to unbounded memory.
    """
    member_iter = _archive_member_iter(content, filename)

    count = 0
    total_bytes = 0
    for name, data in member_iter:
        if count >= _ARCHIVE_MAX_MEMBERS:
            break
        total_bytes += len(data)
        if total_bytes > _ARCHIVE_MAX_TOTAL_BYTES:
            break
        count += 1
        yield name, data


def parse_archive(content: bytes, filename: str) -> str:
    """Expand a zip/tar/gz/rar archive and concatenate the text of its members.

    Each member is prefixed with a `### <path>` header so the LLM can tell files apart.
    Bounded by member-count, per-member-size, total-size, and total-text caps.
    """
    member_iter = _archive_member_iter(content, filename)

    parts: list[str] = []
    total_bytes = 0
    count = 0
    for name, data in member_iter:
        if count >= _ARCHIVE_MAX_MEMBERS:
            parts.append(f"\n(se omitieron miembros adicionales; límite de {_ARCHIVE_MAX_MEMBERS} archivos)")
            break
        total_bytes += len(data)
        if total_bytes > _ARCHIVE_MAX_TOTAL_BYTES:
            parts.append("\n(se detuvo la extracción; el contenido descomprimido excede el límite permitido)")
            break
        count += 1
        text = _extract_member(name, data)
        if text.strip():
            parts.append(f"### {name}\n{text.strip()}")

    combined = "\n\n".join(parts)
    if len(combined) > _ARCHIVE_MAX_TEXT_CHARS:
        combined = combined[:_ARCHIVE_MAX_TEXT_CHARS] + "\n... (contenido truncado)"
    return combined


def parse_document(content: bytes, filename: str) -> str:
    """Auto-detect format and extract text.

    Supports rich formats (PDF/DOCX/XLSX), archives (zip/tar/gz/rar — contents
    expanded recursively), and a plain-text fallback for any other decodable format.
    Raises ValueError only for genuinely non-textual binaries.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(content)
    if ext == ".docx":
        return parse_docx(content)
    if ext in (".xlsx", ".xls"):
        return parse_xlsx(content)
    if is_archive_filename(filename):
        return parse_archive(content, filename)
    if ext in _NON_TEXT_EXTENSIONS:
        msg = f"Unsupported binary file format: {ext}"
        raise ValueError(msg)
    # Unknown extension: try a best-effort text decode (covers .txt/.csv/.md/.json/code).
    return parse_text(content)
