"""Template filler tool — fill Jinja2 DOCX/HTML templates with data."""

from __future__ import annotations

import io

from jinja2 import BaseLoader, Environment


def fill_template(template_text: str, data: dict[str, str | int | float | None]) -> str:
    """Render a Jinja2 text template with the given data."""
    env = Environment(loader=BaseLoader(), autoescape=True)
    tmpl = env.from_string(template_text)
    return tmpl.render(**data)


def fill_docx_template(template_bytes: bytes, data: dict[str, str | int | float | None]) -> bytes:
    """Fill placeholders in a DOCX template and return the result as bytes."""
    from docx import Document

    doc = Document(io.BytesIO(template_bytes))
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value or ""))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder = "{{" + key + "}}"
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value or ""))

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
